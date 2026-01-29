import os, html2text, html
from typing import Union
import markdown, pypandoc, shutil
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def readTextFile(textFile: str) -> Union[str, None]:
    if not os.path.isfile(textFile):
        return None
    with open(textFile, 'r', encoding='utf8') as fileObj:
        content = fileObj.read()
    return content if content else ""

def writeTextFile(textFile: str, textContent: str) -> None:
    with open(textFile, "w", encoding="utf-8") as fileObj:
        fileObj.write(textContent)

# Function to convert HTML to Markdown
def htmlToMarkdown(html_string):
    # Create an instance of the HTML2Text converter
    converter = html2text.HTML2Text()
    # Convert the HTML string to Markdown
    markdown_string = converter.handle(html_string)
    # Return the Markdown string
    return markdown_string

def plainTextToUrl(text):
    # https://wiki.python.org/moin/EscapingHtml
    text = html.escape(text)
    searchReplace = (
        (" ", "%20"),
        ("\n", "%0D%0A"),
    )
    for search, replace in searchReplace:
        text = text.replace(search, replace)
    return text

def markdownToHtml(markdown_text, output_file):
    if shutil.which("pandoc"):
        pypandoc.convert_text(markdown_text, "html", format="md", outputfile=output_file)
    else:
        writeTextFile(output_file, markdown.markdown(markdown_text))

def markdownToDocx(markdown_text, output_file):
    if shutil.which("pandoc"):
        pypandoc.convert_text(markdown_text, "docx", format="md", outputfile=output_file)
        return None

    # Parse the markdown text
    html = markdown.markdown(markdown_text)

    # Create a new docx document
    document = Document()

    # Add the parsed markdown text to the document
    for element in html.split('\n'):
        if element.startswith('<h1>'):
            # Add a heading
            heading = document.add_heading(text=element.replace('<h1>', '').replace('</h1>', ''), level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif element.startswith('<h2>'):
            # Add a subheading
            subheading = document.add_heading(text=element.replace('<h2>', '').replace('</h2>', ''), level=2)
            subheading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif element.startswith('<p>'):
            # Add a paragraph
            paragraph = document.add_paragraph()
            paragraph.add_run(text=element.replace('<p>', '').replace('</p>', ''))
        elif element.startswith('<ul>'):
            # Add an unordered list
            ul = document.add_paragraph()
            ul.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for item in element.replace('<ul>', '').replace('</ul>', '').split('<li>'):
                if item:
                    ul.add_run(text=item.replace('</li>', '') + '\n')
        elif element.startswith('<ol>'):
            # Add an ordered list
            ol = document.add_paragraph()
            ol.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for i, item in enumerate(element.replace('<ol>', '').replace('</ol>', '').split('<li>')):
                if item:
                    ol.add_run(text=str(i+1) + '. ' + item.replace('</li>', '') + '\n')

    # Save the document
    document.save(output_file)

def set_log_file_max_lines(log_file, max_lines):
    if os.path.isfile(log_file):
        # Read the contents of the log file
        with open(log_file, "r", encoding="utf-8") as fileObj:
            lines = fileObj.readlines()
        # Count the number of lines in the file
        num_lines = len(lines)
        if num_lines > max_lines:
            # Calculate the number of lines to be deleted
            num_lines_to_delete = num_lines - max_lines
            if num_lines_to_delete > 0:
                # Open the log file in write mode and truncate it
                with open(log_file, "w", encoding="utf-8") as fileObj:
                    # Write the remaining lines back to the log file
                    fileObj.writelines(lines[num_lines_to_delete:])
            filename = os.path.basename(log_file)
            print(f"{num_lines_to_delete} old lines deleted from log file '{filename}'.")