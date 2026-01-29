import unittest
from unittest.mock import patch
from bs4 import BeautifulSoup
from docx import Document

from htmltodocx.config import base_config
from htmltodocx.htmltodocx import HTMLtoDocx


class TestHTMLtoDocx(unittest.TestCase):
    def setUp(self):
        self.doc = Document()
        self.converter = HTMLtoDocx(document=self.doc, config=base_config)

    def test_parse_html_empty(self):
        """Тест: пустой HTML"""
        html = ""
        self.converter.parse_html(html)
        self.assertEqual(len(self.doc.paragraphs), 0)

    def test_parse_html_single_paragraph(self):
        """Тест: один абзац <p>"""
        html = "<p>Простой текст</p>"
        self.converter.parse_html(html)
        self.assertEqual(len(self.doc.paragraphs), 1)
        self.assertEqual(self.doc.paragraphs[0].text, "Простой текст")

    def test_handle_paragraph_with_indent(self):
        """Тест: абзац с отступом"""
        paragraph = self.doc.add_paragraph()
        tag = BeautifulSoup("<p>Текст</p>", "html.parser").p
        self.converter.handle_paragraph(tag, paragraph, with_indent=True)
        self.assertEqual(paragraph.paragraph_format.first_line_indent.cm, 1.2505972222222221)

    def test_handle_paragraph_with_message_start(self):
        """Тест: первый абзац с префиксом сообщения"""
        self.converter.message_start = "Сообщение: «"
        tag = BeautifulSoup("<p>Текст</p>", "html.parser").p
        self.converter.handle_paragraph(tag)
        self.assertTrue(self.doc.paragraphs[0].text.startswith("Сообщение: «"))
        self.assertTrue(self.doc.paragraphs[0].runs[0].underline)

    def test_handle_list_unordered(self):
        """Тест: неупорядоченный список"""
        html = "<ul><li>Пункт 1</li><li>Пункт 2</li></ul>"
        soup = BeautifulSoup(html, "html.parser")
        self.converter.handle_tag(soup.ul)
        self.assertEqual(len(self.doc.paragraphs), 2)
        self.assertIn("Пункт 1", self.doc.paragraphs[0].text)
        self.assertIn("Пункт 2", self.doc.paragraphs[1].text)

    def test_handle_list_ordered(self):
        """Тест: упорядоченный список"""
        html = "<ol><li>Первый</li><li>Второй</li></ol>"
        soup = BeautifulSoup(html, "html.parser")
        self.converter.handle_tag(soup.ol)
        self.assertEqual(len(self.doc.paragraphs), 2)
        self.assertIn("Первый", self.doc.paragraphs[0].text)
        self.assertIn("Второй", self.doc.paragraphs[1].text)

    def test_add_text_and_style_simple(self):
        """Тест: добавление текста с тегами стилей"""
        paragraph = self.doc.add_paragraph()
        html = "<p><strong>Жирный</strong> <em>курсив</em></p>"
        soup = BeautifulSoup(html, "html.parser")
        self.converter.add_text_and_style(soup.p.children, paragraph)
        runs = paragraph.runs
        self.assertEqual(len(runs), 3)  # "Жирный", " ", "курсив"
        self.assertTrue(runs[0].bold)
        self.assertTrue(runs[2].italic)

    def test_parse_styled_text_with_nested_styles(self):
        """Тест: вложенные стили (strong + em + u)"""
        paragraph = self.doc.add_paragraph()
        html = "<p><i><strong><u>Текст</u></strong></i></p>"
        soup = BeautifulSoup(html, "html.parser")
        self.converter._parse_styled_text(soup.p, paragraph)
        runs = paragraph.runs
        self.assertEqual(len(runs), 1)
        self.assertTrue(runs[0].italic)
        self.assertTrue(runs[0].bold)
        self.assertTrue(runs[0].underline)

    def test_handle_table_simple(self):
        """Тест: простая таблица 2x2"""
        html = """
        <table>
            <tbody>
                <tr><td>1</td><td>2</td></tr>
                <tr><td>3</td><td>4</td></tr>
            </tbody>
        </table>
        """
        soup = BeautifulSoup(html, "html.parser")
        self.converter.handle_tag(soup.table)
        table = self.doc.tables[0]
        self.assertEqual(table.cell(0, 0).text, "1")
        self.assertEqual(table.cell(0, 1).text, "2")
        self.assertEqual(table.cell(1, 0).text, "3")
        self.assertEqual(table.cell(1, 1).text, "4")

    def test_handle_table_with_colspan(self):
        """Тест: таблица с colspan"""
        html = """
        <table>
            <tbody>
                <tr><td colspan="2">Объединено</td></tr>
                <tr><td>1</td><td>2</td></tr>
            </tbody>
        </table>
        """
        soup = BeautifulSoup(html, "html.parser")
        self.converter.handle_tag(soup.table)
        table = self.doc.tables[0]
        cell = table.cell(0, 0)
        self.assertEqual(cell.text, "Объединено")

    def test_handle_table_with_rowspan(self):
        """Тест: таблица с rowspan"""
        html = """
        <table>
            <tbody>
                <tr><td rowspan="2">Вертикаль</td><td>1</td></tr>
                <tr><td>2</td></tr>
            </tbody>
        </table>
        """
        soup = BeautifulSoup(html, "html.parser")
        self.converter.handle_tag(soup.table)
        table = self.doc.tables[0]
        cell = table.cell(0, 0)
        self.assertEqual(cell.text, "Вертикаль")

    def test_handle_image(self):
        """Тест: вставка изображения"""
        paragraph = self.doc.add_paragraph()
        with patch("htmltodocx.add_image_to_docx") as mock_add_image:
            html = '<figure><img src="image.png"></figure>'
            soup = BeautifulSoup(html, "html.parser")
            self.converter.handle_image(soup.figure, paragraph)
            mock_add_image.assert_called_once()

    def test_handle_tag_figure_image(self):
        """Тест: обработка figure как изображения"""
        with patch.object(self.converter, "handle_image") as mock_handle:
            tag = BeautifulSoup("<figure></figure>", "html.parser").figure
            self.converter.handle_tag(tag)
            mock_handle.assert_called_once()

    def test_handle_tag_unknown_with_text(self):
        """Тест: неизвестный тег с текстом"""
        paragraph = self.doc.add_paragraph()
        tag = BeautifulSoup("<div>Текст</div>", "html.parser").div
        with patch.object(self.converter, "add_text_and_style") as mock_add:
            self.converter.handle_tag(tag, paragraph)
            mock_add.assert_called_once_with(tag, paragraph)
