from __future__ import annotations

import base64
import re
from io import BytesIO
from typing import TYPE_CHECKING

import requests
from PIL import Image, ImageDraw, ImageFilter
from docx.shared import Inches

from htmltodocx.parse_style import css_length_to_inches, parse_style_css

if TYPE_CHECKING:
    from bs4.element import Tag
    from docx.text.paragraph import Paragraph


# ==========================
# Публичная точка входа
# ==========================

def add_image_to_docx(tag: Tag, paragraph: Paragraph) -> None:
    """
    Вставляет <img> в абзац Word с учётом:
    - src / srcset / data URI
    - относительных путей (через <base href="...">)
    - width / height / %
    - aspect-ratio
    - object-fit: contain / cover
    - ограничений страницы
    - webp → png
    - border-radius / box-shadow (если заданы)
    """
    img = tag if tag.name == "img" else tag.find("img")
    if not img:
        return

    base_url = detect_base_url(tag)
    src = resolve_image_url(img, base_url)
    if not src:
        return

    image_data, mime_type = load_image(src)
    if not image_data:
        return

    # размеры страницы
    section = paragraph.part.document.sections[0]
    page_width = section.page_width
    left_margin = section.left_margin
    right_margin = section.right_margin
    max_width_inches = (page_width - left_margin - right_margin) / 914400  # EMU → inches

    # стили
    style = parse_style_css(tag.get("style", ""))

    # width / height из CSS
    target_width_inches: float | None = None
    target_height_inches: float | None = None

    if "width" in style:
        w = style["width"]
        if "%" in w:
            target_width_inches = max_width_inches * float(w.strip("%")) / 100
        else:
            target_width_inches = css_length_to_inches(w)

    if "height" in style:
        h = style["height"]
        if "%" not in h:
            target_height_inches = css_length_to_inches(h)

    # aspect-ratio
    aspect_ratio: float | None = None
    css_ar = style.get("aspect-ratio", "").split("/")
    if len(css_ar) == 2:
        try:
            aspect_ratio = float(css_ar[0]) / float(css_ar[1])
        except ZeroDivisionError:
            aspect_ratio = None

    # object-fit
    object_fit = style.get("object-fit")

    insert_image_with_styles(
        paragraph=paragraph,
        image_data=image_data,
        max_width_inches=max_width_inches,
        target_width_inches=target_width_inches,
        target_height_inches=target_height_inches,
        aspect_ratio=aspect_ratio,
        object_fit=object_fit,
        style=style,
    )


# ==========================
# Вставка с учётом стилей
# ==========================

def insert_image_with_styles(
    paragraph: Paragraph,
    image_data: bytes,
    max_width_inches: float,
    target_width_inches: float | None,
    target_height_inches: float | None,
    aspect_ratio: float | None,
    object_fit: str | None,
    style: dict[str, str],
) -> None:
    img = Image.open(BytesIO(image_data)).convert("RGBA")
    width_px, height_px = img.size
    dpi = img.info.get("dpi", (96, 96))[0] or 96

    natural_width_inches = width_px / dpi
    natural_height_inches = height_px / dpi

    # начальные размеры
    width_inches = natural_width_inches
    height_inches = natural_height_inches

    # 1. width / height без object-fit
    if target_width_inches and not target_height_inches:
        width_inches = target_width_inches
        scale = width_inches / natural_width_inches
        height_inches = natural_height_inches * scale
    elif target_height_inches and not target_width_inches:
        height_inches = target_height_inches
        scale = height_inches / natural_height_inches
        width_inches = natural_width_inches * scale
    elif target_width_inches and target_height_inches and not object_fit:
        width_inches = target_width_inches
        height_inches = target_height_inches

    # 2. aspect-ratio
    if aspect_ratio:
        if target_width_inches and not target_height_inches:
            height_inches = width_inches / aspect_ratio
        elif target_height_inches and not target_width_inches:
            width_inches = height_inches * aspect_ratio

    # 3. object-fit
    if object_fit in ("contain", "cover") and target_width_inches and target_height_inches:
        if object_fit == "contain":
            scale = min(
                target_width_inches / natural_width_inches,
                target_height_inches / natural_height_inches,
            )
        else:  # cover
            scale = max(
                target_width_inches / natural_width_inches,
                target_height_inches / natural_height_inches,
            )
        width_inches = natural_width_inches * scale
        height_inches = natural_height_inches * scale

    # 4. ограничение по ширине страницы
    if width_inches > max_width_inches:
        scale = max_width_inches / width_inches
        width_inches = max_width_inches
        height_inches *= scale

    # 5. ограничение по высоте страницы
    section = paragraph.part.document.sections[0]
    page_height_inches = section.page_height / 914400
    top_margin = section.top_margin / 914400
    bottom_margin = section.bottom_margin / 914400
    max_height_inches = page_height_inches - top_margin - bottom_margin

    if height_inches > max_height_inches:
        scale = max_height_inches / height_inches
        height_inches = max_height_inches
        width_inches *= scale

    # 6. border-radius
    if "border-radius" in style:
        radius_px = css_length_to_pixels(style["border-radius"], dpi)
        img = apply_border_radius(img, radius_px)

    # 7. box-shadow (упрощённо)
    if "box-shadow" in style:
        img = apply_box_shadow(img)

    # 8. сохраняем в PNG
    stream = BytesIO()
    img.save(stream, format="PNG")
    stream.seek(0)

    run = paragraph.add_run()
    run.add_picture(stream, width=Inches(width_inches), height=Inches(height_inches))

    # 9. display / float / margin (частично)
    display = style.get("display")
    if display == "block":
        paragraph.add_run().add_break()

    # float можно доработать при необходимости:
    # float_value = style.get("float")
    # if float_value in ("left", "right"):
    #     apply_float(run, float_value)


# ==========================
# Загрузка изображений
# ==========================

def load_image(src: str) -> tuple[bytes | None, str | None]:
    # data URI
    if src.startswith("data:image/"):
        match = re.match(r"data:image/(\w+);base64,(.*)", src)
        if not match:
            return None, None
        mime = match.group(1)
        data = base64.b64decode(match.group(2))
        return data, mime

    # http/https
    if src.startswith(("http://", "https://")):
        r = requests.get(src, headers={"User-Agent": "Mozilla/5.0"}, stream=True)
        r.raise_for_status()
        content_type = r.headers.get("Content-Type", "")
        if not content_type.startswith("image/"):
            return None, None
        mime = content_type.split("/")[-1]
        data = r.content
        if mime.lower() == "webp":
            data = convert_webp_to_png(data)
            mime = "png"
        return data, mime

    return None, None


def convert_webp_to_png(webp_bytes: bytes) -> bytes:
    img = Image.open(BytesIO(webp_bytes)).convert("RGBA")
    output = BytesIO()
    img.save(output, format="PNG")
    return output.getvalue()


# ==========================
# URL: src / srcset / base
# ==========================

def detect_base_url(tag: Tag) -> str | None:
    """
    Ищет <base href="..."> в документе.
    """
    html = tag.find_parent("html")
    if not html:
        return None
    base_tag = html.find("base", href=True)
    if not base_tag:
        return None
    return base_tag.get("href")


def resolve_image_url(tag: Tag, base_url: str | None) -> str | None:
    # srcset
    srcset = tag.get("srcset")
    if srcset:
        candidates: list[tuple[int, str]] = []
        for part in srcset.split(","):
            url, _, size = part.strip().partition(" ")
            width = int(size[:-1]) if size.endswith("w") and size[:-1].isdigit() else 0
            candidates.append((width, url))
        candidates.sort(reverse=True)
        src = candidates[0][1]
    else:
        src = tag.get("src")

    if not src:
        return None

    # относительные пути
    if base_url and not src.startswith(("http://", "https://", "data:")):
        from urllib.parse import urljoin
        src = urljoin(base_url, src)

    return src


# ==========================
# Визуальные эффекты
# ==========================

def css_length_to_pixels(value: str, dpi: int) -> int:
    """
    Простейшая конвертация CSS длины в пиксели.
    Поддерживает px, pt, em (условно), без единиц — px.
    """
    value = value.strip()
    if value.endswith("px"):
        return int(float(value[:-2]))
    if value.endswith("pt"):
        return int(float(value[:-2]) * dpi / 72)
    if value.endswith("em"):
        return int(float(value[:-2]) * 16)  # условно
    try:
        return int(float(value))
    except ValueError:
        return 0


def apply_border_radius(img: Image.Image, radius_px: int) -> Image.Image:
    if radius_px <= 0:
        return img
    mask = Image.new("L", img.size, 0)
    draw = ImageDraw.Draw(mask)
    draw.rounded_rectangle((0, 0, img.size[0], img.size[1]), radius_px, fill=255)
    img = img.copy()
    img.putalpha(mask)
    return img


def apply_box_shadow(img: Image.Image, shadow_size: int = 20, shadow_color=(0, 0, 0, 80)) -> Image.Image:
    shadow = Image.new("RGBA", (img.width + shadow_size * 2, img.height + shadow_size * 2), (0, 0, 0, 0))
    shadow_layer = Image.new("RGBA", img.size, shadow_color)
    shadow.paste(shadow_layer, (shadow_size, shadow_size))
    shadow = shadow.filter(ImageFilter.GaussianBlur(shadow_size / 2))
    shadow.paste(img, (shadow_size, shadow_size), img)
    return shadow
