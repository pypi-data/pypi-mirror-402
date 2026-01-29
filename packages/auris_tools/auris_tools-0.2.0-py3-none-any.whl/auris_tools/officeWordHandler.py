import io
import logging
import re
from typing import List, Optional

import boto3
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.text.paragraph import Paragraph

from auris_tools.configuration import AWSConfiguration


class OfficeWordHandler:
    """
    Handler for DOCX operations including text extraction and manipulation.

    This class provides methods to interact with Microsoft Word documents (DOCX)
    stored in S3, including reading, extracting text, and text replacement operations.
    """

    def __init__(self, config=None):
        """
        Initialize the Office Word handler with AWS configuration.

        Args:
            config: An AWSConfiguration object, or None to use environment variables
        """
        if config is None:
            config = AWSConfiguration()

        # Create a boto3 session with the configuration
        session = boto3.session.Session(**config.get_boto3_session_args())

        # Create an S3 client with additional configuration if needed
        self.s3_client = session.client('s3', **config.get_client_args())
        logging.info(f'Initialized S3 client in region {config.region}')

    def read_from_s3(self, bucket_name, object_name, as_bytes_io=False):
        """
        Read a DOCX file from S3 and return its bytes.

        Args:
            bucket_name: Name of the S3 bucket containing the document
            object_name: Object key of the document in the S3 bucket
            as_bytes_io: If True, return a BytesIO object instead of raw bytes

        Returns:
            bytes or BytesIO: The document content

        Raises:
            Exception: If there is an error retrieving the document
        """
        try:
            response = self.s3_client.get_object(
                Bucket=bucket_name, Key=object_name
            )
            content = response['Body'].read()

            if as_bytes_io:
                return io.BytesIO(content)
            return content
        except Exception as e:
            logging.error(
                f'Error reading document from {bucket_name}/{object_name}: {str(e)}'
            )
            raise Exception(f'Error reading file from S3: {str(e)}')

    def upload_docx(self, docx_document, bucket_name, object_name):
        """
        Upload a DOCX document to S3.

        Args:
            docx_document: The Document object to upload
            bucket_name: Name of the S3 bucket
            object_name: Object key for the document in S3

        Returns:
            bool: True if upload was successful, False otherwise

        Raises:
            Exception: If there is an error uploading the document
        """
        try:
            logging.info(f'Starting upload to S3: {bucket_name}/{object_name}')

            # Convert document to bytes
            temp_stream = io.BytesIO()
            docx_document.save(temp_stream)
            temp_stream.seek(0)
            document_size = len(temp_stream.getvalue())

            # Upload to S3
            self.s3_client.upload_fileobj(
                temp_stream,
                Bucket=bucket_name,
                Key=object_name,
                ExtraArgs={
                    'ContentType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                },
            )

            logging.info(
                f'Upload finished successfully. Size: {document_size} bytes'
            )
            return True
        except Exception as e:
            logging.error(f'Failed to upload to S3: {str(e)}')
            raise Exception(f'Error uploading file to S3: {str(e)}')

    def get_text_from_bytes(self, bytes_data):
        """
        Extract text from a DOCX file bytes.

        Args:
            bytes_data: The document bytes

        Returns:
            str: Extracted text from the document

        Raises:
            ValueError: If there is an error extracting the text
        """
        try:
            doc = Document(io.BytesIO(bytes_data))
            full_text = []

            # Extract text from paragraphs
            for para in doc.paragraphs:
                full_text.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)

            return '\n'.join(full_text)
        except Exception as e:
            logging.error(f'Error extracting text from DOCX: {str(e)}')
            raise ValueError(f'Error extracting text from DOCX: {str(e)}')

    def clean_text(self, text):
        """
        Clean extracted text from a DOCX file.

        Args:
            text: Text to clean

        Returns:
            str: Cleaned text
        """
        if not text:
            return ''

        # Basic cleaning (can be extended)
        cleaned_text = text.strip()
        return cleaned_text

    def collect_all_paragraphs(self, document: Document) -> List[Paragraph]:
        """
        Collect all paragraphs from a Document object.

        This method collects paragraphs from the main document body,
        tables, headers, and footers.

        Args:
            document: The Document object

        Returns:
            List[Paragraph]: List of all paragraphs in the document
        """
        paragraphs = list(document.paragraphs)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)

        for section in document.sections:
            paragraphs.extend(section.header.paragraphs)
            paragraphs.extend(section.footer.paragraphs)

        return paragraphs

    def replace_placeholder_by_text(
        self,
        paragraphs: List[Paragraph],
        document: Document,
        placeholder: str,
        replacement: str,
        max_count: Optional[int] = None,
    ) -> int:
        """
        Replace placeholder text with replacement in document's XML w:t nodes.

        Args:
            paragraphs: List of paragraphs to process
            document: Document object
            placeholder: Text to find and replace
            replacement: Text to insert instead of placeholder
            max_count: Maximum number of replacements, or None for unlimited

        Returns:
            int: Number of replacements made

        Note:
            This method works at the XML level to ensure proper formatting is preserved.
        """
        count = 0
        WORD_NAMESPACE = (
            'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        )
        T_TAG = f'{{{WORD_NAMESPACE}}}t'

        if placeholder in replacement:
            logging.warning(
                f'Replacement skipped to avoid recursion: {placeholder} -> {replacement}'
            )
            return 0

        def replace_in_element(element):
            nonlocal count
            for node in element.iter(tag=T_TAG):
                if node.text and placeholder in node.text:
                    remaining = (
                        None if max_count is None else max_count - count
                    )
                    new_text, n = re.subn(
                        re.escape(placeholder),
                        replacement,
                        node.text,
                        count=remaining if remaining else 0,
                    )
                    if n > 0:
                        node.text = new_text
                        count += n
                        if max_count is not None and count >= max_count:
                            return True
            return False

        # Main paragraphs
        for para in paragraphs:
            if replace_in_element(para._element):
                return count

        # Headers/footers
        for section in document.sections:
            for container in [section.header, section.footer]:
                for para in container.paragraphs:
                    if replace_in_element(para._element):
                        return count

        # Tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if replace_in_element(para._element):
                            return count

        # Hyperlinks
        for rel in document.part.rels.values():
            if rel.reltype == RT.HYPERLINK and placeholder in rel.target_ref:
                logging.info(
                    f'Replacing hyperlink: {rel.target_ref} -> {rel.target_ref.replace(placeholder, replacement)}'
                )
                rel._target = rel.target_ref.replace(placeholder, replacement)
                count += 1

        return count
