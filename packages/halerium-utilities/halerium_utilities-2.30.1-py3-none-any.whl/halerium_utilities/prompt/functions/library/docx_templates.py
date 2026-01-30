from pydantic.v1 import BaseModel, Field


def _find_substrings(text):
    """
    Find placeholder substrings in document
    """
    import re

    # Define the regex pattern to find substrings between {{ and }}
    pattern = r'\{\{.*?\}\}'

    # Use re.findall() to find all matches in the text
    matches = re.findall(pattern, text)

    return matches


def _collect_placeholders_from_item(paragraphs, placeholders):
    """
    Get all placeholders from items
    """
    for paragraph in paragraphs:
        full_text = ''.join(run.text for run in paragraph.runs)
        matches = _find_substrings(full_text)
        placeholders.extend(matches)
    return placeholders


class ReturnPlaceholdersArgs(BaseModel):
    template_path: str = Field(desciption="The path to the docx template.")


async def get_docx_placeholders(data: ReturnPlaceholdersArgs):
    """
    Return list of placeholders of the pattern '{{key}}' in the docx template.
    """

    from docx import Document
    import asyncio

    template_path = data.get('template_path')

    cfg = data.get("config_parameters", {})

    doc = await asyncio.to_thread(Document, template_path)
    placeholders = []

    # look for placeholders in paragraphs
    placeholders = _collect_placeholders_from_item(doc.paragraphs, placeholders)

    # look for placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                placeholders = _collect_placeholders_from_item(cell.paragraphs, placeholders)

    # look for placeholders in headers and footers
    for section in doc.sections:
        # headers
        placeholders = _collect_placeholders_from_item(section.header.paragraphs, placeholders)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    placeholders = _collect_placeholders_from_item(cell.paragraphs, placeholders)

        # footers
        placeholders = _collect_placeholders_from_item(section.footer.paragraphs, placeholders)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    placeholders = _collect_placeholders_from_item(cell.paragraphs, placeholders)
    return str(placeholders)


async def _replace_text_preserving_formatting(paragraph, search_text, replacement_text):
    # Join all runs to handle placeholders spanning multiple runs
    full_text = ''.join(run.text for run in paragraph.runs)
    if search_text in full_text:
        # Replace the text in the full string
        new_text = full_text.replace(search_text, replacement_text)
        # Clear the existing runs
        for run in paragraph.runs:
            run.text = ''
        # Add the new text back into the runs
        paragraph.runs[0].text = new_text


async def _replace_in_tables(tables, search_text, replacement_text):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    await _replace_text_preserving_formatting(paragraph, search_text, replacement_text)


async def _replace_in_headers_footers(section, search_text, replacement_text):
    # Replace in headers
    for header in section.header.paragraphs:
        await _replace_text_preserving_formatting(header, search_text, replacement_text)
    await _replace_in_tables(section.header.tables, search_text, replacement_text)

    # Replace in footers
    for footer in section.footer.paragraphs:
        await _replace_text_preserving_formatting(footer, search_text, replacement_text)
    await _replace_in_tables(section.footer.tables, search_text, replacement_text)


class FillTemplateArgs(BaseModel):
    template_path: str = Field(description="The path to the DOCX template.")
    output_path: str = Field(description="The path of the output file.")
    context: str = Field(
        description="The JSON string of the dictionary containing the placeholders as keys and their corresponding values to be replaced in the template. Multiple placeholders can be filled in a single call.")


async def fill_docx_template(data: FillTemplateArgs):
    """
    This function creates a copy of the template with the placeholders replaced by the context.
    The context parameter should be a JSON string containing all the placeholders as keys and their corresponding values to be replaced in the template.

    Note:
    - For in-place updates, ensure that the output_path is the same as the template_path to maintain previously filled placeholders.
    - Example for context: '{"{{Title}}": "Hello", "{{Body}}": "world"}'.
    """

    import json
    from docx import Document
    import asyncio

    template_path = data.get('template_path')
    output_path = data.get('output_path')
    context = data.get("context", "{}")
    context = json.loads(context)

    cfg = data.get("config_parameters", {})

    # Load the template document
    doc = await asyncio.to_thread(Document, template_path)

    # Replace placeholders in the document body while preserving formatting
    for paragraph in doc.paragraphs:
        for key, value in context.items():
            await _replace_text_preserving_formatting(paragraph, key, value)

    # Replace placeholders in tables within the document body
    for key, value in context.items():
        await _replace_in_tables(doc.tables, key, value)

    # Replace placeholders in headers and footers, including tables within them
    for section in doc.sections:
        for key, value in context.items():
            await _replace_in_headers_footers(section, key, value)

    # Save the modified document
    await asyncio.to_thread(doc.save, output_path)
    return "Done"
