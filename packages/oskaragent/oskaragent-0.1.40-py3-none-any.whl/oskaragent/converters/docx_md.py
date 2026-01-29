"""Conversor de arquivos DOCX para Markdown com exportação de mídias."""

from __future__ import annotations

from docx.document import Document
from docx import Document
from pathlib import Path
import re


def convert_docx_to_markdown(docx_path: str, md_path: str, media_dir: str = None):
    """Converte um arquivo DOCX em Markdown, salvando mídias associadas.

    Args:
        docx_path: Caminho absoluto ou relativo para o arquivo `.docx`.
        md_path: Caminho de saída para o arquivo Markdown gerado.
        media_dir: Diretório onde imagens extraídas serão persistidas.

    Returns:
        Caminho do arquivo Markdown escrito no disco.
    """
    if media_dir is None:
        media_dir = Path(md_path).parent.name

    Path(media_dir).mkdir(parents=True, exist_ok=True)
    doc = Document(docx_path)

    md_lines = []
    images = extract_images_from_docx(doc, media_dir)

    for element in doc.element.body:
        if element.tag.endswith("p"):
            para = next(p for p in doc.paragraphs if p._p is element)
            md_lines.append(paragraph_to_md_docx(para))

        elif element.tag.endswith("tbl"):
            table = next(t for t in doc.tables if t._tbl is element)
            md_lines.append(table_to_md_docx(table))
            md_lines.append("")

    for img in images:  # append non-inline images
        md_lines.append(f"![image]({img})")

    Path(md_path).write_text("\n".join(md_lines), encoding="utf-8")
    return md_path


def clean_text(text):
    """Normaliza quebras e tabulações para manter o Markdown consistente.

    Args:
        text: Conteúdo cru do parágrafo ou run.

    Returns:
        Texto limpo, sem tabulações e com espaços trimados.
    """
    return text.replace("\t", "    ").strip()


def run_to_md_docx(run):
    """Aplica formatação Markdown a um run de texto do DOCX.

    Args:
        run: Fragmento de texto (`Run`) com estilos aplicados.

    Returns:
        Texto em Markdown com negrito, itálico e sublinhado refletidos.
    """
    text = clean_text(run.text)

    if run.bold:
        text = f"**{text}**"
    if run.italic:
        text = f"*{text}*"
    if run.underline:
        text = f"__{text}__"

    return text

def detect_list_level(paragraph):
    """Detecta nível de indentação aproximado para listas aninhadas.

    Args:
        paragraph: Parágrafo cujo recuo será analisado.

    Returns:
        Nível de indentação estimado em relação ao recuo do parágrafo.
    """
    indent = paragraph.paragraph_format.left_indent
    if indent is None:
        return 0
    return int(indent.pt // 20)  # escala aproximada

def paragraph_to_md_docx(para):
    """Converte parágrafos em Markdown preservando títulos e listas.

    Args:
        para: Parágrafo do DOCX com estilo e runs de texto.

    Returns:
        Representação em Markdown do parágrafo, incluindo hierarquia de listas.
    """
    style = para.style.name.lower()
    text = "".join(run_to_md_docx(r) for r in para.runs).strip()

    if not text:
        return ""

    # heading
    if "heading" in style:
        level = int(re.findall(r"\d+", style)[0])
        return "#" * level + " " + text

    # listas aninhadas
    level = detect_list_level(para)
    indent = "  " * level

    if "list" in style or level > 0:
        prefix = "- "
        return indent + prefix + text

    return text

def table_to_md_docx(table):
    """Converte uma tabela DOCX em Markdown.

    Args:
        table: Tabela proveniente do documento DOCX.

    Returns:
        String em Markdown contendo cabeçalhos e linhas da tabela.
    """
    md = []
    rows = table.rows
    headers = rows[0].cells
    md.append("| " + " | ".join(c.text.strip() for c in headers) + " |")
    md.append("| " + " | ".join("---" for _ in headers) + " |")

    for row in rows[1:]:
        md.append("| " + " | ".join(c.text.strip() for c in row.cells) + " |")

    return "\n".join(md)

def extract_images_from_docx(doc, output_folder):
    """Extrai imagens do DOCX e salva no disco.

    Args:
        doc: Documento carregado pelo `python-docx`.
        output_folder: Diretório onde as imagens serão escritas.

    Returns:
        Lista de caminhos das imagens extraídas.
    """
    images = []
    for rel in doc.part._rels:
        rel_obj = doc.part._rels[rel]
        if "image" in rel_obj.target_ref:
            img_part = rel_obj.target_part
            img_data = img_part.blob

            img_name = rel_obj.target_ref.split("/")[-1]
            img_path = Path(output_folder) / img_name
            img_path.write_bytes(img_data)

            images.append(str(img_path))

    return images
