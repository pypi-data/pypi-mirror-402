# encoding: utf-8
from . import sx
#pip install python-docx==0.8.11
from docx import Document
from docx.shared import Cm,Pt,Inches
import base64
from PIL import Image
from io import BytesIO
import re
class 试题():
    题目=解析=选项=答案=题型=''
    题目图片=解析图片=[]
    def __str__(self):
        s='''{}\n{}{}\n{}\n{}\n{}\n'''.format(self.题目,
                                              '图片:' + str(self.题目图片)+'\n' if self.题目图片 else '',
                                              self.选项,
                                              self.答案,
                                              self.解析,
                                              '图片:' + str(self.解析图片)+'\n' if self.解析图片 else '')
        return s
def 解析试题到文档带图片(st:试题, 段落, 过滤字符列表:list=['\x08', '\x07'])->None:
    '''
    :param st:试题对象
    :param 段落: doc = document.add_paragraph('')  # 添加段落
    :return:
    '''
    def add_row(s:str):
        nonlocal 段落
        assert isinstance(s,str),'非文本对象,无法写入到文档'
        for x in 过滤字符列表:
            s = s.replace(x, '')
        段落.add_run(f'{s}')
    def replace_img(x:str):
        '''x="aaa<img src='' />"'''
        nonlocal 段落
        x=re.sub('<img.*?src="(.*?)".*?\/>',r'@image=@image_src=\1@image=',x)
        lst=x.split('@image=')
        for s in lst:
            if s[:len('@image_src=')]=='@image_src=':
                #处理图片
                try:
                    img_url=s[len('@image_src='):]
                    pic = 段落.add_run()
                    if 'data:image' in img_url:
                        img = sx.base64图片转PIL(img_url.strip())
                    else:
                        img = sx.获取_网络图片(图片网址=img_url.strip(),pil=True)
                    img.save('temp.png')
                    width, height = img.size
                    run = 段落.add_run()
                    w = int(width)
                    h = int(height)
                    bili=width/height
                    if width>=500:
                        w=500
                        h=int(w/bili)
                    print('下载图片-->', f'{int(w)}x{int(h)}',img_url)
                    #75像素1英寸
                    picture = run.add_picture('temp.png', width=Inches(w/75), height=Inches(w/75))
                    #picture.alignment = 1  # 图片布局居中
                    picture.alignment = 0  # 图片布局居左
                    paragraph_format = 段落.paragraph_format
                    paragraph_format.keep_with_next = True  # keep_with_next 属性设置为 True
                    paragraph_format.keep_together = True  # keep_together 属性设置为 True
                except Exception as e:
                    sx.打印错误(e)
            else:
                add_row(s)
    replace_img(st.题目)
    段落.add_run('\n')
    replace_img(st.选项)
    段落.add_run('\n')
    replace_img(st.答案)
    段落.add_run('\n')
    replace_img(st.解析)
    段落.add_run('\n')
def 解析试题到文档(st:试题, 段落, 过滤字符列表:list=['\x08', '\x07'],题目图片尺寸:float=10,解析图片尺寸:float=10)->None:
    '''
    :param st:试题对象
    :param 段落: doc = document.add_paragraph('')  # 添加段落
    :return:
    '''
    def add_row(s:str):
        assert isinstance(s,str),'非文本对象,无法写入到文档'
        for x in 过滤字符列表:
            s = s.replace(x, '')
        段落.add_run(f'{s}\n')
    add_row(sx.XPATH(st.题目).取文本())#去掉html标签
    if st.题目图片:
        for img_url in st.题目图片:
            try:
                pic = 段落.add_run()
                if 'data:image' in img_url:
                    img=sx.base64图片转PIL(img_url.strip())
                else:
                    img = sx.获取_网络图片(图片网址=img_url.strip(), pil=1)
                img.save('temp.png')
                pic.add_picture('temp.png', width=Cm(题目图片尺寸))
                段落.add_run('\n')
            except Exception as e:
                sx.打印错误(e)
    add_row(st.选项)
    add_row(st.答案)
    add_row(sx.XPATH(st.解析).取文本()) #去掉html标签
    if st.解析图片:
        for img_url in st.解析图片:
            try:
                pic = 段落.add_run()
                if 'data:image' in img_url:
                    img = sx.base64图片转PIL(img_url.strip())
                else:
                    img = sx.获取_网络图片(图片网址=img_url.strip(), pil=1)
                img.save('temp.png')
                pic.add_picture('temp.png', width=Cm(解析图片尺寸))
                段落.add_run('\n')
            except Exception as e:
                sx.打印错误(e)
    段落.add_run('\n')
def html_to_docx( html,xpath_表达式, save_name):
    # from docx import Document
    # from docx.shared import Cm, Pt
    # import base64
    # from PIL import Image
    # from io import BytesIO
    # import re
    pp = sx.创建目录(save_name)
    if pp.可创建:
        html = re.sub('<img.*?src="(.*?)".*?>', r'@imgsrc=\1@img', html)
        html = re.sub('<br.*?>', '\n', html)
        xp = sx.XPATH(html)
        #去掉script
        dd = xp.xpath('.//script')
        for d in dd:
            d.getparent().remove(d)
        #div = xp.xpath('//div[@class="blog_cont"]')[0]
        div = xp.xpath(xpath_表达式)[0]
        div = div.xpath('string(.)')
        document = Document()
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal'].font.size = Pt(12)  # 小四
        p = document.add_paragraph('')  # 添加段落
        div = div.split('@img')
        for d in div:
            if 'src=' == d[:4]:
                d = d[4:]
                try:
                    if 'data:image' in d:
                        img = sx.base64图片转PIL(d.strip())
                    else:
                        img = sx.获取_网络图片(图片网址=d.strip(), pil=1)
                    img.save('temp.png')
                    pic = p.add_run()
                    pic.add_picture('temp.png', width=Cm(15))
                except Exception as e:
                    sx.打印错误(e)
            else:
                p.add_run(d)
            print(d)
        document.save(pp.文件路径)
        print('保存路径', pp.文件路径)
if __name__ == '__main__':
    #使用实例
    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal'].font.size = Pt(12)  # 小四
    p = document.add_paragraph('')  # 添加段落
    st=试题()
    st.题目='题目'
    st.解析='解析'
    st.选项='A.1\nB.2'
    st.答案='A'
    st.题型='选择题'
    st.解析图片=['http://ip.wgnms.top:9980/static/images/default.jpg']
    st.题目图片=['http://ip.wgnms.top:9980/static/images/default.jpg']
    print(st)
    解析试题到文档(st,p)
    document.save('abc.docx')