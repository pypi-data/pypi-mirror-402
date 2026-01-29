import argparse
import sys
import pathlib
import os
import copy
import subprocess
import csv
import numpy as np
import re
from xyz_py.atomic import elements as atomic_elements

from . import job
from . import utils as ut
from . import constants as cst
from .exceptions import DataNotFoundError, DataFormattingError

_SHOW_CONV = {
    'on': True,
    'save': False,
    'show': True,
    'off': False
}

_SAVE_CONV = {
    'on': True,
    'save': True,
    'show': False,
    'off': False
}


def extract_coords_func(uargs, save=True):
    '''
    Wrapper for extract_coords function

    Parameters
    ----------
    uargs: argparser object
        command line arguments
    save: bool, default=True
        If True, saves data to file. If False, prints to stdout.

    Returns
    -------
        None
    '''
    from . import extractor as oe
    import xyz_py as xyzp

    # Open file and extract coordinates
    labels, coords = oe.get_coords(
        uargs.output_file,
        coord_type=uargs.type,
        index_style=uargs.index_style
    )

    if save:
        # Save to new .xyz file
        xyzp.save_xyz(
            f'{uargs.output_file.stem}_coords.xyz',
            labels,
            coords,
            comment=f'Coordinates extracted from {uargs.output_file}'
        )
    else:
        for lbl, crd in zip(labels, coords):
            print(f'{lbl} {crd[0]:.4f} {crd[1]:.4f} {crd[2]:.4f}')

    return


def extract_sf_energies_func(uargs):
    '''
    Wrapper for cli call to extract spin-free energies
    '''
    from . import extractor as oe
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.shared import Pt

    cas_energies = oe.SpinFreeEnergyExtractor().extract(
        uargs.output_file,
        before='NEVPT2 TRANSITION ENERGIES'
    )

    # Extract NEVPT2 data
    try:
        nev_energies = oe.SpinFreeEnergyExtractor().extract(
            uargs.output_file,
            after='NEVPT2 TRANSITION ENERGIES'
        )
    except DataNotFoundError:
        nev_energies = []
    except (DataFormattingError, ValueError) as e:
        ut.red_exit(str(e))

    # Section names
    names = ['CASSCF Results'] * len(cas_energies)
    names += ['NEVPT2 Results'] * len(nev_energies)

    if uargs.output_format == 'txt':
        out_name = f'{uargs.output_file.stem}_sfenergies.txt'
        with open(out_name, 'w') as f:
            for it, data in enumerate(cas_energies + nev_energies):
                f.write(f'{names[it]}\n')

                f.write('State, Root, Multiplicity, Relative Energy (cm⁻¹)\n')

                f.write(
                    f'1, {data['root'][0]:d}, {data['multiplicity'][0]:d}, 0\n' # noqa
                )
                for rit in range(1, len(data['root'])):
                    f.write(
                        f'{rit + 1:d}, {data['root'][rit]:d}, {data['multiplicity'][rit]:d}, {data['delta energy (cm^-1)'][rit - 1]:.2f}\n' # noqa
                    )
                if len(names) > 1:
                    f.write('-------------------\n')

    if uargs.output_format == 'docx':
        out_name = f'{uargs.output_file.stem}_sfenergies.docx'

        # Create document
        doc = Document()

        # Add style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)

        title = 'Spin-Free energies'
        doc.add_paragraph(title, style='Normal')

        # Add style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)

        # For each extracted section, print matrix, vectors, and values
        for it, data in enumerate(cas_energies + nev_energies):
            doc.add_paragraph(f'\n {names[it]}\n')

            # Table of data
            table = doc.add_table(rows=len(data['root']) + 1, cols=4)

            table.cell(0, 0).text = 'State'
            table.cell(0, 1).text = 'Root'
            table.cell(0, 2).text = 'Multiplicity'
            table.cell(0, 3).text = 'Relative Energy (cm'
            ss = table.cell(0, 3).paragraphs[0].add_run('-1')
            ss.font.superscript = True
            table.cell(0, 3).paragraphs[0].add_run(')')

            single = {"sz": 12, "color": "#000000", "val": "single"}
            double = {"sz": 12, "color": "#000000", "val": "double"}

            ut.set_cell_border(table.cell(0, 0), bottom=double, top=single)
            ut.set_cell_border(table.cell(0, 1), bottom=double, top=single)
            ut.set_cell_border(table.cell(0, 2), bottom=double, top=single)
            ut.set_cell_border(table.cell(0, 3), bottom=double, top=single)

            # Add data
            for rit in range(len(data['root'])):
                table.cell(rit + 1, 0).text = f'{rit + 1:d}'
                table.cell(rit + 1, 1).text = f'{data['root'][rit]:d}'
                table.cell(rit + 1, 2).text = f'{data['multiplicity'][rit]:d}'
                if rit == 0:
                    table.cell(1, 3).text = '0'
                else:
                    table.cell(rit + 1, 3).text = f'{data['delta energy (cm^-1)'][rit - 1]:.2f}' # noqa

            for cit in range(4):
                ut.set_cell_border(
                    table.cell(len(data['root']), cit),
                    bottom=single
                )

            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER # noqa
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell.paragraphs[0].style = 'Normal'

        try:
            doc.save(out_name)
        except PermissionError:
            ut.red_exit(
                f'Cannot write to {out_name}\n'
                'is the file open somewhere else?'
            )

    ut.cprint(f'Data written to {out_name}', 'cyan')

    return


def extract_so_energies_func(uargs):
    '''
    Wrapper for cli call to extract spin-orbit energies
    '''
    from . import extractor as oe

    cas_energies = oe.SpinOrbitEnergyExtractor().extract(
        uargs.output_file,
        before='QDPT WITH NEVPT2 DIAGONAL ENERGIES'
    )

    # Extract NEVPT2 data
    try:
        nev_energies = oe.SpinOrbitEnergyExtractor().extract(
            uargs.output_file,
            after='QDPT WITH NEVPT2 DIAGONAL ENERGIES'
        )
    except DataNotFoundError:
        nev_energies = []
    except (DataFormattingError, ValueError) as e:
        ut.red_exit(str(e))

    # Section names
    names = ['CASSCF Results'] * len(cas_energies)
    names += ['NEVPT2 Results'] * len(nev_energies)

    out_name = f'{uargs.output_file.stem}_soenergies.txt'
    with open(out_name, 'w') as f:
        for it, energies in enumerate(cas_energies + nev_energies):
            f.write(f'{names[it]}\n')

            f.write('State, Relative Energy (cm⁻¹)\n')

            for eit, energy in enumerate(energies):
                f.write(
                    f'{eit + 1:d}, {energy:.4f}\n'
                )
            if len(names) > 1:
                f.write('-------------------\n')

    ut.cprint(f'Data written to {out_name}', 'cyan')

    return


def extract_hyperfine_func(uargs, save=True):
    '''
    Wrapper for cli call to extract hyperfine

    Parameters
    ----------
    uargs: argparser object
        command line arguments
    save: bool, default=True
        If True, saves data to file, else prints to stdout.
    '''

    from . import extractor as oe
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.shared import Pt

    try:
        all_data = oe.HyperfineExtractor.extract(uargs.output_file)
    except (DataFormattingError, ValueError) as e:
        ut.red_exit(str(e))

    if not save:
        for data in all_data:
            for key, val in data.items():
                print(f'{key}:')
                print(val)
        sys.exit(0)

    if uargs.output_format == 'txt':
        out_name = f'{uargs.output_file.stem}_hyperfine.txt'
        with open(out_name, 'w') as f:
            f.write(f'Data from {uargs.output_file}\n')
            f.write('All values are in MHz\n')
            for data in all_data:
                if len(all_data) > 1:
                    f.write('=================\n')
                for key, val in data.items():
                    f.write(f'{key}:\n')
                    f.write(str(val).replace('[', '').replace(']', ''))
                    f.write('\n')

    if uargs.output_format == 'docx':
        out_name = f'{uargs.output_file.stem}_hyperfine.docx'

        title = 'Hyperfine coupling data from DFT'
        title += (f'Data from {uargs.output_file}\n')

        # Create document
        doc = Document()

        doc.add_heading(title, 0)

        # Add style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)

        # For each extracted section, print all data
        for data in all_data:
            if len(all_data) > 1:
                doc.add_paragraph(
                    f'Nucleus: {data['nucleus']}, Isotope: {data['isotope']}'
                )

            # Full matrix
            matrix = doc.add_table(rows=5, cols=3)

            matrix.cell(0, 0).merge(
                matrix.cell(0, 1)
            ).merge(
                matrix.cell(0, 2)
            )

            matrix.cell(0, 0).text = 'Full Hyperfine Tensor / MHz'

            matrix.cell(1, 0).text = '{:.4f}'.format(data['matrix'][0, 0])
            matrix.cell(2, 0).text = '{:.4f}'.format(data['matrix'][0, 1])
            matrix.cell(3, 0).text = '{:.4f}'.format(data['matrix'][0, 2])

            matrix.cell(1, 1).text = '{:.4f}'.format(data['matrix'][1, 0])
            matrix.cell(2, 1).text = '{:.4f}'.format(data['matrix'][1, 1])
            matrix.cell(3, 1).text = '{:.4f}'.format(data['matrix'][1, 2])

            matrix.cell(1, 2).text = '{:.4f}'.format(data['matrix'][2, 0])
            matrix.cell(2, 2).text = '{:.4f}'.format(data['matrix'][2, 1])
            matrix.cell(3, 2).text = '{:.4f}'.format(data['matrix'][2, 2])

            matrix.cell(4, 0).text = 'Isotropic / MHz'

            # Merge three cells for isotropic value
            matrix.cell(4, 1).merge(
                matrix.cell(4, 2)
            )
            matrix.cell(4, 1).text = f'{data['iso']:.4f}'

            doc.add_paragraph('\n')

            for row in matrix.rows:
                for cell in row.cells:
                    cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER # noqa
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell.paragraphs[0].style = 'Normal'

            # g values and g vectors
            eigs = doc.add_table(rows=4, cols=4)
            eigs.cell(0, 1).merge(eigs.cell(0, 1)).merge(eigs.cell(0, 2)).merge(eigs.cell(0, 3)) # noqa

            eigs.cell(0, 0).text = 'Values / MHz'
            eigs.cell(0, 1).text = 'Vectors'

            eigs.cell(1, 0).text = '{:.4f}'.format(data['values'][0])
            eigs.cell(2, 0).text = '{:.4f}'.format(data['values'][1])
            eigs.cell(3, 0).text = '{:.4f}'.format(data['values'][2])

            eigs.cell(1, 1).text = '{:.4f}'.format(data['vectors'][0, 0])
            eigs.cell(2, 1).text = '{:.4f}'.format(data['vectors'][0, 1])
            eigs.cell(3, 1).text = '{:.4f}'.format(data['vectors'][0, 2])

            eigs.cell(1, 2).text = '{:.4f}'.format(data['vectors'][1, 0])
            eigs.cell(2, 2).text = '{:.4f}'.format(data['vectors'][1, 1])
            eigs.cell(3, 2).text = '{:.4f}'.format(data['vectors'][1, 2])

            eigs.cell(1, 3).text = '{:.4f}'.format(data['vectors'][2, 0])
            eigs.cell(2, 3).text = '{:.4f}'.format(data['vectors'][2, 1])
            eigs.cell(3, 3).text = '{:.4f}'.format(data['vectors'][2, 2])

            for row in eigs.rows:
                for cell in row.cells:
                    cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER # noqa
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell.paragraphs[0].style = 'Normal'
            doc.add_paragraph('\n')

            components = doc.add_table(rows=5, cols=5)
            components.cell(1, 0).text = 'Fermi Contact'
            components.cell(2, 0).text = 'Spin Dipole'
            components.cell(3, 0).text = 'Diamagnetic (Gauge Corr.)'
            components.cell(4, 0).text = 'Spin-Orbit'

            components.cell(0, 1).text = 'x / MHz'
            components.cell(0, 2).text = 'y / MHz'
            components.cell(0, 3).text = 'z / MHz'
            components.cell(0, 4).text = 'Average / MHz'

            if len(data['fc']):
                components.cell(1, 1).text = f'{data['fc'][0]:.4f}'
                components.cell(1, 2).text = f'{data['fc'][1]:.4f}'
                components.cell(1, 3).text = f'{data['fc'][2]:.4f}'
                components.cell(1, 4).text = f'{np.mean(data['fc']):.4f}'
            if len(data['sd']):
                components.cell(2, 1).text = f'{data['sd'][0]:.4f}'
                components.cell(2, 2).text = f'{data['sd'][1]:.4f}'
                components.cell(2, 3).text = f'{data['sd'][2]:.4f}'
                components.cell(2, 4).text = f'{np.mean(data['sd']):.4f}'
            if len(data['dia']):
                components.cell(3, 1).text = f'{data['dia'][0]:.4f}'
                components.cell(3, 2).text = f'{data['dia'][1]:.4f}'
                components.cell(3, 3).text = f'{data['dia'][2]:.4f}'
                components.cell(3, 4).text = f'{np.mean(data['dia']):.4f}'
            if len(data['orb']):
                components.cell(4, 1).text = f'{data['orb'][0]:.4f}'
                components.cell(4, 2).text = f'{data['orb'][1]:.4f}'
                components.cell(4, 3).text = f'{data['orb'][2]:.4f}'
                components.cell(4, 4).text = f'{np.mean(data['orb']):.4f}'

            for tab in [matrix, eigs, components]:
                for row in tab.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER # noqa
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        cell.paragraphs[0].style = 'Normal'

            doc.add_page_break()

        try:
            doc.save(out_name)
        except PermissionError:
            ut.red_exit(
                f'Cannot write to {out_name}\n'
                'is the file open somewhere else?'
            )

    ut.cprint(f'Data written to {out_name}', 'cyan')

    return


def extract_gmatrix_func(uargs, save=True):
    '''
    Wrapper for cli call to extract gmatrix

    Parameters
    ----------
    uargs: argparser object
        command line arguments
    save: bool, default=True
        If True, saves data to file. If False, prints to stdout.

    Returns
    -------
        None
    '''
    from . import extractor as oe
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.shared import Pt, RGBColor

    choices = {
        'total': oe.GMatrixExtractor,
        'L': oe.GMatrixLExtractor,
        'S': oe.GMatrixSExtractor,
        'eff': oe.GMatrixEffectiveExtractor,
        'dft': oe.GMatrixDFTExtractor
    }

    # Check for linear response hyperfines
    if oe.EPRNMRDetector(uargs.output_file):
        uargs.type = 'dft'

    if uargs.type == 'dft':
        cas_data, nev_data = [], []

        # Extract DFT data
        try:
            dft_data = choices[uargs.type]().extract(
                uargs.output_file
            )
        except DataNotFoundError:
            dft_data = []
        except (DataFormattingError, ValueError) as e:
            ut.red_exit(str(e))
    else:
        dft_data = []

        # Extract CASSCF data
        try:
            cas_data = choices[uargs.type]().extract(
                uargs.output_file,
                before='QDPT WITH NEVPT2 DIAGONAL ENERGIES',
                after='QDPT WITH CASSCF DIAGONAL ENERGIES'
            )
        except (DataFormattingError, ValueError) as e:
            ut.red_exit(str(e))

        # Extract NEVPT2 data
        try:
            nev_data = choices[uargs.type]().extract(
                uargs.output_file,
                after='QDPT WITH NEVPT2 DIAGONAL ENERGIES'
            )
        except DataNotFoundError:
            nev_data = []
        except (DataFormattingError, ValueError) as e:
            ut.red_exit(str(e))

    if not save:
        for it, data in enumerate(dft_data):
            print('\nDFT Results')
            for key, val in data.items():
                print(f'{key}:')
                print(val)

        for it, data in enumerate(cas_data):
            print('\nCASSCF Results')
            for key, val in data.items():
                print(f'{key}:')
                print(val)

        for it, data in enumerate(nev_data):
            print('\nNEVPT2 Results')
            for key, val in data.items():
                print(f'{key}:')
                print(val)

        sys.exit(0)

    titles = {
        'total': 'ELECTRONIC G-MATRIX',
        'L': 'ELECTRONIC G-MATRIX: L contribution',
        'S': 'ELECTRONIC G-MATRIX: S contribution',
        'eff': 'ELECTRONIC G-MATRIX FROM EFFECTIVE HAMILTONIAN',
        'dft': 'ELECTRONIC G-MATRIX FROM DFT LINEAR RESPONSE'
    }

    if uargs.output_format == 'txt':
        out_name = f'{uargs.output_file.stem}_gmatrix.txt'
        with open(out_name, 'w') as f:
            f.write(f'G Matrix data from {uargs.output_file}\n')
            f.write(f'{titles[uargs.type]}\n')
            f.write('=======================================\n')

            if len(dft_data):
                f.write('\nDFT Results\n')
                for it, data in enumerate(dft_data):
                    for key, val in data.items():
                        f.write(f'{key}:\n')
                        f.write(str(val).replace('[', '').replace(']', ''))
                        f.write('\n')

            if len(cas_data):
                f.write('\nCASSCF Results\n')
                for it, data in enumerate(cas_data):
                    for key, val in data.items():
                        f.write(f'{key}:\n')
                        f.write(str(val).replace('[', '').replace(']', ''))
                        f.write('\n')

            if len(nev_data):
                f.write('\nNEVPT2 Results\n')
                for it, data in enumerate(nev_data):
                    for key, val in data.items():
                        f.write(f'{key}:\n')
                        f.write(str(val).replace('[', '').replace(']', ''))
                        f.write('\n')

    if uargs.output_format == 'docx':
        out_name = f'{uargs.output_file.stem}_gmatrix.docx'

        # Create document
        doc = Document()

        # Add style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)

        title = titles[uargs.type]
        doc.add_paragraph(title, style='Normal')
        doc.add_paragraph(f'Data from {uargs.output_file}', style='Normal')

        # Section names
        names = ['CASSCF Results'] * len(cas_data)
        names += ['NEVPT2 Results'] * len(nev_data)
        names += ['DFT Results'] * len(dft_data)

        # For each extracted section, print matrix, vectors, and values
        for it, data in enumerate(cas_data + nev_data + dft_data):
            doc.add_paragraph(f'{names[it]}', style='Normal')

            if uargs.type == 'eff':
                doc.add_paragraph(f'Effective S={data['spin_mult']:2d}')

            # Full matrix
            doc.add_paragraph('Full Matrix', style='Normal')

            matrix = doc.add_table(rows=3, cols=3)

            matrix.cell(0, 0).text = '{:.4f}'.format(data['matrix'][0, 0])
            matrix.cell(1, 0).text = '{:.4f}'.format(data['matrix'][0, 1])
            matrix.cell(2, 0).text = '{:.4f}'.format(data['matrix'][0, 2])

            matrix.cell(0, 1).text = '{:.4f}'.format(data['matrix'][1, 0])
            matrix.cell(1, 1).text = '{:.4f}'.format(data['matrix'][1, 1])
            matrix.cell(2, 1).text = '{:.4f}'.format(data['matrix'][1, 2])

            matrix.cell(0, 2).text = '{:.4f}'.format(data['matrix'][2, 0])
            matrix.cell(1, 2).text = '{:.4f}'.format(data['matrix'][2, 1])
            matrix.cell(2, 2).text = '{:.4f}'.format(data['matrix'][2, 2])

            doc.add_paragraph('\n')

            for row in matrix.rows:
                for cell in row.cells:
                    cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER # noqa
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell.paragraphs[0].style = 'Normal'

            # g values and g vectors
            doc.add_paragraph('Eigenpairs', style='Normal')

            table = doc.add_table(rows=4, cols=4)
            table.cell(0, 1).merge(table.cell(0, 1)).merge(table.cell(0, 2)).merge(table.cell(0, 3)) # noqa

            single = {"sz": 12, "color": "#000000", "val": "single"}
            double = {"sz": 12, "color": "#000000", "val": "double"}
            table.cell(0, 0).text = 'Values'
            ut.set_cell_border(table.cell(0, 0), bottom=double, top=single)
            table.cell(0, 1).text = 'Vectors'
            ut.set_cell_border(table.cell(0, 1), bottom=double, top=single)

            table.cell(1, 0).text = '{:.4f}'.format(data['values'][0])
            table.cell(2, 0).text = '{:.4f}'.format(data['values'][1])
            table.cell(3, 0).text = '{:.4f}'.format(data['values'][2])

            table.cell(1, 1).text = '{:.4f}'.format(data['vectors'][0, 0])
            table.cell(2, 1).text = '{:.4f}'.format(data['vectors'][0, 1])
            table.cell(3, 1).text = '{:.4f}'.format(data['vectors'][0, 2])

            table.cell(1, 2).text = '{:.4f}'.format(data['vectors'][1, 0])
            table.cell(2, 2).text = '{:.4f}'.format(data['vectors'][1, 1])
            table.cell(3, 2).text = '{:.4f}'.format(data['vectors'][1, 2])

            table.cell(1, 3).text = '{:.4f}'.format(data['vectors'][2, 0])
            table.cell(2, 3).text = '{:.4f}'.format(data['vectors'][2, 1])
            table.cell(3, 3).text = '{:.4f}'.format(data['vectors'][2, 2])

            ut.set_cell_border(table.cell(3, 0), bottom=single)
            ut.set_cell_border(table.cell(3, 1), bottom=single)
            ut.set_cell_border(table.cell(3, 2), bottom=single)
            ut.set_cell_border(table.cell(3, 3), bottom=single)

            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER # noqa
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell.paragraphs[0].style = 'Normal'

            if it != len(names) - 1:
                doc.add_page_break()
        try:
            doc.save(out_name)
        except PermissionError:
            ut.red_exit(
                f'Cannot write to {out_name}\n'
                'is the file open somewhere else?'
            )

    ut.cprint(f'Data written to {out_name}', 'cyan')

    return


def gen_orbs_func(uargs):
    '''
    Wrapper for CLI gen orbs call

    Parameters
    ----------
    uargs: argparser object
        User arguments

    Returns
    -------
    None
    '''
    from subto.job import SlurmJob

    # Check orca module
    if len(uargs.orca_load):
        orca_load_val = copy.copy(uargs.orca_load)
    elif os.getenv('orto_orca_load'):
        try:
            if len(os.getenv('orto_orca_load')):
                orca_load_val = os.getenv('orto_orca_load')
        except ValueError:
            ut.red_exit(
                (
                    'Error in orto_orca_load environment variable'
                )
            )
    else:
        ut.red_exit(
            (
                'Missing orto_orca_load environment variable or '
                '--orca_load argument'
            )
        )

    # Create string of command numbers for orca plot to follow
    command_nums = '1\\n1\\n4'
    command_nums += f'\\n{uargs.n_pts:d}'
    command_nums += '\\n5\\n7'

    if uargs.beta:
        command_nums += '\\n3\\n1'

    for orbital_number in uargs.orb_numbers:
        command_nums += '\\n2'
        command_nums += f'\\n{orbital_number:d}'
        command_nums += '\\n11'
    command_nums += '\\n12\\n'

    # Generate job file for orca_plot
    orb_job = SlurmJob(
        pathlib.Path(f'{uargs.gbw_file.stem}_orbital_job.sh')
    )

    # Create content of job
    orb_job.content_block = ''

    orb_job.content_block += '# Load orca module\n'
    orb_job.content_block += f'module load {orca_load_val}\n'
    orb_job.content_block += '# Run orca_plot to generate orbital cube file(s)\n' # noqa
    orb_job.content_block += f'time mpirun -np {uargs.n_procs:d} $(which orca_plot_mpi) {uargs.gbw_file.name} -i aa <<< $\'{command_nums}\'\n\n' # noqa

    # Set job name
    orb_job.job_name = f'{uargs.gbw_file.stem}_orbitals'

    orb_job.ntasks_per_node = str(uargs.n_procs)
    orb_job.mem_per_cpu = str(uargs.memory)

    orb_job.error = f'{uargs.gbw_file.stem}_orbitals.%j.e'
    orb_job.output = f'{uargs.gbw_file.stem}_orbitals.%j.o'

    # Write job script
    # with submitter configuration options specified
    orb_job.write_script(True)

    # Submit to queue
    if not uargs.no_sub:
        subprocess.call(
            'cd {}; {} "{}"; cd ../'.format(
                uargs.gbw_file.parents[0],
                orb_job.SUBMIT_COMMAND,
                orb_job.file_path
                ),
            shell=True
        )

    return


def gen_spden_func(uargs):
    '''
    Wrapper for CLI gen spin_density call

    Parameters
    ----------
    uargs: argparser object
        User arguments

    Returns
    -------
    None
    '''
    from subto.job import SlurmJob

    # Check orca module
    if len(uargs.orca_load):
        orca_load_val = copy.copy(uargs.orca_load)
    elif os.getenv('orto_orca_load'):
        try:
            if len(os.getenv('orto_orca_load')):
                orca_load_val = os.getenv('orto_orca_load')
        except ValueError:
            ut.red_exit(
                (
                    'Error in orto_orca_load environment variable'
                )
            )
    else:
        ut.red_exit(
            (
                'Missing orto_orca_load environment variable or '
                '--orca_load argument'
            )
        )

    # Create string of command numbers for orca plot to follow
    command_nums = '1\\n3\\ny\\n4'
    command_nums += f'\\n{uargs.n_pts:d}'
    command_nums += '\\n5\\n7\\n11\\n12\\n'

    # Generate job file for orca_plot
    spin_job = SlurmJob(
        pathlib.Path(f'{uargs.gbw_file.stem}_spin_density_job.sh')
    )

    # Create content of job
    spin_job.content_block = ''

    spin_job.content_block += '# Load orca module\n'
    spin_job.content_block += f'module load {orca_load_val}\n'
    spin_job.content_block += '# Run orca_plot to generate spin density cube file\n' # noqa
    spin_job.content_block += f'time mpirun -np {uargs.n_procs:d} $(which orca_plot_mpi) {uargs.gbw_file.name} -i aa <<< $\'{command_nums}\'\n\n' # noqa

    # Set job name
    spin_job.job_name = f'{uargs.gbw_file.stem}_spin_density'

    spin_job.ntasks_per_node = str(uargs.n_procs)
    spin_job.mem_per_cpu = str(uargs.memory)

    spin_job.error = f'{uargs.gbw_file.stem}_spin_density.%j.e'
    spin_job.output = f'{uargs.gbw_file.stem}_spin_density.%j.o'

    # Write job script
    # with submitter configuration options specified
    spin_job.write_script(True)

    # Submit to queue
    if not uargs.no_sub:
        subprocess.call(
            'cd {}; {} "{}"; cd ../'.format(
                uargs.gbw_file.parents[0],
                spin_job.SUBMIT_COMMAND,
                spin_job.file_path
                ),
            shell=True
        )

    return


def gen_trunc_molden_func(uargs):
    '''
    Wrapper for CLI gen truncmolden call

    Parameters
    ----------
    uargs: argparser object
        User arguments

    Returns
    -------
    None
    '''
    from mmap import mmap, ACCESS_READ
    from shutil import move as shutilmove

    # Read molden file in as binary string
    # and find number of MOs by counting number of
    # occurrences of 'Occup='
    _patt = re.compile(b'Sym=')
    with open(uargs.input_file, mode="r") as file_obj:
        with mmap(file_obj.fileno(), length=0, access=ACCESS_READ) as mmap_obj:
            n_MO = len(_patt.findall(mmap_obj))

    _patt = re.compile(b'Occup= 0.000000')
    with open(uargs.input_file, mode="r") as file_obj:
        with mmap(file_obj.fileno(), length=0, access=ACCESS_READ) as mmap_obj:
            n_virt = len(_patt.findall(mmap_obj))

    ut.cprint(
        (
            f'{n_virt}/{n_MO} MOs are virtual...'
        ),
        'cyan'
    )

    if uargs.output_file is None:
        # If no output file specified, use input file name
        # with .molden extension
        uargs.output_file = '.tmp.molden'

    # Trim file
    _count = 0
    final = False
    with open(uargs.input_file, mode="r") as old:
        with open(uargs.output_file, mode="w") as new:
            # Read in molden file line by line
            for line in old:
                if 'Occup= 0.000000' in line:
                    _count += 1
                if _count == uargs.n_virt:
                    final = True
                if line.startswith(f'{n_MO:d}') and final:
                    new.write(line)
                    break
                else:
                    new.write(line)

    ut.cprint(f'... trimming to {uargs.n_virt} virtual orbitals\n', 'cyan')

    # If no output file given
    if uargs.output_file == '.tmp.molden':
        # Copy new file to original name
        shutilmove(uargs.output_file, uargs.input_file)
        uargs.output_file = uargs.input_file

    ut.cprint(f'New molden file written to {uargs.output_file}', 'cyan')

    return


def gen_job_func(uargs):
    '''
    Wrapper for CLI gen job call

    Parameters
    ----------
    uargs: argparser object
        User arguments

    Returns
    -------
    None
    '''
    from . import input as inp

    for input_file in uargs.input_files:

        # Check input exists
        if not input_file.exists:
            ut.red_exit('Cannot locate {}'.format(input_file.name))

        oj = job.OrcaJob(
            input_file
        )

        # Get orca module load command
        orca_args = [
            'orca_load'
        ]

        required = [
            'orca_load'
        ]

        for oarg in orca_args:
            uarg_val = getattr(uargs, oarg)
            if len(uarg_val):
                oarg_val = copy.copy(uarg_val)
            elif os.getenv(f'orto_{oarg}'):
                try:
                    if len(os.getenv(f'orto_{oarg}')):
                        oarg_val = os.getenv(f'orto_{oarg}')
                except ValueError:
                    ut.red_exit(
                        (
                            f'Error in orto_{oarg} environment variable'
                        )
                    )
            elif oarg in required:
                ut.red_exit(
                    (
                        f'Missing orto_{oarg} environment variable or '
                        f'--{oarg} argument'
                    )
                )
            else:
                oarg_val = ''

            if oarg == 'orca_load':
                oarg = 'load'
                if 'module load' not in oarg_val:
                    oarg_val = f'module load {oarg_val}'

            setattr(oj, oarg, oarg_val)

        # Check xyz file is present
        try:
            inp.check_coord(oj.input_file, uargs.skip_xyz)
        except (DataNotFoundError, DataFormattingError) as e:
            ut.red_exit(str(e))

        # Check for moread and moinp
        try:
            inp.check_moinp_moread(oj.input_file)
        except (DataNotFoundError, DataFormattingError) as e:
            ut.red_exit(str(e))

        # Submitter configuration options
        # currently hardcoded for slurm!
        config = {}

        # Get nprocs and maxcore
        try:
            n_procs = inp.get_nprocs(oj.input_file)
            maxcore = inp.get_maxcore(oj.input_file)
        except (DataNotFoundError, DataFormattingError) as e:
            ut.red_exit(str(e))

        # If memory and procs specified as arguments, give warning when
        # these are smaller than the number in the input file
        if uargs.n_procs:
            if n_procs > uargs.n_procs:
                ut.red_exit('Too few processors requested for input file')
            # Use cli value
            config['ntasks_per_node'] = uargs.n_procs
        else:
            # Use orca file value
            config['ntasks_per_node'] = n_procs

        if uargs.memory:
            if uargs.memory * uargs.n_procs < n_procs * maxcore:
                ut.red_exit('Requested too little memory for orca input')
            config['mem_per_cpu'] = uargs.memory
        else:
            # Use orca file value
            config['mem_per_cpu'] = maxcore

        # Check if NBO is requested
        if inp.get_nbo(oj.input_file):
            # Check if NBO module has been provided to orto
            try:
                if os.getenv('orto_nbo_load') is not None:
                    nbo_module = os.getenv('orto_nbo_load')
                else:
                    ut.red_exit(
                        'Missing orto_nbo_load environment variable'
                    )
            except ValueError:
                ut.red_exit(
                    (
                        'Missing or malformed orto_nbo_load'
                        'environment variable'
                    )
                )
            oj.pre_orca += f'module load {nbo_module}\n'
            oj.pre_orca += f'export NBOFIL={oj.input_file.stem}\n'
        else:
            nbo_module = None

        # Check structure in input file or xyz file
        # unless skip_structure flag is set
        if not uargs.skip_structure:
            try:
                inp.check_structure(oj.input_file)
            except (DataNotFoundError, DataFormattingError) as e:
                ut.red_exit(str(e))

        # Set SLURM error and output file names
        config['error'] = 'slurm.%j.e'
        config['output'] = 'slurm.%j.o'

        # Add call to orca_2mkl to create molden file from gbw
        if not uargs.no_molden:
            oj.post_orca += 'orca_2mkl {} -molden'.format(oj.input_file.stem)
            oj.post_orca += '\norto gen trunc_molden {}.molden.input'.format(
                oj.input_file.stem
            )

        # Write job script
        # with submitter configuration options specified
        oj.write_script(True, **config)

        # Submit to queue
        if not uargs.no_sub:
            subprocess.call(
                'cd {}; {} "{}"; cd ../'.format(
                    oj.input_file.parents[0],
                    oj.Job.SUBMIT_COMMAND,
                    oj.job_file
                    ),
                shell=True
            )
    return


def plot_xes_func(uargs, save_data_only=False):
    '''
    Wrapper for CLI plot abs call\n\n

    Plots ABSORPTION blocks from orca output\n
        - UVVIS (TDDFT)\n
        - XAS (TDDFT)\n\n

    Parameters
    ----------
    uargs: argparser object
        User arguments

    save_data_only: bool, default=False
        If True, saves generated spectrum data to file only.

    Returns
    -------
    None
    '''
    from . import plotter
    from . import extractor as oe
    from . import data as d

    # Process x_shift argument
    if uargs.x_shift is None:
        uargs.x_shift = [0.0 for _ in uargs.output_file]
    elif len(uargs.x_shift) != len(uargs.output_file):
        ut.red_exit(
            'Number of x_shift values must match number of output files'
        )

    # Create dictionary to hold absorption data for plotting later
    spectra_dict = {}

    # Extract data from each output file
    for it, output_file in enumerate(uargs.output_file):

        version = oe.OrcaVersionExtractor.extract(output_file)

        if not len(version):
            ut.cprint(
                'Warning: Cannot find version number in Orca output file',
                'black_yellowbg'
            )
            version = [6, 0, 0]

        # Extract absorption data from file
        # using appropriate extractor for version
        # and intensity type
        if version[0] >= 6:
            if uargs.intensities == 'electric':
                all_datasets = oe.XESElectricDipoleExtractor.extract(
                    output_file
                )
            elif uargs.intensities == 'velocity':
                all_datasets = oe.XESVelocityDipoleExtractor.extract(
                    output_file
                )
            elif uargs.intensities == 'semi-classical':
                all_datasets = oe.XESSemiClassicalDipoleExtractor.extract(
                    output_file
                )
        else:
            ut.red_exit(
                'Orca version < 6 not supported for XES spectrum generation'
            )

        ut.cprint('Using intensities: {}'.format(uargs.intensities), 'cyan')

        # Create absorption data object
        # one per dataset extracted
        # (doubtful multiple datasets will be present, but just in case)
        all_abs_data = [
            d.AbsorptionData.from_extractor_dataset(
                dataset,
                uargs.intensities,
                remove_zero_osc=uargs.zero_osc
            )
            for dataset in all_datasets
        ]

        # Check absorption data was found
        if len(all_abs_data) == 0:
            ut.red_exit(
                f'No ABSORPTION data found in file {output_file}', 'red'
            )
        # report if multiple absorption sections found
        elif len(all_abs_data) > 1:
            ut.cprint(
                f'\nFound {len(all_abs_data)} ABSORPTION sections in '
                f'file {output_file}\n'
                f'Plotting final section ONLY\n',
                'cyan'
            )
        # and only use final one
        # (again, doubtful multiple sections will be present, but just in case)
        abs_data = all_abs_data[-1]

        # Determine x values for setting x limits of computed spectrum
        if uargs.x_unit == 'wavenumber':
            x_vals = abs_data.wavenumbers
            min_factor = 0.8
            max_factor = 1.1
            x_reversed = False
        elif uargs.x_unit == 'wavelength':
            x_vals = abs_data.wavelengths
            max_factor = 1.1
            min_factor = 0.8
            x_reversed = False
        elif uargs.x_unit == 'wavelength_rev':
            x_vals = abs_data.wavelengths
            max_factor = 1.1
            min_factor = 0.8
            x_reversed = True
            uargs.x_unit = 'wavelength'
        elif uargs.x_unit == 'energy':
            x_vals = abs_data.energies
            min_factor = 0.9995
            max_factor = 1.0005
            x_reversed = False

        # Set x_min
        # based on user arguments
        if isinstance(uargs.xlim[0], str):
            if uargs.xlim[0] == 'auto':
                x_min = min(x_vals) * min_factor
            elif ut.is_floatable(uargs.xlim[0]):
                x_min = float(uargs.xlim[0])
            else:
                raise ValueError(f'Invalid x_min value: {uargs.xlim[0]}')
        else:
            x_min = uargs.xlim[0]

        # Set x_max
        # based on user arguments
        if isinstance(uargs.xlim[1], str):
            if uargs.xlim[1] == 'auto':
                x_max = max(x_vals) * max_factor
            elif ut.is_floatable(uargs.xlim[1]):
                x_max = float(uargs.xlim[1])
            else:
                raise ValueError(f'Invalid x_max value: {uargs.xlim[1]}')
        else:
            x_max = uargs.xlim[1]

        # Trim transitions to a specific number of states if requested
        if uargs.trim_transitions is not None:
            abs_data.trim_to_n(uargs.trim_transitions)

        # Generate spectrum data
        abs_data.generate_spectrum(
            fwhm=uargs.linewidth,
            lineshape=uargs.lineshape,
            x_type=uargs.x_unit,
            num_points=10000,
            x_min=x_min,
            x_max=x_max,
            comment=output_file,
            x_reversed=x_reversed
        )

        # Save spectrum data to file
        if save_data_only:
            abs_data.save_spectrum_data(
                f'absorption_spectrum_{output_file.stem}.csv',
                comments='Data from {}\nfwhm={}, lineshape={}\nintensities={}'.format( # noqa
                    output_file,
                    uargs.linewidth,
                    uargs.lineshape,
                    uargs.intensities
                )
            )
            abs_data.save_transition_data(
                f'transition_data_{output_file.stem}.csv',
                comments='Data from {}\nintensities={}'.format(
                    output_file,
                    uargs.intensities
                )
            )
            ut.cprint(
                f'Saved absorption spectrum to absorption_spectrum_{output_file.stem}.csv', # noqa
                'cyan'
            )
            ut.cprint(
                f'Saved absorption data to transition_data_{output_file.stem}.csv', # noqa
                'cyan'
            )
        else:
            # Add to dictionary of absorption data for plotting later
            spectra_dict[output_file] = abs_data

    # Exit if only saving data
    if save_data_only:
        return

    # Plot all spectra
    import matplotlib.pyplot as plt
    import matplotlib.colors as mcolors

    # Set font name and size
    ut.check_font_envvar()
    plt.rcParams['font.size'] = 10
    plt.rcParams['legend.fontsize'] = 9
    plt.rcParams['legend.loc'] = 'center right'

    # Create list of colours for plotting
    if len(uargs.output_file) == 1:
        colours = ['black']
    else:
        colours = list(mcolors.TABLEAU_COLORS.values())

    # Width of figure in inches
    # Make wider if multiple spectra to plot
    # to include legend
    if len(spectra_dict) > 1:
        width = 6.
    else:
        width = 4.
    # width in cm
    width_cm = 2.54 * width

    # Create figure and axis with
    # height based on golden ratio
    golden = (1 + np.sqrt(5))/2
    fig, ax = plt.subplots(
        1,
        1,
        num='Absorption Spectrum',
        figsize=(width, 4. / golden)
    )
    # Secondary axis for oscillator strength sticks
    oax = ax.twinx()

    # Plot absorption spectrum for each file
    for it, (output_file, abs_data) in enumerate(spectra_dict.items()): # noqa

        plotter.plot_absorption_spectrum(
            abs_data,
            linecolor=colours[it],
            stickcolor=colours[it],
            osc_style=uargs.osc_style,
            normalise=uargs.normalise,
            window_title='',
            fig=fig,
            ax=ax,
            oax=oax,
            show=False,
            save=False,
            xlim=[x_min, x_max],
            ylim=uargs.ylim,
            x_shift=uargs.x_shift[it],
            legend=False
        )

    if uargs.xlim != ['auto', 'auto']:
        ax.set_xlim(float(uargs.xlim[0]), float(uargs.xlim[1]))

    fig.tight_layout()

    if len(spectra_dict) > 1:
        fig.subplots_adjust(right=3.3/6.)
        if uargs.legend is None:
            legend_labels = [
                f'{output_file.stem}' for output_file in spectra_dict.keys()
            ]
            fig.legend(legend_labels, loc=7)
        else:
            fig.legend(uargs.legend, loc=7)

    if _SAVE_CONV[uargs.plot]:
        savename = 'absorption_spectrum.png'
        plt.savefig(savename, dpi=500)
        ut.cprint(f'Saved image to {savename}', 'cyan')
        ut.cprint(
            f'Use width={width:.1f} in. or {width_cm:.1f} cm',
            'cyan'
        )

    if _SHOW_CONV[uargs.plot]:
        plt.show()

    return


def plot_abs_func(uargs, save_data_only=False):
    '''
    Wrapper for CLI plot abs call\n\n

    Plots ABSORPTION blocks from orca output\n
        - UVVIS (TDDFT)\n
        - XAS (TDDFT)\n\n

    Parameters
    ----------
    uargs: argparser object
        User arguments

    save_data_only: bool, default=False
        If True, saves generated spectrum data to file only.

    Returns
    -------
    None
    '''
    from . import plotter
    from . import extractor as oe
    from . import data as d

    # Process x_shift argument
    if uargs.x_shift is None:
        uargs.x_shift = [0.0 for _ in uargs.output_file]
    elif len(uargs.x_shift) != len(uargs.output_file):
        ut.red_exit(
            'Number of x_shift values must match number of output files'
        )

    # Create dictionary to hold absorption data for plotting later
    spectra_dict = {}

    # Extract data from each output file
    for it, output_file in enumerate(uargs.output_file):

        version = oe.OrcaVersionExtractor.extract(output_file)

        if not len(version):
            ut.cprint(
                'Warning: Cannot find version number in Orca output file',
                'black_yellowbg'
            )
            version = [6, 0, 0]

        # Extract absorption data from file
        # using appropriate extractor for version
        # and intensity type
        if version[0] < 6:
            if uargs.intensities == 'electric':
                all_datasets = oe.OldAbsorptionElectricDipoleExtractor.extract(
                    output_file
                )
            elif uargs.intensities == 'velocity':
                all_datasets = oe.OldAbsorptionVelocityDipoleExtractor.extract(
                    output_file
                )
        elif version[0] >= 6:
            if uargs.intensities == 'electric':
                all_datasets = oe.AbsorptionElectricDipoleExtractor.extract(
                    output_file
                )
            elif uargs.intensities == 'velocity':
                all_datasets = oe.AbsorptionVelocityDipoleExtractor.extract(
                    output_file
                )
            elif uargs.intensities == 'semi-classical':
                all_datasets = oe.AbsorptionSemiClassicalDipoleExtractor.extract( # noqa
                    output_file
                )

        ut.cprint('Using intensities: {}'.format(uargs.intensities), 'cyan')

        # Create absorption data object
        # one per dataset extracted
        # (doubtful multiple datasets will be present, but just in case)
        all_abs_data = [
            d.AbsorptionData.from_extractor_dataset(
                dataset,
                uargs.intensities,
                remove_zero_osc=uargs.zero_osc
            )
            for dataset in all_datasets
        ]

        # Check absorption data was found
        if len(all_abs_data) == 0:
            ut.red_exit(
                f'No ABSORPTION data found in file {output_file}', 'red'
            )
        # report if multiple absorption sections found
        elif len(all_abs_data) > 1:
            ut.cprint(
                f'\nFound {len(all_abs_data)} ABSORPTION sections in '
                f'file {output_file}\n'
                f'Plotting final section ONLY\n',
                'cyan'
            )
        # and only use final one
        # (again, doubtful multiple sections will be present, but just in case)
        abs_data = all_abs_data[-1]

        # Determine x values for setting x limits of computed spectrum
        if uargs.x_unit == 'wavenumber':
            x_vals = abs_data.wavenumbers
            min_factor = 0.8
            max_factor = 1.1
            x_reversed = False
        elif uargs.x_unit == 'wavelength':
            x_vals = abs_data.wavelengths
            max_factor = 1.1
            min_factor = 0.8
            x_reversed = False
        elif uargs.x_unit == 'wavelength_rev':
            x_vals = abs_data.wavelengths
            max_factor = 1.1
            min_factor = 0.8
            x_reversed = True
            uargs.x_unit = 'wavelength'
        elif uargs.x_unit == 'energy':
            x_vals = abs_data.energies
            min_factor = 0.9995
            max_factor = 1.0005
            x_reversed = False

        # Set x_min
        # based on user arguments
        if isinstance(uargs.xlim[0], str):
            if uargs.xlim[0] == 'auto':
                x_min = min(x_vals) * min_factor
            elif ut.is_floatable(uargs.xlim[0]):
                x_min = float(uargs.xlim[0])
            else:
                raise ValueError(f'Invalid x_min value: {uargs.xlim[0]}')
        else:
            x_min = uargs.xlim[0]

        # Set x_max
        # based on user arguments
        if isinstance(uargs.xlim[1], str):
            if uargs.xlim[1] == 'auto':
                x_max = max(x_vals) * max_factor
            elif ut.is_floatable(uargs.xlim[1]):
                x_max = float(uargs.xlim[1])
            else:
                raise ValueError(f'Invalid x_max value: {uargs.xlim[1]}')
        else:
            x_max = uargs.xlim[1]

        # Trim transitions to a specific number of states if requested
        if uargs.trim_transitions is not None:
            abs_data.trim_to_n(uargs.trim_transitions)

        # Generate spectrum data
        abs_data.generate_spectrum(
            fwhm=uargs.linewidth,
            lineshape=uargs.lineshape,
            x_type=uargs.x_unit,
            num_points=10000,
            x_min=x_min,
            x_max=x_max,
            comment=output_file,
            x_reversed=x_reversed
        )

        # Save spectrum data to file
        if save_data_only:
            abs_data.save_spectrum_data(
                f'absorption_spectrum_{output_file.stem}.csv',
                comments='Data from {}\nfwhm={}, lineshape={}\nintensities={}'.format( # noqa
                    output_file,
                    uargs.linewidth,
                    uargs.lineshape,
                    uargs.intensities
                )
            )
            abs_data.save_transition_data(
                f'transition_data_{output_file.stem}.csv',
                comments='Data from {}\nintensities={}'.format(
                    output_file,
                    uargs.intensities
                )
            )
            ut.cprint(
                f'Saved absorption spectrum to absorption_spectrum_{output_file.stem}.csv', # noqa
                'cyan'
            )
            ut.cprint(
                f'Saved absorption data to transition_data_{output_file.stem}.csv', # noqa
                'cyan'
            )
        else:
            # Add to dictionary of absorption data for plotting later
            spectra_dict[output_file] = abs_data

    # Exit if only saving data
    if save_data_only:
        return

    # Plot all spectra
    import matplotlib.pyplot as plt
    import matplotlib.colors as mcolors

    # Set font name and size
    ut.check_font_envvar()
    plt.rcParams['font.size'] = 10
    plt.rcParams['legend.fontsize'] = 9
    plt.rcParams['legend.loc'] = 'center right'

    # Create list of colours for plotting
    if len(uargs.output_file) == 1:
        colours = ['black']
    else:
        colours = list(mcolors.TABLEAU_COLORS.values())

    # Width of figure in inches
    # Make wider if multiple spectra to plot
    # to include legend
    if len(spectra_dict) > 1:
        width = 6.
    else:
        width = 4.
    # width in cm
    width_cm = 2.54 * width

    # Create figure and axis with
    # height based on golden ratio
    golden = (1 + np.sqrt(5))/2
    fig, ax = plt.subplots(
        1,
        1,
        num='Absorption Spectrum',
        figsize=(width, 4. / golden)
    )
    # Secondary axis for oscillator strength sticks
    oax = ax.twinx()

    # Plot absorption spectrum for each file
    for it, (output_file, abs_data) in enumerate(spectra_dict.items()): # noqa

        plotter.plot_absorption_spectrum(
            abs_data,
            linecolor=colours[it],
            stickcolor=colours[it],
            osc_style=uargs.osc_style,
            normalise=uargs.normalise,
            window_title='',
            fig=fig,
            ax=ax,
            oax=oax,
            show=False,
            save=False,
            xlim=[x_min, x_max],
            ylim=uargs.ylim,
            x_shift=uargs.x_shift[it],
            legend=False
        )

    if uargs.xlim != ['auto', 'auto']:
        ax.set_xlim(float(uargs.xlim[0]), float(uargs.xlim[1]))

    fig.tight_layout()

    if len(spectra_dict) > 1:
        fig.subplots_adjust(right=3.3/6.)
        if uargs.legend is None:
            legend_labels = [
                f'{output_file.stem}' for output_file in spectra_dict.keys()
            ]
            fig.legend(legend_labels, loc=7)
        else:
            fig.legend(uargs.legend, loc=7)

    if _SAVE_CONV[uargs.plot]:
        savename = 'absorption_spectrum.png'
        plt.savefig(savename, dpi=500)
        ut.cprint(f'Saved image to {savename}', 'cyan')
        ut.cprint(
            f'Use width={width:.1f} in. or {width_cm:.1f} cm',
            'cyan'
        )

    if _SHOW_CONV[uargs.plot]:
        plt.show()

    return


def plot_ir_func(uargs):
    '''
    Wrapper for CLI plot_ir call

    Parameters
    ----------
    uargs: argparser object
        User arguments

    Returns
    -------
    None
    '''
    import matplotlib.pyplot as plt
    from . import plotter
    from . import extractor as oe

    # Set user specified font name
    ut.check_font_envvar()

    # Change matplotlib font size to be larger
    plt.rcParams['font.size'] = 10
    plt.rcParams['legend.fontsize'] = 9

    # Extract frequency information
    data = oe.FrequencyExtractor.extract(uargs.output_file)

    if not len(data):
        ut.red_exit(f'Cannot find frequencies in file {uargs.output_file}')

    data = data[0]

    # Plot infrared spectrum
    plotter.plot_ir(
        data['energy (cm^-1)'],
        data['epsilon (L mol^-1 cm^-1)'],
        linewidth=uargs.linewidth,
        lineshape=uargs.lineshape,
        window_title=f'Infrared Spectrum from {uargs.output_file}',
        show=True
    )

    return


def distort_func(uargs):
    '''
    Distorts molecule along specified normal mode

    Parameters
    ----------
    uargs: argparser object
        command line arguments

    Returns
    -------
        None

    '''
    import xyz_py as xyzp
    from . import extractor as oe

    # Open file and extract coordinates
    labels, coords = oe.get_coords(
        uargs.output_file
    )

    # Extract frequency information
    data = oe.FrequencyExtractor.extract(uargs.output_file)

    ut.cprint(
        'Distorting along mode #{}:  {: .2f} cm⁻¹'.format(
            uargs.mode_number,
            data[0]['energy (cm^-1)'][uargs.mode_number]
            ),
        'cyan'
        )

    coords += uargs.scale * data[0]['displacements'][:, uargs.mode_number]

    comment = (
        f'Coordinates from {uargs.output_file} distorted by {uargs.scale:f} unit of' # noqa
        f' Mode #{uargs.mode_number}'
    )

    labels_nn = xyzp.remove_label_indices(labels)

    xyzp.save_xyz('distorted.xyz', labels_nn, coords, comment=comment)

    return


def extract_orbs_func(uargs, save=True) -> None:
    '''
    Extracts Loewdin Orbital contributions from orca output file

    Parameters
    ----------
    uargs: argparser object
        command line arguments
    save: bool, default=True
        If True, saves data to file. If False, prints to stdout.

    Returns
    -------
        None
    '''
    from . import extractor as oe

    # Check for spin in file, if present
    try:
        mult = oe.MultiplicityInputExtractor.extract(uargs.output_file)[0]
        # Disable spin if not present in file
        if mult == 1:
            uargs.spin = None
    except DataNotFoundError:
        pass

    extractors = {
        'orb_comp': lambda x: oe.LoewdinCompositionExtractor.extract(x),
        'redorb_pop': lambda x: oe.LoewdinReducedOrbitalPopulationExtractor.extract( # noqa
            x, spin=uargs.spin
        ),
        'orb_pop': lambda x: oe.LoewdinOrbitalPopulationExtractor.extract(
            x, spin=uargs.spin
        )
    }

    extractor_names = {
        'orb_comp': 'Loewdin Orbital Composition',
        'orb_pop': 'Loewdin Orbital Population',
        'redorb_pop': 'Loewdin Reduced Orbital Population'
    }

    if uargs.flavour == 'first_match':
        failed = 0
        for name, extractor in extractors.items():
            try:
                data = extractor(
                    uargs.output_file
                )
                uargs.flavour = name
                break
            except DataNotFoundError:
                failed += 1
        if failed == len(extractors):
            ut.red_exit(
                'Cannot find Loewdin orbital contributions in file'
            )
    else:
        try:
            data = extractors[uargs.flavour](
                uargs.output_file
            )
        except DataNotFoundError as dne:
            ut.red_exit(str(dne))

    # Unpack data
    contributions, occupancies, energies = data[0]

    # Trim contributions to selected range of MOs
    if uargs.active:
        keep = [
            it for it, val in enumerate(occupancies)
            if 0 < val < 2
        ]
    # HOMO and LUMO plus specified number of orbitals
    # either side
    elif uargs.homo_lumo is not None:
        keep = np.concatenate([
            np.arange(
                np.argmin(occupancies) - 1 - uargs.homo_lumo,
                np.argmin(occupancies),
                1,
                dtype=int
            ),
            np.arange(
                np.argmin(occupancies),
                np.argmin(occupancies) + uargs.homo_lumo + 1,
                1,
                dtype=int
            )
        ])
    elif uargs.num is not None:
        keep = uargs.num
    else:
        keep = range(len(contributions))

    # Remove orbital indices which do not exist
    keep = [val for val in keep if val < len(contributions)]

    if not len(keep):
        ut.red_exit(
            (
                r'Selected orbital indices do not exist!'
                f'\nNORBS = {len(contributions):d}'
            )
        )

    contributions = contributions.loc[:, keep]

    # Remove contributions from unwanted orbitals
    # and shells
    _ORB = {
        's': ['s'],
        'p': ['px', 'py', 'pz'],
        'd': ['dx2y2', 'dz2', 'dxy', 'dxz', 'dyz'],
        'f': ['f0', 'f1', 'f2', 'f3', 'f-1', 'f-2', 'f-3'],
    }
    _orbs_to_use = [
        val
        for chosen in uargs.orb
        for val in _ORB[chosen]
    ]
    if uargs.flavour not in ['redorb_pop', 'orb_comp']:
        _shellorbs_to_use = [
            f'{shell:1d}{orb}'
            for orb in _orbs_to_use
            for shell in uargs.shell
        ]
    else:
        _shellorbs_to_use = _orbs_to_use
    _orb_query = f"AO in {_shellorbs_to_use}"
    contributions = contributions.query(_orb_query)

    # Print info to screen
    ut.cprint(f'Using {extractor_names[uargs.flavour]}', 'cyan')
    if uargs.flavour in ['redorb_pop', 'orb_pop'] and uargs.spin is not None:
        ut.cprint(f'For SPIN={uargs.spin}', 'cyan')
    for mo_num, mo in contributions.items():
        # if no rows greater than threshold, skip
        if not len(mo[mo > uargs.threshold]):
            continue
        _output = ''
        total = 0.
        for row, val in mo.items():
            if val > uargs.threshold and row[1] in uargs.elements:
                _output += f'  {row[1]+str(row[0]):5} {row[2]:5}: {val:>5.1f} %\n' # noqa
                total += val
        if len(_output):
            print(f'MO #{mo_num} (Occ={occupancies[mo_num]}, E={energies[mo_num]: .5f}):') # noqa
            print(_output)
            print(f'  Total:        {total:>5.1f} %\n')

    return


def parse_cutoffs(cutoffs):

    if len(cutoffs) % 2:
        raise argparse.ArgumentTypeError('Error, cutoffs should come in pairs')

    for it in range(1, len(cutoffs), 2):
        try:
            float(cutoffs[it])
        except ValueError:
            raise argparse.ArgumentTypeError(
                'Error, second part of cutoff pair should be float'
            )

    parsed = {}

    for it in range(0, len(cutoffs), 2):
        parsed[cutoffs[it].capitalize()] = float(cutoffs[it + 1])

    return parsed


def extract_freq_func(uargs, save=True):
    '''
    Wrapper for command line frequency extract/print

    Parameters
    ----------
    uargs: argparser object
        command line arguments
    save: bool, default=True
        If True, saves data to file. If False, prints to stdout.

    Returns
    -------
        None
    '''
    from . import extractor as oe

    # Extract frequency information
    data = oe.FrequencyExtractor.extract(uargs.output_file)

    if not len(data):
        ut.red_exit(f'Cannot find frequencies in file {uargs.output_file}')

    if uargs.num is None:
        uargs.num = len(data[0]['energy (cm^-1)'])

    if not save:
        print('Frequencies (cm⁻¹) and intensities (km/mol)')
        for frq, inty in zip(
            data[0]['energy (cm^-1)'][:uargs.num],
            data[0]['IR Intensity (km mol^-1)'][:uargs.num]
        ):
            print(f'{frq:.5f} {inty:.5f}')
    else:
        # Save to new .csv file
        out_name = f'{uargs.output_file.stem}_frequencies.csv'
        with open(out_name, 'w') as f:
            f.write(
                f'# Frequencies and intensities from {uargs.output_file}\n')
            f.write('# Frequency (cm⁻¹), Intensity (km/mol)\n')
            for frq, inty in zip(
                data[0]['energy (cm^-1)'][:uargs.num],
                data[0]['IR Intensity (km mol^-1)'][:uargs.num]
            ):
                f.write(f'{frq:.5f}, {inty:.5f}\n')

        ut.cprint(f'Data written to {out_name}', 'cyan')

    return


def extract_pop_func(uargs, save=True) -> None:
    '''
    Wrapper for command line frequency extract/print

    Parameters
    ----------
    uargs: argparser object
        User arguments
    save: bool, default=False
        If True, saves data to file. If False, prints to stdout.
    Returns
    -------
        None
    '''
    import xyz_py as xyzp
    from . import extractor as oe

    if uargs.flavour in ['loewdin', 'lowdin']:
        data = oe.LoewdinPopulationExtractor.extract(
            uargs.output_file
        )
    elif uargs.flavour in ['mulliken']:
        try:
            data = oe.MullikenPopulationExtractorDensities.extract(
                uargs.output_file
            )
        except DataNotFoundError:
            data = oe.MullikenPopulationExtractorPopulations.extract(
                uargs.output_file
            )

    # Extract structure
    labels, coords = oe.get_coords(uargs.output_file)

    labels = xyzp.add_label_indices(
        labels,
        start_index=0,
        style='sequential'
    )

    if uargs.cutoffs:
        cutoffs = parse_cutoffs(uargs.cutoffs)
    else:
        cutoffs = {}

    # Generate dictionary of entities
    entities_dict = xyzp.find_entities(
        labels, coords, adjust_cutoff=cutoffs, non_bond_labels=uargs.no_bond
    )

    if len(data) > 1:
        ut.cprint(f'Found {len(data)} population blocks in file...', 'green')

    for it, datum in enumerate(data):

        if len(data) > 1:
            ut.cprint(f'\nBlock {it + 1:d}/{len(data):d}', 'green')

        # Calculate charge and spin density of each fragment
        ut.cprint(f'{uargs.flavour.capitalize()} Population Analysis', 'cyan')
        ut.cprint('Entity: Charge Spin', 'cyan')
        ut.cprint('-------------------', 'cyan')
        print()
        for entity_name, entities in entities_dict.items():
            for entity in entities:
                _chg = sum([datum[0][labels[ind]] for ind in entity])
                _spin = sum([datum[1][labels[ind]] for ind in entity])
                ut.cprint(
                    f'{entity_name}: {_chg:.4f}  {_spin:.4f}',
                    'cyan'
                )

    return


def plot_susc_func(uargs) -> None:
    '''
    Plots susceptibility data from output file

    Parameters
    ----------
    uargs: argparser object
        command line arguments

    Returns
    -------
        None

    '''
    from . import extractor as oe
    from . import data as d

    # Extract data from file before NEV section
    try:
        data = oe.SusceptibilityExtractor.extract(
            uargs.output_file,
            before='QDPT WITH NEVPT2 DIAGONAL ENERGIES'
        )
    except DataNotFoundError:
        data = []

    try:
        data_nev = oe.SusceptibilityExtractor.extract(
            uargs.output_file,
            after='QDPT WITH NEVPT2 DIAGONAL ENERGIES'
        )
    except DataNotFoundError:
        data_nev = []

    if not len(data + data_nev):
        ut.red_exit(
            f'Cannot find susceptibility data in {uargs.output_file}'
        )

    # Create susceptibilitydata objects, one per data section
    susc_data = d.SusceptibilityData.from_extractor_data(data)
    nev_susc_data = d.SusceptibilityData.from_extractor_data(data_nev)

    if uargs.nev_only:
        susc_data = []

    # Plot all data
    import matplotlib.pyplot as plt
    import matplotlib.colors as mcolors
    from . import plotter

    # Set font name and size
    ut.check_font_envvar()
    plt.rcParams['font.size'] = 10
    plt.rcParams['legend.fontsize'] = 9

    # Create list of colours for plotting
    if len(susc_data + nev_susc_data) == 1:
        colours = ['black']
    else:
        colours = list(mcolors.TABLEAU_COLORS.values())

    # Width of figure in inches
    # Make wider if multiple spectra to plot
    # to include legend
    if len(susc_data + nev_susc_data) > 1 and uargs.legend_separate:
        width = 6.
    else:
        width = 3.425
    # width in cm
    width_cm = 2.54 * width

    # Create figure and axis with
    # height based on golden ratio for 4 inch wide figure
    golden = (1 + np.sqrt(5))/2
    height = 4. / golden
    fig, ax = plt.subplots(
        1,
        1,
        num='Magnetic Susceptibility',
        figsize=(width, height)
    )

    for it, sd in enumerate(susc_data):
        plotter.plot_susceptibility(
            sd,
            uargs.y_style,
            linecolor=colours[it],
            xlim=uargs.xlim,
            ylim=uargs.ylim,
            save=False,
            show=False,
            fig=fig,
            ax=ax
        )

    for it, sd in enumerate(nev_susc_data):
        plotter.plot_susceptibility(
            sd,
            uargs.y_style,
            linecolor=colours[it + len(susc_data)],
            label=f'Calculated NEVPT2 ($H$ = {sd.field:.1f} Oe)',
            save=False,
            show=False,
            fig=fig,
            ax=ax
        )

    if uargs.exp_file is not None:

        exp_data = {'Temperature (K)': [], 'chi*T (cm3*K/mol)': []}
        with open(uargs.exp_file, newline='') as csvfile:
            reader = csv.DictReader(
                [row for row in csvfile if not row.startswith('#')],
                skipinitialspace=True
            )
            for row in reader:
                exp_data['Temperature (K)'].append(
                    float(row['Temperature (K)'])
                )
                exp_data['chi*T (cm3*K/mol)'].append(
                    float(row['chi*T (cm3*K/mol)'])
                )

        # Conversion factors to cm3 K mol^-1 from ...
        convs = {
            'A3 K': 1E24 / cst.AVOGADRO,
            'A3 mol-1 K': 1E24,
            'cm3 K': 1 / cst.AVOGADRO,
            'cm3 mol-1 K': 1,
            'emu K': 1 / (4 * np.pi * cst.AVOGADRO),
            'emu mol-1 K': 1 / (4 * np.pi)
        }

        ax.plot(
            exp_data['Temperature (K)'],
            [
                val * convs[uargs.esusc_units]
                for val in exp_data['chi*T (cm3*K/mol)']
            ],
            lw=0,
            marker='o',
            fillstyle='none',
            mew=.3,
            color='k',
            label='Experiment'
        )

    # Set x and y limits
    plotter.set_axlims(ax, 'x', uargs.xlim)
    plotter.set_axlims(ax, 'y', uargs.ylim)

    fig.tight_layout()

    if len(susc_data + nev_susc_data) > 1 or uargs.exp_file is not None:
        if uargs.legend_separate:
            fig.subplots_adjust(right=3.3/6.)
            fig.legend(loc='center right')
        else:
            ax.legend(loc='best')

    if _SAVE_CONV[uargs.plot]:
        savename = 'susceptibility.png'
        plt.savefig(savename, dpi=500)
        ut.cprint(f'Saved image to {savename}', 'cyan')
        ut.cprint(
            f'Use width={width:.1f} in. or {width_cm:.1f} cm',
            'cyan'
        )

    if _SHOW_CONV[uargs.plot]:
        plt.show()

    return


def plot_ailft_func(uargs) -> None:
    '''
    Plots AI-LFT orbital energies

    Parameters
    ----------
    uargs: argparser object
        command line arguments

    Returns
    -------
        None
    '''
    import matplotlib.pyplot as plt
    from . import plotter
    from . import extractor as oe

    # Set user specified font name
    ut.check_font_envvar()
    # Change matplotlib font size to be larger
    plt.rcParams['font.size'] = 10
    plt.rcParams['legend.fontsize'] = 9

    # Create extractor
    data = oe.AILFTOrbEnergyExtractor.extract(uargs.output_file)

    # Conversion factors from cm-1 to ...
    convs = {
        'cm-1': 1,
        'K': 1E24,
    }

    unit_labels = {
        'cm-1': r'cm^{-1}',
        'K': r'K',
    }

    for dit, dataframe in enumerate(data):

        if len(data) > 1:
            print()
            print(f'Section {dit+1:d}')
            print('---------')

        wfuncs = 100 * np.abs(dataframe['eigenvectors']) ** 2

        for e, wf in zip(dataframe['energies (cm^-1)'], wfuncs.T):
            print(f'E = {e * convs[uargs.units]} {uargs.units}:')
            for it, pc in enumerate(wf):
                if pc > 5.:
                    print(f'{pc:.1f} % {dataframe['orbitals'][it]}')
            if dit == len(wf):
                print('******')
            else:
                print()

        # mm_orbnames = ut.orbname_to_mathmode(dataframe['orbitals'])
        plotter.plot_ailft_orb_energies(
            dataframe['energies (cm^-1)'] * convs[uargs.units],
            groups=uargs.groups,
            occupations=uargs.occupancies,
            # labels=mm_orbnames, # convert these to %ages
            window_title=f'AI-LFT orbitals from {uargs.output_file}',
            y_unit=unit_labels[uargs.units],
            show=_SHOW_CONV[uargs.plot],
            save=_SAVE_CONV[uargs.plot],
            save_name=f'{uargs.output_file.stem}_ailft_orbs_set_{dit+1:d}.png'
        )

        plt.show()

    return


class CustomErrorArgumentParser(argparse.ArgumentParser):
    '''
    Custom ArgumentParser to handle errors and print usage\n
    This is required to avoid the default behavior of argparse which
    modifies the usage message when it prints, conflicting with the preset
    values used in the subparsers.
    '''
    def error(self, message):
        self.print_usage(sys.stderr)
        sys.stderr.write(f"error: {message}.\n")
        sys.stderr.write("       Use -h to see all options.\n")
        sys.exit(2)


def read_args(arg_list=None):
    '''
    Reader for command line arguments. Uses subReaders for individual programs

    Parameters
    ----------
    uargs: argparser object
        command line arguments

    Returns
    -------
        None
    '''

    description = 'OrcaTools (orto) - A package for working with Orca'

    epilog = 'Type\n'
    epilog += ut.cstring('orto <subprogram> -h\n', 'cyan')
    epilog += 'for help with a specific subprogram.\n'

    parser = CustomErrorArgumentParser(
        usage=ut.cstring('orto <subprogram> [options]', 'cyan'),
        description=description,
        epilog=epilog,
        formatter_class=argparse.RawTextHelpFormatter
    )

    parser._positionals.title = 'Subprograms'

    all_subparsers = parser.add_subparsers(dest='prog_grp')

    extract_subprog = all_subparsers.add_parser(
        'extract',
        description='Extract information from Orca file(s)',
        formatter_class=argparse.RawTextHelpFormatter,
        usage=ut.cstring('orto extract <section>', 'cyan'),
    )

    extract_parser = extract_subprog.add_subparsers(dest='extract_grp')

    extract_subprog._positionals.title = 'Sections'

    # If argument list is empty then call help function
    extract_subprog.set_defaults(func=lambda _: extract_subprog.print_help())

    extract_coords = extract_parser.add_parser(
        'coords',
        aliases=['coord'],
        description='Extracts coordinates from Orca output file',
        formatter_class=argparse.RawTextHelpFormatter,
        usage=ut.cstring('orto extract coords <output_file> [options]', 'cyan')
    )
    extract_coords._positionals.title = 'Mandatory Arguments'
    extract_coords.set_defaults(func=extract_coords_func)

    extract_coords.add_argument(
        'output_file',
        type=pathlib.Path,
        help='Orca output file name'
    )

    extract_coords.add_argument(
        '--type',
        type=str,
        help='Which coordinates to extract',
        choices=['opt', 'init'],
        default='init'
    )

    extract_coords.add_argument(
        '--index_style',
        type=str,
        help='Style of indexing used for output atom labels',
        choices=['per_element', 'sequential', 'sequential_orca', 'none'],
        default='per_element'
    )

    extract_freq = extract_parser.add_parser(
        'freq',
        description='Extracts frequencies',
        usage=ut.cstring('orto extract freq <output_file> [options]', 'cyan'),
        formatter_class=argparse.RawTextHelpFormatter
    )

    extract_freq._positionals.title = 'Mandatory Arguments'
    extract_freq.set_defaults(func=extract_freq_func)

    extract_freq.add_argument(
        'output_file',
        type=pathlib.Path,
        help='Orca output file name - must contain Frequencies section'
    )

    extract_freq.add_argument(
        '-n',
        '--num',
        type=int,
        default=None,
        help='Number of frequencies to print, default is all'
    )

    extract_hyperfine = extract_parser.add_parser(
        'hyperfine',
        description='Extracts hyperfine couplings from Orca output file',
        usage=ut.cstring(
            'orto extract hyperfine <output_file> [options]',
            'cyan'
        ),
        formatter_class=argparse.RawTextHelpFormatter
    )
    extract_hyperfine._positionals.title = 'Mandatory Arguments'
    extract_hyperfine.set_defaults(func=extract_hyperfine_func)

    extract_hyperfine.add_argument(
        'output_file',
        type=pathlib.Path,
        help='Path to/Name of Orca output file containing HYPERFINE section'
    )

    extract_hyperfine.add_argument(
        '--output_format',
        '-of',
        type=str,
        help='Format of outputted data file',
        choices=['txt', 'docx'],
        default='txt'
    )

    extract_gmatrix = extract_parser.add_parser(
        'gmatrix',
        description='Extracts coordinates from Orca output file',
        usage=ut.cstring('orto extract gmatrix <output_file> [options]', 'cyan'), # noqa
        formatter_class=argparse.RawTextHelpFormatter
    )
    extract_gmatrix._positionals.title = 'Mandatory Arguments'
    extract_gmatrix.set_defaults(func=extract_gmatrix_func)

    extract_gmatrix.add_argument(
        'output_file',
        type=pathlib.Path,
        help='Path to/Name of Orca output file containing G-MATRIX block'
    )

    extract_gmatrix.add_argument(
        '--type',
        type=str,
        help='Which G-MATRIX block to extract - if DFT then this option is redundant', # noqa
        choices=['total', 'S', 'L', 'eff', 'dft'],
        default='total'
    )

    extract_gmatrix.add_argument(
        '--output_format',
        '-of',
        type=str,
        help='Format of outputted data file',
        choices=['txt', 'docx'],
        default='txt'
    )

    extract_sfenergies = extract_parser.add_parser(
        'sf_energies',
        description='Extract Spin-Free energies from Orca output file',
        usage=ut.cstring('orto extract sf_energies <output_file> [options]', 'cyan'), # noqa
        formatter_class=argparse.RawTextHelpFormatter
    )
    extract_sfenergies._positionals.title = 'Mandatory Arguments'
    extract_sfenergies.set_defaults(func=extract_sf_energies_func)

    extract_sfenergies.add_argument(
        'output_file',
        type=pathlib.Path,
        help='Orca output file name containing TRANSITION ENERGIES block'
    )

    extract_sfenergies.add_argument(
        '--output_format',
        '-of',
        type=str,
        help='Format of outputted data file',
        choices=['txt', 'docx'],
        default='txt'
    )

    extract_soenergies = extract_parser.add_parser(
        'so_energies',
        description='Extract Spin-Free energies from Orca output file',
        usage=ut.cstring('orto extract so_energies <output_file> [options]', 'cyan'), # noqa
        formatter_class=argparse.RawTextHelpFormatter
    )
    extract_soenergies._positionals.title = 'Mandatory Arguments'
    extract_soenergies.set_defaults(func=extract_so_energies_func)

    extract_soenergies.add_argument(
        'output_file',
        type=pathlib.Path,
        help='Orca output file name containing Spin Orbit rel block'
    )

    gen_subprog = all_subparsers.add_parser(
        'gen',
        description='Generate inputs/jobs to/for Orca',
        formatter_class=argparse.RawTextHelpFormatter,
        usage=ut.cstring('orto gen <item>', 'cyan')
    )
    gen_subprog._positionals.title = 'Items'

    gen_parser = gen_subprog.add_subparsers(dest='gen_grp')

    # If argument list is empty then call help function
    gen_subprog.set_defaults(func=lambda _: gen_subprog.print_help())

    gen_trunc_molden = gen_parser.add_parser(
        'trunc_molden',
        description='Generate truncated molden file',
        usage=ut.cstring('orto gen trunc_molden <input_file> [options]', 'cyan'), # noqa
        formatter_class=argparse.RawTextHelpFormatter
    )
    gen_trunc_molden._positionals.title = 'Mandatory Arguments'
    gen_trunc_molden.set_defaults(func=gen_trunc_molden_func)
    gen_trunc_molden.add_argument(
        'input_file',
        type=pathlib.Path,
        help='Name of molden file to truncate'
    )

    gen_trunc_molden.add_argument(
        '--output_file',
        type=pathlib.Path,
        help=(
            'Name of truncated molden file\n'
            'If not specified, molden file will be truncated in place'
        ),
        default=None
    )

    gen_trunc_molden.add_argument(
        '--n_virt',
        type=int,
        default=100,
        help=(
            'Number of virtual orbitals to keep in truncated molden file\n'
            'Default: %(default)s'
        )
    )

    gen_orbs = gen_parser.add_parser(
        'orbs',
        aliases=['orbitals', 'mos', 'mo'],
        description='Generate molecular orbital cube file(s) from gbw file', # noqa
        usage=ut.cstring(
            'orto gen orbs <gbw_file> <orbs> <n_procs> [options]',
            'cyan'
        ),
        formatter_class=argparse.RawTextHelpFormatter
    )
    gen_orbs._positionals.title = 'Mandatory Arguments'
    gen_orbs.set_defaults(func=gen_orbs_func)

    gen_orbs.add_argument(
        'gbw_file',
        type=pathlib.Path,
        help='Orca gbw file name'
    )

    gen_orbs.add_argument(
        'orb_numbers',
        type=int,
        nargs='+',
        help=(
            'Orbital number(s) to generate cube file for'
            '(these are alpha spin by default)'
        )
    )

    gen_orbs.add_argument(
        'n_procs',
        type=int,
        help='Number of processors/cores used in calculation',
        default=1
    )

    gen_orbs.add_argument(
        '--n_pts',
        '-n',
        type=int,
        default=100,
        help=(
            'Number of points in each dimension of cube file\n'
            'Default: %(default)s'
        )
    )

    gen_orbs.add_argument(
        '--memory',
        '-mem',
        type=int,
        default=500,
        help=(
            'Per-core Memory to use in MB\n'
            'Default: %(default)s'
        )
    )

    gen_orbs.add_argument(
        '-om',
        '--orca_load',
        type=str,
        default='',
        help='Orca environment module (overrides ORTO_ORCA_LOAD envvar)'
    )

    gen_orbs.add_argument(
        '--beta',
        action='store_true',
        help=(
            'Plot beta orbitals instead of alpha (for open shell calculations)'
        )
    )

    gen_orbs.add_argument(
        '--no_sub',
        '-ns',
        action='store_true',
        help=(
            'Disables submission of job to queue'
        )
    )

    gen_spden = gen_parser.add_parser(
        'spdens',
        aliases=['spin_density', 'spden'],
        description='Generate spin density cube file from gbw file', # noqa
        usage=ut.cstring(
            'orto gen spdens <gbw_file> <n_procs> [options]',
            'cyan'
        ),
        formatter_class=argparse.RawTextHelpFormatter
    )
    gen_spden._positionals.title = 'Mandatory Arguments'
    gen_spden.set_defaults(func=gen_spden_func)

    gen_spden.add_argument(
        'gbw_file',
        type=pathlib.Path,
        help='Orca gbw file name'
    )

    gen_spden.add_argument(
        'n_procs',
        type=int,
        help='Number of processors/cores used in calculation',
        default=1
    )

    gen_spden.add_argument(
        '--n_pts',
        '-n',
        type=int,
        default=100,
        help=(
            'Number of points in each dimension of cube file\n'
            'Default: %(default)s'
        )
    )

    gen_spden.add_argument(
        '--memory',
        '-mem',
        type=int,
        default=500,
        help=(
            'Per-core Memory to use in MB\n'
            'Default: %(default)s'
        )
    )

    gen_spden.add_argument(
        '-om',
        '--orca_load',
        type=str,
        default='',
        help='Orca environment module (overrides ORTO_ORCA_LOAD envvar)'
    )

    gen_spden.add_argument(
        '--no_sub',
        '-ns',
        action='store_true',
        help=(
            'Disables submission of job to queue'
        )
    )

    gen_job = gen_parser.add_parser(
        'job',
        description=(
            'Generate submission script for orca calculation.\n'
            'Job script should be executed/submitted from its parent directory'
        ),
        usage=ut.cstring('orto gen job <input_file(s)> [options]', 'cyan'),
        formatter_class=argparse.RawTextHelpFormatter
    )
    gen_job._positionals.title = 'Mandatory Arguments'

    gen_job.set_defaults(func=gen_job_func)

    gen_job.add_argument(
        'input_files',
        metavar='<input_file(s)>',
        type=pathlib.Path,
        nargs='+',
        help='Orca input file name(s)'
    )

    gen_job.add_argument(
        '--n_procs',
        type=int,
        default=0,
        help=(
            'Number of cores requested in submission system.\n'
            ' This does not need to match the orca input, but must not be'
            'less\n. If not specified then value is read from input file.')
    )

    gen_job.add_argument(
        '--memory',
        '-mem',
        type=int,
        default=0,
        help=(
            'Per core memory requested in submission system (megabytes).\n'
            ' This does not need to match the orca input, but must not be'
            ' less.\n If not specified then value is read from input file.'
        )
    )

    gen_job.add_argument(
        '--no_sub',
        '-ns',
        action='store_true',
        help=(
            'Disables submission of job to queue'
        )
    )

    gen_job.add_argument(
        '--no_molden',
        '-nm',
        action='store_true',
        help=(
            'Disables orca_2mkl call for molden file generation after calculation' # noqa
        )
    )

    gen_job.add_argument(
        '--skip_xyz',
        '-sx',
        action='store_true',
        help=(
            'Disables xyz file format check'
        )
    )

    gen_job.add_argument(
        '--skip_structure',
        '-ss',
        action='store_true',
        help=(
            'Disables checking of structure in input file or xyz file'
        )
    )

    gen_job.add_argument(
        '-om',
        '--orca_load',
        type=str,
        default='',
        help='Orca environment module (overrides ORTO_ORCA_LOAD envvar)'
    )

    distort = gen_parser.add_parser(
        'distort',
        description='Distorts molecule along given normal mode',
        usage=ut.cstring('orto gen distort <output_file> <mode_number> [options]', 'cyan'), # noqa
        formatter_class=argparse.RawTextHelpFormatter
    )
    distort._positionals.title = 'Mandatory Arguments'

    distort.set_defaults(func=distort_func)

    distort.add_argument(
        'output_file',
        type=pathlib.Path,
        help='Orca output file name - must contain frequency section'
    )

    distort.add_argument(
        'mode_number',
        type=int,
        help='Mode to distort along - uses orca indexing and starts from zero'
    )

    distort.add_argument(
        '--scale',
        type=float,
        default=1,
        help=(
            'Number of units of distortion\n'
            'Default: %(default)s'
        )
    )

    gen_abs = gen_parser.add_parser(
        'abs',
        description='Generates absorption spectrum data from TDDFT/CI calculation output', # noqa
        usage=ut.cstring('orto gen abs <output_file> [options]', 'cyan'), # noqa
        formatter_class=argparse.RawTextHelpFormatter
    )
    gen_abs._positionals.title = 'Mandatory Arguments'

    gen_abs.set_defaults(func=lambda x: plot_abs_func(x, save_data_only=True))

    gen_abs.add_argument(
        'output_file',
        type=pathlib.Path,
        nargs='+',
        help='Orca output file name'
    )

    gen_abs.add_argument(
        '--intensities',
        '-i',
        type=str,
        choices=['velocity', 'electric', 'semi-classical'],
        default='electric',
        help='Type of intensity to plot (orca_mapspc uses electric)'
    )

    gen_abs.add_argument(
        '--linewidth',
        '-lw',
        type=float,
        default=1,
        help=(
            'Width of signal (FWHM for Gaussian, Width for Lorentzian),'
            ' in same unit as x axis'
        )
    )

    gen_abs.add_argument(
        '--lineshape',
        '-ls',
        type=str,
        choices=['gaussian', 'lorentzian'],
        default='lorentzian',
        help='Lineshape to use for each signal'
    )

    gen_abs.add_argument(
        '--x_unit',
        '-xu',
        type=str,
        choices=['energy', 'wavelength', 'wavenumber'],
        default='energy',
        help='x units to use for spectrum'
    )

    gen_abs.add_argument(
        '--x_shift',
        type=float,
        default=None,
        nargs='+',
        help=(
            'Shift spectrum by this amount in x units\n'
            'Default: 0.'
        )
    )

    gen_abs.add_argument(
        '--no_trim',
        action='store_true',
        default=False,
        help=(
            'Do not trim spectrum to non-zero oscillator strength\n'
            'Default: %(default)s'
        )
    )

    gen_abs.add_argument(
        '--xlim',
        nargs=2,
        default=['auto', 'auto'],
        help='x limits of spectrum'
    )

    gen_abs.add_argument(
        '--normalise',
        '-n',
        action='store_true',
        default=False,
        help=(
            'Normalises spectrum to maximum value\n'
            'Default: %(default)s'
        )
    )

    gen_abs.add_argument(
        '--zero_osc',
        default=1E-7,
        type=float,
        help=(
            'Oscillator strengths below this value are treated as zero\n'
            'Default: %(default)s'
        )
    )

    gen_abs.add_argument(
        '--unique_names',
        action='store_true',
        default=False,
        help=(
            'Attempt to shorten file names in plot legend\n'
            'Default: %(default)s'
        )
    )

    gen_abs.add_argument(
        '--trim_transitions',
        '-tt',
        type=int,
        default=None,
        help=(
            'Trim to first n transitions\n'
            'where n is an integer\n'
            'Default: %(default)s'
        )
    )

    gen_xes = gen_parser.add_parser(
        'xes',
        description='Generates XES data from SGS-DEF calculation output',
        usage=ut.cstring('orto gen xes <output_file> [options]', 'cyan'),
        formatter_class=argparse.RawTextHelpFormatter
    )
    gen_xes._positionals.title = 'Mandatory Arguments'

    gen_xes.set_defaults(func=lambda x: plot_xes_func(x, save_data_only=True))

    gen_xes.add_argument(
        'output_file',
        type=pathlib.Path,
        nargs='+',
        help='Orca output file name'
    )

    gen_xes.add_argument(
        '--intensities',
        '-i',
        type=str,
        choices=['velocity', 'electric', 'semi-classical'],
        default='electric',
        help='Type of intensity to plot (orca_mapspc uses electric)'
    )

    gen_xes.add_argument(
        '--linewidth',
        '-lw',
        type=float,
        default=1,
        help=(
            'Width of signal (FWHM for Gaussian, Width for Lorentzian),'
            ' in same unit as x axis'
        )
    )

    gen_xes.add_argument(
        '--lineshape',
        '-ls',
        type=str,
        choices=['gaussian', 'lorentzian'],
        default='lorentzian',
        help='Lineshape to use for each signal'
    )

    gen_xes.add_argument(
        '--x_unit',
        '-xu',
        type=str,
        choices=['energy', 'wavelength', 'wavenumber'],
        default='energy',
        help='x units to use for spectrum'
    )

    gen_xes.add_argument(
        '--x_shift',
        type=float,
        default=None,
        nargs='+',
        help=(
            'Shift spectrum by this amount in x units\n'
            'Default: 0.'
        )
    )

    gen_xes.add_argument(
        '--no_trim',
        action='store_true',
        default=False,
        help=(
            'Do not trim spectrum to non-zero oscillator strength\n'
            'Default: %(default)s'
        )
    )

    gen_xes.add_argument(
        '--xlim',
        nargs=2,
        default=['auto', 'auto'],
        help='x limits of spectrum'
    )

    gen_xes.add_argument(
        '--normalise',
        '-n',
        action='store_true',
        default=False,
        help=(
            'Normalises spectrum to maximum value\n'
            'Default: %(default)s'
        )
    )

    gen_xes.add_argument(
        '--zero_osc',
        default=1E-7,
        type=float,
        help=(
            'Oscillator strengths below this value are treated as zero\n'
            'Default: %(default)s'
        )
    )

    gen_xes.add_argument(
        '--unique_names',
        action='store_true',
        default=False,
        help=(
            'Attempt to shorten file names in plot legend\n'
            'Default: %(default)s'
        )
    )

    gen_xes.add_argument(
        '--trim_transitions',
        '-tt',
        type=int,
        default=None,
        help=(
            'Trim to first n transitions\n'
            'where n is an integer\n'
            'Default: %(default)s'
        )
    )

    plot_subprog = all_subparsers.add_parser(
        'plot',
        description='Plot data from orca file',
        formatter_class=argparse.RawTextHelpFormatter,
        usage=ut.cstring('orto plot <section>', 'cyan'),
    )
    plot_subprog._positionals.title = 'Mandatory Arguments'

    plot_parser = plot_subprog.add_subparsers(dest='plot_grp')

    # If argument list is empty then call help function
    plot_subprog.set_defaults(func=lambda _: plot_subprog.print_help())

    plot_abs = plot_parser.add_parser(
        'abs',
        description='Plots absorption spectrum from TDDFT/CI calculation output', # noqa
        usage=ut.cstring('orto plot abs <output_file> [options]', 'cyan'),
        formatter_class=argparse.RawTextHelpFormatter
    )
    plot_abs._positionals.title = 'Mandatory Arguments'

    plot_abs.set_defaults(func=plot_abs_func)

    plot_abs.add_argument(
        'output_file',
        type=pathlib.Path,
        nargs='+',
        help='Orca output file name'
    )

    plot_abs.add_argument(
        '--intensities',
        '-i',
        type=str,
        choices=['velocity', 'electric', 'semi-classical'],
        default='electric',
        help='Type of intensity to plot (orca_mapspc uses electric)'
    )

    plot_abs.add_argument(
        '--linewidth',
        '-lw',
        type=float,
        default=1,
        help=(
            'Width of signal (FWHM for Gaussian, Width for Lorentzian),'
            ' in same unit as x axis'
        )
    )

    plot_abs.add_argument(
        '--osc_style',
        type=str,
        default='separate',
        help=(
            'Style of oscillators to plot\n'
            ' - \'separate\' plots oscillator strengths as stems on separate axis\n' # noqa
            ' - \'combined\' plots oscillator strengths on intensity axis\n'
            ' - \'off\' does not plot oscillator strengths\n'
            'Default: %(default)s'
        )
    )

    plot_abs.add_argument(
        '--plot',
        '-p',
        choices=['on', 'show', 'save', 'off'],
        metavar='<str>',
        type=str,
        default='on',
        help=(
            'Controls plot appearance/save \n'
            ' - \'on\' shows and saves the plots\n'
            ' - \'show\' shows the plots\n'
            ' - \'save\' saves the plots\n'
            ' - \'off\' neither shows nor saves\n'
            'Default: %(default)s'
        )
    )

    plot_abs.add_argument(
        '--lineshape',
        '-ls',
        type=str,
        choices=['gaussian', 'lorentzian'],
        default='lorentzian',
        help='Lineshape to use for each signal'
    )

    plot_abs.add_argument(
        '--x_unit',
        '-xu',
        type=str,
        choices=['energy', 'wavelength', 'wavelength_rev', 'wavenumber'],
        default='energy',
        help='x units to use for spectrum (rev = reversed axis direction)'
    )

    plot_abs.add_argument(
        '--x_shift',
        type=float,
        default=None,
        nargs='+',
        help=(
            'Shift spectrum by this amount in x units\n'
            'Default: 0.'
        )
    )

    plot_abs.add_argument(
        '--xlim',
        nargs=2,
        default=['auto', 'auto'],
        help='x limits of spectrum'
    )

    plot_abs.add_argument(
        '--ylim',
        nargs=2,
        default=[0., 'auto'],
        help='Epsilon limits of spectrum in cm^-1 mol^-1 L'
    )

    plot_abs.add_argument(
        '--normalise',
        '-n',
        action='store_true',
        default=False,
        help=(
            'Normalises spectrum to maximum value\n'
            'Default: %(default)s'
        )
    )

    plot_abs.add_argument(
        '--zero_osc',
        default=1E-7,
        type=float,
        help=(
            'Oscillator strengths below this value are treated as zero\n'
            'Default: %(default)s'
        )
    )

    plot_abs.add_argument(
        '--legend',
        type=str,
        nargs='+',
        default=None,
        help=(
            'Legend labels for each spectrum plotted\n'
            'If not provided, file names are used\n'
        )
    )

    plot_abs.add_argument(
        '--trim_transitions',
        '-tt',
        type=int,
        default=None,
        help=(
            'Trim to first n transitions\n'
            'where n is an integer\n'
            'Default: %(default)s'
        )
    )

    plot_abs.add_argument(
        '--no_trim',
        action='store_true',
        default=False,
        help=(
            'Do not trim spectrum to non-zero oscillator strength\n'
            'Default: %(default)s'
        )
    )
    plot_xes = plot_parser.add_parser(
        'xes',
        description='Plots X-ray emission spectrum from SGS-DFT calculation output', # noqa
        usage=ut.cstring('orto plot xes <output_file> [options]', 'cyan'),
        formatter_class=argparse.RawTextHelpFormatter
    )
    plot_xes._positionals.title = 'Mandatory Arguments'

    plot_xes.set_defaults(func=plot_xes_func)

    plot_xes.add_argument(
        'output_file',
        type=pathlib.Path,
        nargs='+',
        help='Orca output file name'
    )

    plot_xes.add_argument(
        '--intensities',
        '-i',
        type=str,
        choices=['velocity', 'electric', 'semi-classical'],
        default='electric',
        help='Type of intensity to plot (orca_mapspc uses electric)'
    )

    plot_xes.add_argument(
        '--linewidth',
        '-lw',
        type=float,
        default=1,
        help=(
            'Width of signal (FWHM for Gaussian, Width for Lorentzian),'
            ' in same unit as x axis'
        )
    )

    plot_xes.add_argument(
        '--osc_style',
        type=str,
        default='separate',
        help=(
            'Style of oscillators to plot\n'
            ' - \'separate\' plots oscillator strengths as stems on separate axis\n' # noqa
            ' - \'combined\' plots oscillator strengths on intensity axis\n'
            ' - \'off\' does not plot oscillator strengths\n'
            'Default: %(default)s'
        )
    )

    plot_xes.add_argument(
        '--plot',
        '-p',
        choices=['on', 'show', 'save', 'off'],
        metavar='<str>',
        type=str,
        default='on',
        help=(
            'Controls plot appearance/save \n'
            ' - \'on\' shows and saves the plots\n'
            ' - \'show\' shows the plots\n'
            ' - \'save\' saves the plots\n'
            ' - \'off\' neither shows nor saves\n'
            'Default: %(default)s'
        )
    )

    plot_xes.add_argument(
        '--lineshape',
        '-ls',
        type=str,
        choices=['gaussian', 'lorentzian'],
        default='lorentzian',
        help='Lineshape to use for each signal'
    )

    plot_xes.add_argument(
        '--x_unit',
        '-xu',
        type=str,
        choices=['energy', 'wavelength', 'wavelength_rev', 'wavenumber'],
        default='energy',
        help='x units to use for spectrum (rev = reversed axis direction)'
    )

    plot_xes.add_argument(
        '--x_shift',
        type=float,
        default=None,
        nargs='+',
        help=(
            'Shift spectrum by this amount in x units\n'
            'Default: 0.'
        )
    )

    plot_xes.add_argument(
        '--xlim',
        nargs=2,
        default=['auto', 'auto'],
        help='x limits of spectrum'
    )

    plot_xes.add_argument(
        '--ylim',
        nargs=2,
        default=[0., 'auto'],
        help='Epsilon limits of spectrum in cm^-1 mol^-1 L'
    )

    plot_xes.add_argument(
        '--normalise',
        '-n',
        action='store_true',
        default=False,
        help=(
            'Normalises spectrum to maximum value\n'
            'Default: %(default)s'
        )
    )

    plot_xes.add_argument(
        '--zero_osc',
        default=1E-7,
        type=float,
        help=(
            'Oscillator strengths below this value are treated as zero\n'
            'Default: %(default)s'
        )
    )

    plot_xes.add_argument(
        '--legend',
        type=str,
        nargs='+',
        default=None,
        help=(
            'Legend labels for each spectrum plotted\n'
            'If not provided, file names are used\n'
        )
    )

    plot_xes.add_argument(
        '--trim_transitions',
        '-tt',
        type=int,
        default=None,
        help=(
            'Trim to first n transitions\n'
            'where n is an integer\n'
            'Default: %(default)s'
        )
    )

    plot_xes.add_argument(
        '--no_trim',
        action='store_true',
        default=False,
        help=(
            'Do not trim spectrum to non-zero oscillator strength\n'
            'Default: %(default)s'
        )
    )

    plot_ailft = plot_parser.add_parser(
        'ailft_orbs',
        aliases=['ailft_orb'],
        description='Plots AI-LFT orbital energies from output file',
        usage=ut.cstring(
            'orto plot ailft_orbs <output_file> [options]',
            'cyan'
        ),
        formatter_class=argparse.RawTextHelpFormatter
    )
    plot_ailft._positionals.title = 'Mandatory Arguments'

    plot_ailft.set_defaults(func=plot_ailft_func)

    plot_ailft.add_argument(
        'output_file',
        type=pathlib.Path,
        help='Orca output file name'
    )

    plot_ailft.add_argument(
        '--groups',
        '-g',
        metavar='<str>',
        nargs='+',
        type=int,
        default=None,
        help=(
            'Group indices for each orbital. e.g. 1 1 1 2 2\n'
            'Controls x-staggering of orbitals'
        )
    )

    plot_ailft.add_argument(
        '--occupancies',
        '-o',
        metavar='<str>',
        nargs='+',
        type=int,
        default=None,
        help=(
            'Occupation number of each orbital\n Adds electrons to each orb\n'
            'Must specify occupation of every orbital as 2, 1, -1, or 0'
        )
    )

    plot_ailft.add_argument(
        '--plot',
        '-p',
        choices=['on', 'show', 'save', 'off'],
        metavar='<str>',
        type=str,
        default='on',
        help=(
            'Controls plot appearance/save \n'
            ' - \'on\' shows and saves the plots\n'
            ' - \'show\' shows the plots\n'
            ' - \'save\' saves the plots\n'
            ' - \'off\' neither shows nor saves\n'
            'Default: %(default)s'
        )
    )

    plot_ailft.add_argument(
        '--units',
        '-u',
        choices=[
            'cm-1',
            'K'
        ],
        metavar='<str>',
        type=str,
        default='cm-1',
        help=(
            'Controls energy units of plot\n'
            'Default: %(default)s'
        )
    )

    plot_susc = plot_parser.add_parser(
        'susc',
        description='Plots susceptibility data from output file',
        usage=ut.cstring('orto plot susc <output_file> [options]', 'cyan'),
        formatter_class=argparse.RawTextHelpFormatter
    )
    plot_susc._positionals.title = 'Mandatory Arguments'

    plot_susc.set_defaults(func=plot_susc_func)

    plot_susc.add_argument(
        'output_file',
        type=pathlib.Path,
        help='Orca output file name'
    )

    plot_susc.add_argument(
        '--y_style',
        '-ys',
        type=str,
        metavar='<str>',
        choices=['X', 'XT', '1/X'],
        default='XT'
    )

    plot_susc.add_argument(
        '--xlim',
        nargs=2,
        default=['auto', 'auto'],
        help='x limits of plot'
    )

    plot_susc.add_argument(
        '--ylim',
        nargs=2,
        default=[0., 'auto'],
        help='y limits of plot'
    )

    plot_susc.add_argument(
        '--susc_units',
        '-su',
        choices=[
            'emu mol-1 K',
            'emu K',
            'cm3 mol-1 K',
            'cm3 K',
            'A3 mol-1 K',
            'A3 K'
        ],
        metavar='<str>',
        type=str,
        default='cm3 mol-1 K',
        help=(
            'Controls susceptibility units of calculated data \n'
            '(wrap with "")\n'
            'Default: %(default)s'
        )
    )

    plot_susc.add_argument(
        '--plot',
        '-p',
        choices=['on', 'show', 'save', 'off'],
        metavar='<str>',
        type=str,
        default='on',
        help=(
            'Controls plot appearance/save \n'
            ' - \'on\' shows and saves the plots\n'
            ' - \'show\' shows the plots\n'
            ' - \'save\' saves the plots\n'
            ' - \'off\' neither shows nor saves\n'
            'Default: %(default)s'
        )
    )

    plot_susc.add_argument(
        '--exp_file',
        type=str,
        help=(
            'Experimental datafile as .csv with two columns:\n'
            '1. "Temperature (K)"\n'
            '2. "chi*T (UNITS FROM --esusc_units)"\n'
        )
    )

    plot_susc.add_argument(
        '--esusc_units',
        '-esu',
        choices=[
            'emu mol-1 K',
            'emu K',
            'cm3 mol-1 K',
            'cm3 K',
            'A3 mol-1 K',
            'A3 K'
        ],
        metavar='<str>',
        type=str,
        default='cm3 mol-1 K',
        help=(
            'Controls susceptibility units of experimental data \n'
            '(wrap with "")\n'
            'Default: %(default)s'
        )
    )

    plot_susc.add_argument(
        '--nev_only',
        '-no',
        action='store_true',
        help='Only plot NEVPT2 data'
    )

    plot_susc.add_argument(
        '--legend_separate',
        '-ls',
        action='store_true',
        help='Plot legend next to plot'
    )

    plot_susc.add_argument(
        '--quiet',
        action='store_true',
        help='Suppresses text output'
    )

    plot_ir = plot_parser.add_parser(
        'ir',
        description='Plots infrared spectrum from frequency calculation output', # noqa
        usage=ut.cstring('orto plot ir <output_file> [options]', 'cyan'),
        formatter_class=argparse.RawTextHelpFormatter
    )
    plot_ir._positionals.title = 'Mandatory Arguments'

    plot_ir.set_defaults(func=plot_ir_func)

    plot_ir.add_argument(
        'output_file',
        type=pathlib.Path,
        help='Orca output file name'
    )

    plot_ir.add_argument(
        '--plot',
        '-p',
        choices=['on', 'show', 'save', 'off'],
        metavar='<str>',
        type=str,
        default='on',
        help=(
            'Controls plot appearance/save \n'
            ' - \'on\' shows and saves the plots\n'
            ' - \'show\' shows the plots\n'
            ' - \'save\' saves the plots\n'
            ' - \'off\' neither shows nor saves\n'
            'Default: %(default)s'
        )
    )

    plot_ir.add_argument(
        '--linewidth',
        '-lw',
        type=float,
        default=5,
        help=(
            'Width of signal (FWHM for Gaussian, Width for Lorentzian),'
            ' in same unit as plot x unit'
        )
    )

    plot_ir.add_argument(
        '--lineshape',
        '-ls',
        type=str,
        choices=['gaussian', 'lorentzian'],
        default='lorentzian',
        help='Lineshape to use for each signal'
    )

    print_subprog = all_subparsers.add_parser(
        'print',
        description='Print information from Orca file to screen',
        formatter_class=argparse.RawTextHelpFormatter,
        usage=ut.cstring('orto print <section>', 'cyan')
    )
    print_subprog._positionals.title = 'Sections'

    print_parser = print_subprog.add_subparsers(dest='print_grp')

    # If argument list is empty then call help function
    print_subprog.set_defaults(func=lambda _: print_subprog.print_help())

    print_hyperfine = print_parser.add_parser(
        'hyperfine',
        description='Prints hyperfine couplings from Orca output file',
        usage=ut.cstring(
            'orto print hyperfine <output_file> [options]',
            'cyan'
        ),
        formatter_class=argparse.RawTextHelpFormatter
    )
    print_hyperfine._positionals.title = 'Mandatory Arguments'

    print_hyperfine.set_defaults(
        func=lambda x: extract_hyperfine_func(x, save=False)
    )

    print_hyperfine.add_argument(
        'output_file',
        type=pathlib.Path,
        help='Path to/Name of Orca output file containing HYPERFINE section'
    )

    print_gmatrix = print_parser.add_parser(
        'gmatrix',
        description='Prints g matrix from Orca output file',
        usage=ut.cstring('orto print gmatrix <output_file> [options]', 'cyan'),
        formatter_class=argparse.RawTextHelpFormatter
    )
    print_gmatrix._positionals.title = 'Mandatory Arguments'

    print_gmatrix.set_defaults(
        func=lambda x: extract_gmatrix_func(x, save=False)
    )

    print_gmatrix.add_argument(
        'output_file',
        type=pathlib.Path,
        help='Path to/Name of Orca output file containing G-MATRIX block'
    )

    print_gmatrix.add_argument(
        '--type',
        type=str,
        help='Which G-MATRIX block to extract.',
        choices=['total', 'S', 'L', 'eff'],
        default='total'
    )

    print_freq = print_parser.add_parser(
        'freq',
        description='Prints frequencies from Orca output file',
        usage=ut.cstring('orto print freq <output_file> [options]', 'cyan'),
        formatter_class=argparse.RawTextHelpFormatter
    )
    print_freq._positionals.title = 'Mandatory Arguments'

    print_freq.set_defaults(func=lambda x: extract_freq_func(x, save=False))

    print_freq.add_argument(
        'output_file',
        type=pathlib.Path,
        help='Path to/Name of Orca output file containing FREQUENCIES section'
    )

    print_freq.add_argument(
        '-n',
        '--num',
        type=int,
        default=None,
        help='Number of frequencies to print, default is all'
    )

    print_orbs = print_parser.add_parser(
        'lorbs',
        description='Prints Loewdin orbital compositions',
        usage=ut.cstring('orto print orbs <output_file> [options]', 'cyan'),
        formatter_class=argparse.RawTextHelpFormatter
    )
    print_orbs._positionals.title = 'Mandatory Arguments'

    print_orbs.set_defaults(
        func=lambda x: extract_orbs_func(x, save=False)
    )

    print_orbs.add_argument(
        'output_file',
        type=pathlib.Path,
        help=(
            'Path to/Name of Orca output file containing\n'
            'one of the following sections\n'
            '   LOEWDIN ORBITAL-COMPOSITIONS\n'
            '   LOEWDIN REDUCED ORBITAL POPULATIONS PER MO\n'
            '   LOEWDIN ORBITAL POPULATIONS PER MO\n'
        )
    )

    print_orbs.add_argument(
        '-t',
        '--threshold',
        type=float,
        default=1.,
        help=(
            'Orbitals with contribution >= threshold are printed.\n'
            'Default: %(default)s'
        )
    )

    print_orbs.add_argument(
        '-e',
        '--elements',
        type=str,
        default=atomic_elements,
        nargs='+',
        help='Only print contributions from specified element(s) e.g. Ni'
    )

    print_orbs.add_argument(
        '-s',
        '--shell',
        type=int,
        nargs='+',
        default=[1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 19, 20], # noqa
        help=(
            'Only print contributions from specified shell numbers e.g. 3\n'
            'Default: %(default)s'
        )
    )

    print_orbs.add_argument(
        '-o',
        '--orb',
        type=str,
        default=['s', 'p', 'd', 'f'],
        choices=['s', 'p', 'd', 'f'],
        nargs='+',
        help=(
            'Only print contributions from specified orbital(s) e.g. d\n'
            'Default: %(default)s'
        )
    )

    print_orbs.add_argument(
        '-f',
        '--flavour',
        type=str,
        choices=['first_match', 'orb_pop', 'redorb_pop', 'orb_comp'],
        default='first_match',
        help=(
            'Which section to print from the output file\n'
            'orb_comp: Loewdin orbital compositions\n'
            'orb_pop: Loewdin orbital populations\n'
            'redorb_pop: Loewdin reduced orbital populations\n'
        )
    )

    print_orbs.add_argument(
        '--spin',
        type=str,
        choices=['UP', 'DOWN'],
        default='UP',
        help=(
            'Which spin to print from the output file\n'
            'If closed shell, then this is ignored\n'
            'Default: %(default)s'
        )
    )

    orb_group = print_orbs.add_mutually_exclusive_group(required=False)

    orb_group.add_argument(
        '-a',
        '--active',
        action='store_true',
        help=(
            'Only print active orbitals (0 < occupation < 2)'
        )
    )

    def gte_zero(x):
        '''
        Custom type for argparse to ensure that the input
        \nis greater than or equal to zero
        '''
        value = int(x)
        if value < 0:
            raise argparse.ArgumentTypeError(
                f'{x} is not a valid index (must be >= 0)'
            )
        return value

    orb_group.add_argument(
        '-n',
        '--num',
        nargs='+',
        type=gte_zero,
        metavar='NUMBER',
        default=None,
        help=(
            'Print specified orbitals using index starting from 0\n'
            '(same as Orca)\n'
        )
    )

    orb_group.add_argument(
        '-hl',
        '--homo_lumo',
        nargs='?',
        type=int,
        metavar='NUMBER',
        default=None,
        help=(
            'Print specified number of orbitals either side of HOMO and LUMO'
        )
    )

    print_pop = print_parser.add_parser(
        'pop',
        description='Prints population analysis (spin, charge), and groups by fragment', # noqa
        usage=ut.cstring('orto print pop <output_file> [options]', 'cyan'),
        formatter_class=argparse.RawTextHelpFormatter
    )
    print_pop._positionals.title = 'Mandatory Arguments'

    print_pop.set_defaults(func=lambda x: extract_pop_func(x, save=False))

    print_pop.add_argument(
        'output_file',
        type=pathlib.Path,
        help='Orca output file name - must contain population analysis section'
    )

    print_pop.add_argument(
        '--flavour',
        '-f',
        type=str,
        choices=['lowdin', 'loewdin', 'mulliken'],
        default='mulliken',
        help='Type of population analysis to print'
    )

    print_pop.add_argument(
        '--no_bond',
        '-nb',
        type=str,
        default=[],
        nargs='+',
        metavar='symbol',
        help='Atom labels specifying atoms to which no bonds can be formed'
    )

    print_pop.add_argument(
        '--cutoffs',
        type=str,
        nargs='+',
        metavar='symbol number',
        help='Modify cutoff used to define bonds between atoms'
    )

    # If argument list is empty then call help function
    parser.set_defaults(func=lambda _: parser.print_help())

    # select parsing option based on sub-parser
    args = parser.parse_args(arg_list)
    args.func(args)
    return args


def main():
    read_args()
    return
