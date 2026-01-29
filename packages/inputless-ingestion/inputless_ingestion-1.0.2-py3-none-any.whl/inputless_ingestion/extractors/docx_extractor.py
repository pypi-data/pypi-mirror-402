"""
Microsoft Word document extractor for DOCX and DOC files.
"""

from typing import Dict, Any
from docx import Document
from pathlib import Path


class DOCXExtractor:
    """Extract text and metadata from DOCX files."""
    
    async def extract_text(self, file_path: str) -> str:
        """
        Extract text from DOCX.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text content
        """
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs]
        return "\n".join(paragraphs)
    
    async def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract DOCX metadata.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary with metadata
        """
        doc = Document(file_path)
        core_props = doc.core_properties
        
        return {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "created_date": str(core_props.created) if core_props.created else "",
            "modified_date": str(core_props.modified) if core_props.modified else "",
        }
    
    async def extract_content(self, file_path: str) -> Dict[str, Any]:
        """
        Extract both text and metadata from DOCX.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary with text and metadata
        """
        text = await self.extract_text(file_path)
        metadata = await self.extract_metadata(file_path)
        
        return {
            "text": text,
            "metadata": metadata,
        }
