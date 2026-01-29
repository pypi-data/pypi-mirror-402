"""DOCX file format handling and table extraction."""
import csv
import datetime
import json

import openpyxl
try:
    import xlwt
except ImportError:  # pragma: no cover - optional dependency
    xlwt = None
try:
    from docx import Document
    from docx.oxml.simpletypes import ST_Merge
    from docx.table import _Cell
except ImportError:  # pragma: no cover - optional dependency
    Document = None
    ST_Merge = None
    _Cell = None


def __extract_table(table, strip_space=False):
    """Extracts table data from table object"""
    results = []
    n = 0
    for tr in table._tbl.tr_lst:
        r = []
        for tc in tr.tc_lst:
            for grid_span_idx in range(tc.grid_span):
                if tc.vMerge == ST_Merge.CONTINUE:
                    value = results[n - 1][len(r) - 1]
                elif grid_span_idx > 0:
                    value = r[-1]
                else:
                    cell = _Cell(tc, table)
                    value = cell.text.replace("\n", " ")
                if strip_space:
                    value = value.strip()
                r.append(value)
        results.append(r)
#        print(r)
        n += 1
    return results


def __store_table(tabdata, filename, output_format="csv"):
    """Saves table data as csv file."""
    if output_format == "csv":
        with open(filename, "w", encoding='utf8') as f:
            w = csv.writer(f, delimiter=",")
            for row in tabdata:
                w.writerow(row)
    elif output_format == 'tsv':
        with open(filename, 'w', encoding='utf8') as f:
            w = csv.writer(f, delimiter='\t')
            for row in tabdata:
                w.writerow(row)
    elif output_format == 'xls':
        if xlwt is None:
            raise RuntimeError("xlwt is required for XLS output")
        workbook = xlwt.Workbook()
        __xls_table_to_sheet(tabdata, workbook.add_sheet("0"))
        workbook.save(filename)
    elif output_format == "xlsx":
        workbook = openpyxl.Workbook()
        __xlsx_table_to_sheet(tabdata, workbook.create_sheet("0"))
        workbook.save(filename)

def __xls_table_to_sheet(table, ws):
    rn = 0
    for row in table:
        cn = 0
        for c in row:
            ws.write(rn, cn, c)
            cn += 1
        rn += 1
    return ws


def __xlsx_table_to_sheet(table, ws):
    rn = 0
    for row in table:
        ws.append(row)
        rn += 1
    return ws


def extract_docx_tables(filename, strip_space=True):
    """Extracts table from .DOCX files"""
    if Document is None:
        raise RuntimeError("python-docx is required for DOCX processing")
    tables = []
    document = Document(filename)
    n = 0
    for table in document.tables:
        n += 1
        info = {}
        info['id'] = n
        info['num_cols'] = len(table.columns)
        info['num_rows'] = len(table.rows)
        info['style'] = table.style.name
        tdata = __extract_table(table, strip_space=strip_space)
        info['data'] = tdata
        tables.append(info)
    return tables





def extract(filename, output_format="csv", sizefilter=0, singlefile=False,
           output=None, strip_space=True):
    """Extracts tables from csv files and saves them as csv, xls or xlsx files."""
    tables = extract_docx_tables(filename, strip_space=strip_space)
    name = filename.rsplit(".", 1)[0]
    output_format = output_format.lower()
    n = 0
    lfilter = int(sizefilter)
    if singlefile:
        if output_format == "xls":
            if xlwt is None:
                raise RuntimeError("xlwt is required for XLS output")
            workbook = xlwt.Workbook()
            for t in tables:
                if lfilter >= len(t):
                    continue
                n += 1
                __xls_table_to_sheet(t['data'], workbook.add_sheet(str(n)))
            destname = output if output else f"{name}.{output_format}"
            workbook.save(destname)
        elif output_format == "xlsx":
            workbook = openpyxl.Workbook()
            for t in tables:
                if lfilter >= len(t):
                    continue
                n += 1
                __xlsx_table_to_sheet(t['data'], workbook.create_sheet(str(n)))
            destname = output if output else f"{name}.{output_format}"
            workbook.save(destname)
        elif output_format == "json":
            report = {'filename': filename,
                      'timestamp': datetime.datetime.now().isoformat(),
                      'num_tables': len(tables),
                      'tables': tables}
            destname = output if output else f"{name}.{output_format}"
            with open(destname, 'w', encoding='utf8') as f:
                json.dump(report, f, ensure_ascii=False, indent=4)

    else:
        for t in tables:
            if lfilter >= len(t):
                continue
            n += 1
            destname = output if output else f"{name}_{n}.{output_format}"
            __store_table(t['data'], destname, output_format)


def analyze_docx(filename, extract_data=None, strip_space=True):
    """Analyzes docx file and extracts data if requested."""
    # extract_data parameter kept for API compatibility but not used
    if Document is None:
        raise RuntimeError("python-docx is required for DOCX processing")
    tableinfo = []
    document = Document(filename)
    n = 0
    for table in document.tables:
        n += 1
        info = {}
        info['id'] = n
        info['num_cols'] = len(table.columns)
        info['num_rows'] = len(table.rows)
        info['style'] = table.style.name
        tdata = __extract_table(table, strip_space=strip_space)
        info['data'] = tdata
        tableinfo.append(info)
    return tableinfo
