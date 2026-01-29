import io
from pathlib import Path
from typing import ClassVar, List, Union

from gwenflow.logger import logger
from gwenflow.readers.base import Reader
from gwenflow.types import Document


class DocxReader(Reader):
    trans: ClassVar[dict[int, int | None]] = {
        0x00A0: 0x20,
        0x202F: 0x20,
        0x2007: 0x20,
        0x200B: None,
        0x200C: None,
        0x200D: None,
        0xFEFF: None,
    }

    def get_text(self, file_obj) -> str:
        try:
            import docx
        except ImportError as e:
            raise ImportError("python-docx is not installed. Please install it with `pip install python-docx`") from e
        doc = docx.Document(file_obj)
        return "\n".join((p.text.translate(self.trans) if p.text else "") for p in doc.paragraphs)

    def get_tables(self, file_obj):
        try:
            import docx
        except ImportError as e:
            raise ImportError("python-docx is not installed. Please install it with `pip install python-docx`") from e
        doc = docx.Document(file_obj)

        tables = []
        for t in doc.tables:
            rows = []
            for r in t.rows:
                cells = []
                for c in r.cells:
                    txt = "\n".join(p.text for p in c.paragraphs) if c.paragraphs else ""
                    txt = txt.translate(self.trans) if txt else ""
                    cells.append(txt)
                rows.append(cells)
            tables.append(rows)
        return tables

    def read(self, file: Union[Path, io.BytesIO]) -> List[Document]:
        try:
            filename = self.get_file_name(file)
            content = self.get_file_content(file)

            data = content.getvalue() if isinstance(content, io.BytesIO) else content

            text = self.get_text(io.BytesIO(data))
            tables = self.get_tables(io.BytesIO(data))

            doc = Document(
                id=self.key(f"{filename}"),
                content=text,
                metadata={"filename": filename, "tables": tables},
            )
            return [doc]

        except Exception as e:
            logger.exception(f"Error reading file: {e}")
            return []
