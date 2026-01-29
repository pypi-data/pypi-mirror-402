import argparse
import gzip
import io
import ipaddress
import json
import os
import re
import sys
import tempfile
import time
import urllib.error
import urllib.request
import zipfile
import zlib
from dataclasses import dataclass
from typing import Dict, Iterable, List, Optional, Set, Tuple
import trafilatura
from colorama import Fore, Style, init as colorama_init
from publicsuffix2 import get_sld


# ========= #
# Constants #
# ========= #

APP_NAME = "iocscrape"
APP_VERSION = "0.2.1"
PROJECT_URL = "https://github.com/fwalbuloushi/iocscrape"
APP_DESC = "CTI tool to extract IOCs from CTI reports (URLs or files), and write them to an output file. Low-confidence items are grouped at the end."

DEFAULT_TIMEOUT_SEC = 15
DEFAULT_MAX_BYTES = 20 * 1024 * 1024  # 20MB
DEFAULT_REDIRECT_LIMIT = 5

IOC_TYPES_ORDER = ["url", "domain", "ipv4", "ipv6", "email", "md5", "sha1", "sha256", "cve"]

FILENAME_LIKE_EXTS = {
    "exe", "dll", "js", "vbs", "ps1", "bat", "cmd", "zip", "rar", "7z",
    "pdf", "doc", "docx", "xls", "xlsx", "ppt", "pptx",
    "png", "jpg", "jpeg", "gif", "svg", "css", "woff", "woff2", "ico",
    "conf", "cfg", "rule", "dit", "vmx",
    "py", "sh", "pl", "rs", "md", "so",     # ccTLDs that look like file extensions
}

STATIC_ASSET_EXTS = {"css", "png", "jpg", "jpeg", "gif", "svg", "woff", "woff2", "ico"}


# ================ #
# Cache / Datasets #
# ================ #

def _cache_dir() -> str:
    base = os.environ.get("XDG_CACHE_HOME") or os.path.join(os.path.expanduser("~"), ".cache")
    return os.path.join(base, APP_NAME)

CACHE_DIR = _cache_dir()
CACHE_META_FILE = os.path.join(CACHE_DIR, "meta.json")

PACKAGE_DIR = os.path.dirname(__file__)
BUNDLED_DATA_DIR = os.path.join(PACKAGE_DIR, "data")

# Vendor snapshots into iocscrape/data/
BUNDLED_PSL_FILE = os.path.join(BUNDLED_DATA_DIR, "public_suffix_list.dat")
BUNDLED_WARNINGLISTS_DIR = os.path.join(BUNDLED_DATA_DIR, "warninglists")  # expects .../lists/*.json

MISP_WARNINGLISTS_ZIP = "https://github.com/MISP/misp-warninglists/archive/refs/heads/main.zip"
PSL_URL = "https://publicsuffix.org/list/public_suffix_list.dat"

CACHE_PSL_FILE = os.path.join(CACHE_DIR, "public_suffix_list.dat")
CACHE_WARNINGLISTS_DIR = os.path.join(CACHE_DIR, "warninglists")


# =================== #
# LowConfidence Model #
# =================== #

@dataclass(frozen=True)
class LowConfidenceItem:
    ioc_type: str
    value: str
    reason: str


# =========== #
# CLI styling #
# =========== #

def c_info(s: str) -> str:
    return f"{Style.BRIGHT}{Fore.CYAN}{s}{Style.RESET_ALL}"

def c_ok(s: str) -> str:
    return f"{Style.BRIGHT}{Fore.GREEN}{s}{Style.RESET_ALL}"

def _cli_line(msg: str) -> None:
    print(c_info("[#]"), msg)

def _now_date_time_local() -> Tuple[str, str]:
    return time.strftime("%Y-%m-%d"), time.strftime("%H:%M:%S")


# ======= #
# Helpers #
# ======= #

def _ensure_dir(p: str) -> None:
    os.makedirs(p, exist_ok=True)

def _load_json_file(path: str) -> Optional[dict]:
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None


# ===================================== #
# Fetch URL (stdlib) with decompression #
# ===================================== #

def _decompress(data: bytes, content_encoding: str) -> bytes:
    enc = (content_encoding or "").lower()
    if "gzip" in enc:
        return gzip.decompress(data)
    if "deflate" in enc:
        try:
            return zlib.decompress(data)
        except zlib.error:
            return zlib.decompress(data, -zlib.MAX_WBITS)
    return data

def fetch_url_bytes(url: str, timeout: int, max_bytes: int, redirect_limit: int) -> Tuple[bytes, str, Dict[str, str]]:
    current_url = url
    accept_encoding = "gzip, deflate, identity"

    for _ in range(redirect_limit + 1):
        req = urllib.request.Request(
            current_url,
            headers={
                "User-Agent": f"{APP_NAME}/{APP_VERSION} (+{PROJECT_URL})",
                "Accept-Encoding": accept_encoding,
            },
            method="GET",
        )

        try:
            with urllib.request.urlopen(req, timeout=timeout) as resp:
                status = getattr(resp, "status", None)
                final_url = resp.geturl()
                headers = {k.lower(): v for k, v in resp.headers.items()}

                if status in (301, 302, 303, 307, 308):
                    loc = resp.headers.get("Location")
                    if not loc:
                        break
                    current_url = urllib.request.urljoin(final_url, loc)
                    continue

                data = resp.read(max_bytes + 1)
                if len(data) > max_bytes:
                    raise ValueError(f"Response exceeded max size ({max_bytes} bytes).")

                ce = headers.get("content-encoding", "")
                if "br" in (ce or "").lower() and accept_encoding != "identity":
                    accept_encoding = "identity"
                    continue

                data = _decompress(data, ce)
                return data, final_url, headers

        except urllib.error.HTTPError as e:
            raise RuntimeError(f"HTTP error fetching URL: {e.code} {e.reason}") from e
        except urllib.error.URLError as e:
            raise RuntimeError(f"Network error fetching URL: {e.reason}") from e

    raise RuntimeError(f"Too many redirects (limit={redirect_limit}).")

def decode_bytes(data: bytes, headers: Dict[str, str]) -> str:
    ct = headers.get("content-type", "")
    m = re.search(r"charset=([^\s;]+)", ct, re.IGNORECASE)
    if m:
        cs = m.group(1).strip()
        try:
            return data.decode(cs, errors="replace")
        except LookupError:
            pass
    return data.decode("utf-8", errors="replace")


# ============ #
# File reading #
# ============ #

def read_file_text(path: str) -> str:
    if not os.path.isfile(path):
        raise FileNotFoundError(f"File not found: {path}")

    ext = os.path.splitext(path)[1].lower().strip(".")

    if ext in ("txt", "html", "htm"):
        with open(path, "rb") as f:
            return f.read().decode("utf-8", errors="replace")

    if ext == "pdf":
        try:
            from pdfminer.high_level import extract_text as pdf_extract_text
        except Exception as e:
            raise RuntimeError("pdfminer.six is required for PDF parsing. Install dependencies.") from e
        return pdf_extract_text(path) or ""

    if ext == "docx":
        try:
            import docx  # python-docx
        except Exception as e:
            raise RuntimeError("python-docx is required for DOCX parsing. Install dependencies.") from e

        d = docx.Document(path)
        parts = []
        for p in d.paragraphs:
            if p.text:
                parts.append(p.text)
        for table in d.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    t = (cell.text or "").strip()
                    if t:
                        row_text.append(t)
                if row_text:
                    parts.append(" | ".join(row_text))
        return "\n".join(parts)

    if ext == "xlsx":
        try:
            import openpyxl
        except Exception as e:
            raise RuntimeError("openpyxl is required for XLSX parsing. Install dependencies.") from e

        wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
        parts = []
        for sheet in wb.worksheets:
            parts.append(f"[SHEET] {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = []
                for v in row:
                    if v is None:
                        continue
                    s = str(v).strip()
                    if s:
                        cells.append(s)
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    raise RuntimeError(f"Unsupported file extension: .{ext} (supported: txt, html, pdf, docx, xlsx)")


# ========================================= #
# Trafi extraction (no links / no metadata) #
# ========================================= #

def extract_article_text_from_html(html: str) -> str:
    text = trafilatura.extract(
        html,
        output_format="txt",
        include_links=False,
        include_images=False,
        include_tables=True,
        favor_precision=True,
        deduplicate=True,
    )
    return (text or "").strip()

def looks_like_html(s: str) -> bool:
    sl = s.lower()
    return "<html" in sl or "</" in sl


# ============================= #
# Deobfuscation & Normalization #
# ============================= #

def deobfuscate_text(text: str) -> str:
    t = text
    t = re.sub(r"\bhxxps\b", "https", t, flags=re.IGNORECASE)
    t = re.sub(r"\bhxxp\b", "http", t, flags=re.IGNORECASE)
    t = t.replace("[.]", ".").replace("(.)", ".").replace("{.}", ".")
    t = t.replace("[:]", ":").replace("[://]", "://")
    t = re.sub(r"\s+\(dot\)\s+", ".", t, flags=re.IGNORECASE)
    t = re.sub(r"\s+dot\s+", ".", t, flags=re.IGNORECASE)
    return t

def strip_trailing_punct(s: str) -> str:
    return s.strip().strip(".,;:!?)]}\"'\u201d\u2019")

def normalize_domain(d: str) -> str:
    return strip_trailing_punct(d).lower().strip(".")

def normalize_url(u: str) -> str:
    return strip_trailing_punct(u)

def normalize_hash(h: str) -> str:
    return strip_trailing_punct(h).lower()


# =========== #
# IOC regexes #
# =========== #

URL_RE = re.compile(r"\bhttps?://[^\s<>\"]+", re.IGNORECASE)

DOMAIN_RE = re.compile(
    r"\b(?:[a-z0-9](?:[a-z0-9-]{0,61}[a-z0-9])?\.)+(?:[a-z]{2,63})\b",
    re.IGNORECASE,
)

EMAIL_RE = re.compile(r"\b[a-z0-9._%+-]+@[a-z0-9.-]+\.[a-z]{2,63}\b", re.IGNORECASE)

IPV4_RE = re.compile(r"\b(?:\d{1,3}\.){3}\d{1,3}\b")
IPV6_RE = re.compile(r"\b(?:[0-9a-f]{0,4}:){2,7}[0-9a-f]{0,4}\b", re.IGNORECASE)

MD5_RE = re.compile(r"\b[a-f0-9]{32}\b", re.IGNORECASE)
SHA1_RE = re.compile(r"\b[a-f0-9]{40}\b", re.IGNORECASE)
SHA256_RE = re.compile(r"\b[a-f0-9]{64}\b", re.IGNORECASE)

CVE_RE = re.compile(r"\bCVE-\d{4}-\d{4,}\b", re.IGNORECASE)

def extract_iocs(text: str) -> Dict[str, Set[str]]:
    out: Dict[str, Set[str]] = {k: set() for k in IOC_TYPES_ORDER}

    for m in URL_RE.finditer(text):
        out["url"].add(normalize_url(m.group(0)))

    for m in EMAIL_RE.finditer(text):
        out["email"].add(strip_trailing_punct(m.group(0)).lower())

    for m in SHA256_RE.finditer(text):
        out["sha256"].add(normalize_hash(m.group(0)))
    for m in SHA1_RE.finditer(text):
        out["sha1"].add(normalize_hash(m.group(0)))
    for m in MD5_RE.finditer(text):
        out["md5"].add(normalize_hash(m.group(0)))

    for m in CVE_RE.finditer(text):
        out["cve"].add(m.group(0).upper())

    for m in IPV4_RE.finditer(text):
        out["ipv4"].add(strip_trailing_punct(m.group(0)))
    for m in IPV6_RE.finditer(text):
        out["ipv6"].add(strip_trailing_punct(m.group(0)))

    for m in DOMAIN_RE.finditer(text):
        out["domain"].add(normalize_domain(m.group(0)))

    return out


# =========================== #
# URL host helpers (IP-aware) #
# =========================== #

def url_host(url: str) -> Optional[str]:
    m = re.match(r"^https?://([^/]+)", url, flags=re.IGNORECASE)
    if not m:
        return None
    host = m.group(1)
    host = host.split("@")[-1]
    host = host.split(":")[0]
    return host.lower().strip(".")

def is_ip_literal(s: str) -> bool:
    try:
        ipaddress.ip_address(s)
        return True
    except ValueError:
        return False


# ============== #
# PSL Validation #
# ============== #

def psl_file_path() -> Optional[str]:
    if os.path.isfile(CACHE_PSL_FILE):
        return CACHE_PSL_FILE
    if os.path.isfile(BUNDLED_PSL_FILE):
        return BUNDLED_PSL_FILE
    return None

def psl_invalid_reason(domain: str) -> Optional[str]:
    d = normalize_domain(domain)

    if is_ip_literal(d):
        return None

    psl_path = psl_file_path()
    if not psl_path:
        return None

    if not d or "." not in d:
        return "invalid public suffix (PSL)"

    try:
        sld = get_sld(d, strict=True, psl_file=psl_path)
    except Exception:
        sld = None

    if not sld:
        return "invalid public suffix (PSL)"

    return None


# =================== #
# Warninglists Loader #
# =================== #

class WarninglistIndex:
    """
    Minimal warninglist matcher:
    - exact string matches (domains/hosts/urls)
    - CIDR matches for IPs
    """

    def __init__(self) -> None:
        self._exact_map: Dict[str, str] = {}      # value -> list name
        self._cidr: List[Tuple[ipaddress._BaseNetwork, str]] = []  # (network, list name)

    @staticmethod
    def _iter_warninglist_files(base_dir: str) -> Iterable[str]:
        # MISP structure: warninglists/lists/*.json
        lists_dir = os.path.join(base_dir, "lists")
        if not os.path.isdir(lists_dir):
            return []
        out = []
        for root, _, files in os.walk(lists_dir):
            for fn in files:
                if fn.lower().endswith(".json"):
                    out.append(os.path.join(root, fn))
        return out

    @staticmethod
    def _pick_warninglists_base_dir() -> Optional[str]:
        # Cache, bundled snapshot
        cache = CACHE_WARNINGLISTS_DIR
        bundled = BUNDLED_WARNINGLISTS_DIR
        if os.path.isdir(os.path.join(cache, "lists")):
            return cache
        if os.path.isdir(os.path.join(bundled, "lists")):
            return bundled
        return None

    def load(self) -> None:
        base = self._pick_warninglists_base_dir()
        if not base:
            return

        for path in self._iter_warninglist_files(base):
            doc = _load_json_file(path)
            if not isinstance(doc, dict):
                continue

            name = str(doc.get("name") or doc.get("description") or os.path.basename(path))
            items = doc.get("list")

            if not isinstance(items, list) or not items:
                continue

            for raw in items:
                if not isinstance(raw, str):
                    continue
                v = raw.strip()
                if not v:
                    continue

                # Normalize common forms
                v_norm = v.lower().strip()

                # CIDR or IP
                try:
                    if "/" in v_norm:
                        net = ipaddress.ip_network(v_norm, strict=False)
                        self._cidr.append((net, name))
                        continue
                    # single IP
                    ipaddress.ip_address(v_norm)
                    self._exact_map.setdefault(v_norm, name)
                    continue
                except ValueError:
                    pass

                # domain/host/url-ish
                # Store normalized domain too
                self._exact_map.setdefault(v_norm, name)

        # Sort CIDR for stable behavior
        self._cidr.sort(key=lambda x: (x[0].version, x[0].prefixlen, str(x[0])))

    def match_domain(self, domain: str) -> Optional[str]:
        d = normalize_domain(domain)
        
        # 1. Check exact match
        if d in self._exact_map:
            return self._exact_map[d]

        # 2. Check parent domains (suffix match)
        parts = d.split(".")
        for i in range(1, len(parts)):
            parent = ".".join(parts[i:])
            if parent in self._exact_map:
                return self._exact_map[parent]
                
        return None

    def match_url(self, url: str) -> Optional[str]:
        u = normalize_url(url).strip()
        u_low = u.lower()
        if u_low in self._exact_map:
            return self._exact_map[u_low]

        h = url_host(u)
        if not h:
            return None
        if is_ip_literal(h):
            return self.match_ip(h)
        return self.match_domain(h)

    def match_ip(self, ip_str: str) -> Optional[str]:
        s = ip_str.strip().lower()
        if s in self._exact_map:
            return self._exact_map[s]
        try:
            ip_obj = ipaddress.ip_address(s)
        except ValueError:
            return None
        for net, name in self._cidr:
            if ip_obj.version != net.version:
                continue
            if ip_obj in net:
                return name
        return None


def load_warninglists() -> WarninglistIndex:
    wl = WarninglistIndex()
    wl.load()
    return wl


# ================================ #
# IP low-confidence classification #
# ================================ #

def ip_low_conf_reason(ip_str: str) -> Optional[str]:
    try:
        ip_obj = ipaddress.ip_address(ip_str)
    except ValueError:
        return "invalid ip"

    if ip_obj.is_private:
        return "private ip"
    if ip_obj.is_loopback:
        return "loopback ip"
    if ip_obj.is_link_local:
        return "link-local ip"
    if ip_obj.is_multicast:
        return "multicast ip"
    if ip_obj.is_reserved:
        return "reserved ip"
    if ip_obj.is_unspecified:
        return "unspecified ip"

    if isinstance(ip_obj, ipaddress.IPv6Address) and getattr(ip_obj, "is_site_local", False):
        return "site-local ip (ipv6)"

    return None


# ========================== #
# Low-Confidence Computation #
# ========================== #

def compute_low_confidence(iocs: Dict[str, Set[str]], wl: WarninglistIndex) -> List[LowConfidenceItem]:
    low: List[LowConfidenceItem] = []

    # Domains
    for d in sorted(iocs["domain"]):
        parts = d.rsplit(".", 1)
        if len(parts) == 2 and parts[1].lower() in FILENAME_LIKE_EXTS:
            low.append(LowConfidenceItem("domain", d, "looks like filename (extension)"))

        hit = wl.match_domain(d)
        if hit:
            low.append(LowConfidenceItem("domain", d, f"Warninglist: {hit}"))

        pr = psl_invalid_reason(d)
        if pr:
            low.append(LowConfidenceItem("domain", d, pr))

    # Emails
    for e in sorted(iocs["email"]):
        if "@" in e:
            domain_part = e.split("@", 1)[1]
            hit = wl.match_domain(domain_part)
            # Check Warninglists
            if hit:
                low.append(LowConfidenceItem("email", e, f"Warninglist: {hit}"))
                continue

            # Check PSL
            pr = psl_invalid_reason(domain_part)
            if pr:
                low.append(LowConfidenceItem("email", e, f"Invalid domain: {pr}"))

    # URLs
    for u in sorted(iocs["url"]):
        # Do NOT PSL-check URLs whose host is an IP
        host = url_host(u)
        if host and (not is_ip_literal(host)):
            pr = psl_invalid_reason(host)
            if pr:
                low.append(LowConfidenceItem("url", u, pr))

        # static asset URLs -> low
        path_part = u.split("?", 1)[0].lower()
        pext = path_part.rsplit(".", 1)[-1] if "." in path_part else ""
        if pext in STATIC_ASSET_EXTS:
            low.append(LowConfidenceItem("url", u, "static asset url (likely noise)"))

        hit = wl.match_url(u)
        if hit:
            low.append(LowConfidenceItem("url", u, f"Warninglist: {hit}"))

    # IPs
    for t in ("ipv4", "ipv6"):
        for ip_str in sorted(iocs[t]):
            r = ip_low_conf_reason(ip_str)
            if r:
                low.append(LowConfidenceItem(t, ip_str, r))
                continue
            hit = wl.match_ip(ip_str)
            if hit:
                low.append(LowConfidenceItem(t, ip_str, f"Warninglist: {hit}"))

    # De-duplicate per (type,value) keep first reason
    seen = set()
    deduped: List[LowConfidenceItem] = []
    for item in low:
        key = (item.ioc_type, item.value)
        if key in seen:
            continue
        seen.add(key)
        deduped.append(item)

    return deduped


# ======= #
# Writers #
# ======= #

def write_txt(out_path: str, source: str, iocs: Dict[str, Set[str]], low_conf: List[LowConfidenceItem], user_agent: str) -> None:
    date_s, time_s = _now_date_time_local()
    title = f"{APP_NAME} Run Log"
    underline = "=" * len(title)

    with open(out_path, "w", encoding="utf-8") as f:
        f.write(f"{title}\n{underline}\n\n")
        f.write(f"[#] {'Target:':<12} {source}\n")
        f.write(f"[#] {'Date:':<12} {date_s}\n")
        f.write(f"[#] {'Time:':<12} {time_s}\n")
        f.write(f"[#] {'User-Agent:':<12} {user_agent}\n")
        f.write(f"[#] {'Output File:':<12} {out_path}\n\n")

        f.write("-------\n")
        f.write("Results\n")
        f.write("-------\n\n")

        for t in IOC_TYPES_ORDER:
            vals = sorted(iocs.get(t, set()))
            if not vals:
                continue
            f.write(f"[#] {t.upper()} ({len(vals)})\n")
            for v in vals:
                f.write(f"{v}\n")
            f.write("\n")

        f.write("-----------------------\n")
        f.write("Low-Confidence (Review)\n")
        f.write("-----------------------\n\n")

        if not low_conf:
            f.write("(none)\n")
            return

        grouped: Dict[str, List[LowConfidenceItem]] = {}
        for item in low_conf:
            grouped.setdefault(item.ioc_type, []).append(item)

        for t in sorted(grouped.keys()):
            f.write(f"[#] {t.upper()} ({len(grouped[t])})\n")
            for item in grouped[t]:
                f.write(f"{item.value} >> {item.reason}\n")
            f.write("\n")


def write_json(out_path: str, source: str, iocs: Dict[str, Set[str]], low_conf: List[LowConfidenceItem]) -> None:
    payload = {
        "source": source,
        "generated_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
        "counts": {t: len(iocs.get(t, set())) for t in IOC_TYPES_ORDER},
        "iocs": {t: sorted(iocs.get(t, set())) for t in IOC_TYPES_ORDER},
        "low_confidence": [{"type": x.ioc_type, "value": x.value, "reason": x.reason} for x in low_conf],
        "notice": "This tool may produce false positives. Review output before ingestion.",
    }

    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(payload, f, indent=2, ensure_ascii=False)


# ================================================== #
# Updates (--update) for PSL + warninglists snapshot #
# ================================================== #

def update_psl_to_cache(timeout: int = DEFAULT_TIMEOUT_SEC) -> None:
    _ensure_dir(CACHE_DIR)
    req = urllib.request.Request(
        PSL_URL,
        headers={"User-Agent": f"{APP_NAME}/{APP_VERSION}", "Accept-Encoding": "identity"},
    )
    with urllib.request.urlopen(req, timeout=timeout) as resp:
        data = resp.read()
    if len(data) < 1000:
        raise RuntimeError("PSL update failed: unexpected content downloaded.")
    with open(CACHE_PSL_FILE, "wb") as f:
        f.write(data)

def update_warninglists_to_cache(timeout: int = DEFAULT_TIMEOUT_SEC) -> None:
    _ensure_dir(CACHE_WARNINGLISTS_DIR)
    req = urllib.request.Request(
        MISP_WARNINGLISTS_ZIP,
        headers={"User-Agent": f"{APP_NAME}/{APP_VERSION}", "Accept-Encoding": "identity"},
    )
    with urllib.request.urlopen(req, timeout=timeout) as resp:
        data = resp.read()

    extracted = 0
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        for m in z.namelist():
            if not m.endswith(".json"):
                continue
            if "/lists/" not in m:
                continue
            rel = m.split("/lists/", 1)[1]
            out_path = os.path.join(CACHE_WARNINGLISTS_DIR, "lists", rel)
            _ensure_dir(os.path.dirname(out_path))
            with z.open(m) as src, open(out_path, "wb") as dst:
                dst.write(src.read())
            extracted += 1

    if extracted == 0:
        raise RuntimeError("Warninglists update failed: no lists/**/*.json extracted.")

def write_cache_meta() -> None:
    _ensure_dir(CACHE_DIR)
    meta = _load_json_file(CACHE_META_FILE) or {}
    meta["updated_at"] = time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())
    meta["warninglists_zip"] = MISP_WARNINGLISTS_ZIP
    meta["psl_url"] = PSL_URL
    with open(CACHE_META_FILE, "w", encoding="utf-8") as f:
        json.dump(meta, f, indent=2, ensure_ascii=False)


# === #
# CLI #
# === #

def parse_args(argv: Optional[List[str]] = None) -> argparse.Namespace:
    p = argparse.ArgumentParser(
        prog=APP_NAME,
        description=f"{APP_NAME} v{APP_VERSION}: {APP_DESC}",
    )

    p.add_argument("--version", action="version", version=f"{APP_NAME} {APP_VERSION}")
    g = p.add_mutually_exclusive_group(required=False)
    g.add_argument("--url", help="URL to fetch and extract IOCs from")
    g.add_argument("--file", help="Local file path to extract IOCs from (txt, html, pdf, docx, xlsx)")

    p.add_argument("--out", help="Output file path (mandatory unless using --update alone)")
    p.add_argument("--format", default="txt", choices=["txt", "json"], help="Output format (default: txt)")
    p.add_argument("--update", action="store_true", help=f"Update datasets (MISP warninglists + PSL) in cache ({CACHE_DIR})")

    args = p.parse_args(argv)

    if not args.update:
        if not (args.url or args.file):
            p.error("one of --url or --file is required")
        if not args.out:
            p.error("--out is required")
    else:
        if (args.url or args.file) and not args.out:
            p.error("--out is required when using --update with --url/--file")

    return args


def main(argv: Optional[List[str]] = None) -> None:
    colorama_init(autoreset=True)
    args = parse_args(argv)

    if args.update:
        try:
            _cli_line(f"Updating datasets into: {CACHE_DIR}")
            update_warninglists_to_cache()
            update_psl_to_cache()
            write_cache_meta()
            _cli_line("Update complete.\n")
        except Exception as e:
            print(f"[ERROR] Update failed: {e}", file=sys.stderr)
            sys.exit(1)

        if not (args.url or args.file):
            return

    source = args.url if args.url else args.file
    user_agent = f"{APP_NAME}/{APP_VERSION} (+{PROJECT_URL})"

    _cli_line(f"{APP_NAME} v{APP_VERSION}")
    _cli_line(PROJECT_URL)
    _cli_line(APP_DESC)
    _cli_line(f"Target: {source}\n")

    try:
        if args.url:
            raw_bytes, final_url, headers = fetch_url_bytes(
                args.url,
                timeout=DEFAULT_TIMEOUT_SEC,
                max_bytes=DEFAULT_MAX_BYTES,
                redirect_limit=DEFAULT_REDIRECT_LIMIT,
            )
            # Determine extension from URL or Content-Type
            ctype = headers.get("content-type", "").lower()
            url_ext = os.path.splitext(final_url)[1].lower().strip(".")
            # Force extension if headers are explicit, otherwise trust URL
            use_ext = url_ext
            if "application/pdf" in ctype and use_ext != "pdf":
                use_ext = "pdf"
            elif "wordprocessingml" in ctype and use_ext != "docx":
                use_ext = "docx"
            elif "spreadsheetml" in ctype and use_ext != "xlsx":
                use_ext = "xlsx"

            # Delegate binary formats to a temp file
            if use_ext in ("pdf", "docx", "xlsx"):
                _cli_line(f"Detected binary format ({use_ext}). Downloading to temp file...")
                with tempfile.NamedTemporaryFile(delete=False, suffix=f".{use_ext}") as tmp:
                    tmp.write(raw_bytes)
                    tmp_path = tmp.name

                try:
                    raw = read_file_text(tmp_path)
                finally:
                    if os.path.exists(tmp_path):
                        os.remove(tmp_path)
            else:
                # Default text/html handling
                raw = decode_bytes(raw_bytes, headers)

            source = final_url
        else:
            raw = read_file_text(args.file)

        processed = raw
        if looks_like_html(raw):
            processed = extract_article_text_from_html(raw)

        text = deobfuscate_text(processed)
        iocs = extract_iocs(text)

        wl = load_warninglists()
        low_conf = compute_low_confidence(iocs, wl)

        # Move low-confidence items out of main results (so they appear only once)
        low_unique = len({(x.ioc_type, x.value) for x in low_conf})
        for item in low_conf:
            if item.ioc_type in iocs:
                iocs[item.ioc_type].discard(item.value)

        final_total = sum(len(iocs.get(t, set())) for t in IOC_TYPES_ORDER)
        total_unique = final_total + low_unique

        out_path = args.out
        if args.format == "txt":
            write_txt(out_path, source, iocs, low_conf, user_agent=user_agent)
        else:
            write_json(out_path, source, iocs, low_conf)

        print(f"\n{c_info('Total Unique IOCs:')} {Style.BRIGHT}{total_unique}{Style.RESET_ALL}\n")
        print(c_ok("IOC results and log file have been saved into:"))
        print(out_path)

    except Exception as e:
        print(f"[ERROR] {e}", file=sys.stderr)
        sys.exit(1)
