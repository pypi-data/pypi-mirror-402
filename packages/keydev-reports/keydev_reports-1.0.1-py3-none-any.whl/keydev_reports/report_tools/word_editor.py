import io
import re
from typing import Optional, List

from docx import Document
from docx.text.paragraph import Paragraph as P
from python_docx_replace.paragraph import Paragraph

from .base_editor import BaseEditor


class WordEditor(BaseEditor):
    """
    Класс для редактирования Word отчетов.
    """

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        """
        Конструктор класса. Принимает имя файла и загружает его.
        :param file_name: Имя файла в формате DOCX.
        """
        self.book = Document(self.file_name) if self.file_name else Document()

    def save(self, new_file_name: Optional[str] = None) -> None:
        """
        Override save to support BytesIO mode.
        If new_file_name is provided, saves to file.
        If no file_name is set, saves to BytesIO (self.output).
        """
        if new_file_name:
            self.file_name = new_file_name
            self.book.save(self.file_name)
        elif self.file_name:
            self.book.save(self.file_name)
        else:
            # BytesIO mode - save to self.output
            self.output.seek(0)
            self.output.truncate()
            self.book.save(self.output)
            self.output.seek(0)

    @staticmethod
    def replace_text_in_paragraph(paragraph: P, key: str, value: str) -> None:
        """
        Статический метод для замены текста по ключу и значению.
        :param paragraph: Текст, в котором производится замена.
        :param key: Ключ, который нужно заменить.
        :param value: Значение для замены.
        """
        if key in paragraph.text:
            inline = paragraph.runs
            for item in inline:
                if key in item.text:
                    item.text = item.text.replace(key, value)

    def docx_replace(self, **kwargs: str) -> None:
        """
        Метод для замены ключей в документе DOCX на соответствующие значения.
        :param kwargs: Словарь, где ключи - ключи для замены, значения - новые значения.
        """
        pattern = r'\$\{[^}]*\}'
        for p in Paragraph.get_all(self.book):
            paragraph = Paragraph(p)
            for key, value in kwargs.items():
                key = f'${{{key}}}'
                paragraph.replace_key(key, str(value))
            for match in re.finditer(pattern, paragraph.p.text):
                key = match.group()
                paragraph.replace_key(key, '')

    def add_table(self, data: List[List], table_style: Optional[str] = None) -> None:
        """
        Метод для создания и заполнения таблицы в документе.
        :param data: Данные для таблицы (Лист из листов).
        :param table_style: Стиль для таблицы.
        :return: NoReturn
        """
        table = self.book.add_table(rows=1, cols=len(data[0]))

        # Устанавливаем header
        header_cells = table.rows[0].cells
        for i, field in enumerate(data[0]):
            header_cells[i].text = field

        # Устанавливаем значения
        for row_data in data:
            row = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row[i].text = str(cell_data)

        if table_style:
            table.style = table_style

    def get_bytes_io(self) -> bytes:
        """
        Returns document content as bytes (for S3 upload).
        Must call save() first without arguments to save to BytesIO.
        :return: Bytes content of the document.
        """
        self.output.seek(0)
        return self.output.read()

    def load_from_bytes(self, content: bytes) -> None:
        """
        Загружает Word документ из байтов.

        :param content: Содержимое файла в виде байтов
        """
        self.book = Document(io.BytesIO(content))
