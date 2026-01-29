import re
import json
import docx  # pip install python-docx
from PyPDF2 import PdfReader
from langdetect import detect
from googletrans import Translator as GoogleTranslator
import os
import docx
from fpdf import FPDF
import random
import time
from translate import Translator as TranslateTranslator
import numpy as np
from nltk.tokenize import sent_tokenize
from itertools import pairwise
from tqdm import tqdm
from fuzzywuzzy import fuzz, process


def split_by_sent_n(text, n=10):
    # split text into sentences
    text_split_by_sent = sent_tokenize(text)
    cut_loc_array = np.arange(0, len(text_split_by_sent), n)
    if cut_loc_array[-1] != len(text_split_by_sent):
        cut_loc = np.append(cut_loc_array, len(text_split_by_sent))
    else:
        cut_loc = cut_loc_array
    # get text in section (e.g., every 10 sentences)
    text_section = []
    for i, j in pairwise(cut_loc):
        text_section.append(text_split_by_sent[i:j])
    return text_section


def account_letters(text, n=10):
    len_ = []
    [len_.append(len(i)) for i in split_by_sent_n(text, n)[0]]
    return np.sum(len_)


def auto_chunk_size(txt, verbose=False):
    chunk_size = []
    for i in range(1, 50):
        while 4000 < account_letters(txt, n=i) < 4700:
            if verbose:
                print(f"the optimal chunk_size is {i} sentences")
            chunk_size.append(i)
            break
    return chunk_size[0]


# import pathlib
# import argostranslate.package
# import argostranslate.translate
def get_lang_code_iso639():
    from JFL import netfinder

    url = "https://en.wikipedia.org/wiki/List_of_ISO_639_language_codes"
    # res=netfinder.fetch(url,where="table",what="wikitable sortable jquery-tablesorter")
    res = netfinder.fetch(url, where="tr", extend=0)
    fullname, shortcut = [], []
    for i in range(6, len(res) - 2):
        if len(res[i]) > len(res[i + 1]) and res[i + 1][:2] == res[i + 2][:2]:
            fullname.append(res[i])
            shortcut.append(res[i + 1])
    lang_code_iso639 = dict([*zip(fullname, shortcut)])
    return lang_code_iso639


def detect_lang(text, output="lang", verbose=False):
    dir_curr_script = os.path.dirname(os.path.abspath(__file__))
    dir_lang_code = dir_curr_script + "/data/lang_code_iso639.json"
    with open(dir_lang_code, "r") as file:
        lang_code_iso639 = json.load(file)
    l_lang, l_code = [], []
    [[l_lang.append(v), l_code.append(k)] for v, k in lang_code_iso639.items()]
    try:
        if is_text(text):
            code_detect = detect(text)
            if "c" in output.lower():  # return code
                return l_code[strcmp(code_detect, l_code, verbose=verbose)[1]]
            else:
                return l_lang[strcmp(code_detect, l_code, verbose=verbose)[1]]
        else:
            print(f"{text} is not supported")
            return "no"
    except:
        return "no"


def is_text(s):
    has_alpha = any(char.isalpha() for char in s)
    has_non_alpha = any(not char.isalpha() for char in s)
    # no_special = not re.search(r'[^A-Za-z0-9\s]', s)
    return has_alpha and has_non_alpha


def strcmp(search_term, candidates, ignore_case=True, verbose=False, scorer="WR"):
    """
    Compares a search term with a list of candidate strings and finds the best match based on similarity score.

    Parameters:
    search_term (str): The term to be searched for.
    candidates (list of str): A list of candidate strings to compare against the search term.
    ignore_case (bool): If True, the comparison ignores case differences.
    verbose (bool): If True, prints the similarity score and the best match.

    Returns:
    tuple: A tuple containing the best match and its index in the candidates list.
    """

    def to_lower(s, ignore_case=True):
        # Converts a string or list of strings to lowercase if ignore_case is True.
        if ignore_case:
            if isinstance(s, str):
                return s.lower()
            elif isinstance(s, list):
                return [elem.lower() for elem in s]
        return s

    str1_, str2_ = to_lower(search_term, ignore_case), to_lower(candidates, ignore_case)
    if isinstance(str2_, list):
        if "part" in scorer.lower():
            similarity_scores = [fuzz.partial_ratio(str1_, word) for word in str2_]
        elif "W" in scorer.lower():
            similarity_scores = [fuzz.WRatio(str1_, word) for word in str2_]
        elif "Ratio" in scorer.lower():
            similarity_scores = [fuzz.Ratio(str1_, word) for word in str2_]
        else:
            similarity_scores = [fuzz.WRatio(str1_, word) for word in str2_]
        best_match_index = similarity_scores.index(max(similarity_scores))
        best_match_score = similarity_scores[best_match_index]
    else:
        best_match_index = 0
        if "part" in scorer.lower():
            best_match_score = fuzz.partial_ratio(str1_, str2_)
        elif "W" in scorer.lower():
            best_match_score = fuzz.WRatio(str1_, str2_)
        elif "Ratio" in scorer.lower():
            best_match_score = fuzz.Ratio(str1_, str2_)
        else:
            best_match_score = fuzz.WRatio(str1_, str2_)
    if verbose:
        print(f"\nbest_match is: {candidates[best_match_index],best_match_score}")
        best_match = process.extract(search_term, candidates)
        print(f"建议: {best_match}")
    return candidates[best_match_index], best_match_index


def methods(idx=0):
    methods_ = ["GoogleTrans (default)", "DeepL", "Argos"]
    # print(f"supported methods: {methods_}")
    # print(f"return the selected is: {methods_[idx]}")
    return methods_[idx]


DEFAULT_SERVICE_URLS = ("translate.google.de", "translate.google.fr")


def user_agent():
    # Example of generating a random user-agent string
    user_agents = [
        # Windows (Intel)
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4891.0 Safari/537.36",
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4893.0 Safari/537.36",
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4895.0 Safari/537.36",
        # Windows (ARM)
        "Mozilla/5.0 (Windows NT 10.0; Win64; arm64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4891.0 Safari/537.36",
        "Mozilla/5.0 (Windows NT 10.0; Win64; arm64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4893.0 Safari/537.36",
        "Mozilla/5.0 (Windows NT 10.0; Win64; arm64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4895.0 Safari/537.36",
        # Linux (x86_64)
        "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4891.0 Safari/537.36",
        "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4893.0 Safari/537.36",
        "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4895.0 Safari/537.36",
        # macOS (Intel)
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 12_0_1) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/15.3 Safari/605.1.15",
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 12_0_1) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/15.4 Safari/605.1.15",
        # macOS (ARM)
        "Mozilla/5.0 (Macintosh; ARM Mac OS X 12_0_1) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/15.3 Safari/605.1.15",
        "Mozilla/5.0 (Macintosh; ARM Mac OS X 12_0_1) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/15.4 Safari/605.1.15",
        # iOS Devices
        "Mozilla/5.0 (iPad; CPU OS 15_0 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/15.0 Mobile/15E148 Safari/604.1",
        "Mozilla/5.0 (iPhone; CPU iPhone OS 15_0 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/15.0 Mobile/15E148 Safari/604.1",
        # Android Devices
        "Mozilla/5.0 (Linux; Android 12; Pixel 6) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4891.0 Mobile Safari/537.36",
        "Mozilla/5.0 (Linux; Android 12; Pixel 6) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4893.0 Mobile Safari/537.36",
        "Mozilla/5.0 (Linux; Android 12; Pixel 6) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4895.0 Mobile Safari/537.36",
        # Smart TVs
        "Mozilla/5.0 (SMART-TV; LINUX; Tizen 6.0) AppleWebKit/537.36 (KHTML, like Gecko) SmartTV/1.0",
        "Mozilla/5.0 (SMART-TV; LINUX; Tizen 6.0) AppleWebKit/537.36 (KHTML, like Gecko) WebAppManager/1.0",
        # Game Consoles
        "Mozilla/5.0 (PlayStation 5 3.01) AppleWebKit/605.1.15 (KHTML, like Gecko)",
        "Mozilla/5.0 (Xbox One 10.0) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/88.0.4324.182 Safari/537.36 Edge/44.18363.8740",
    ]
    agents = random.choice(user_agents)
    return agents


def get_language_code(language, translator="google"):
    """
    Get language code for translation services (Google Translate, DeepL).
    """
    deepl_languages = {
        "English": "EN",
        "German": "DE",
        "French": "FR",
        "Spanish": "ES",
        "Italian": "IT",
        "Dutch": "NL",
        "Polish": "PL",
        "Russian": "RU",
        "Japanese": "JA",
        "Chinese": "ZH",
    }
    google_languages = {
        "Afrikaans": "af",
        "Albanian": "sq",
        "Amharic": "am",
        "Arabic": "ar",
        "Armenian": "hy",
        "Azerbaijani": "az",
        "Basque": "eu",
        "Belarusian": "be",
        "Bengali": "bn",
        "Bosnian": "bs",
        "Bulgarian": "bg",
        "Catalan": "ca",
        "Cebuano": "ceb",
        "Chichewa": "ny",
        "Chinese": "zh-CN",
        "Corsican": "co",
        "Croatian": "hr",
        "Czech": "cs",
        "Danish": "da",
        "Dutch": "nl",
        "English": "en",
        "Esperanto": "eo",
        "Estonian": "et",
        "Filipino": "tl",
        "Finnish": "fi",
        "French": "fr",
        "Frisian": "fy",
        "Galician": "gl",
        "Georgian": "ka",
        "German": "de",
        "Greek": "el",
        "Gujarati": "gu",
        "HaitianCreole": "ht",
        "Hausa": "ha",
        "Hawaiian": "haw",
        "Hebrew": "he",
        "Hindi": "hi",
        "Hmong": "hmn",
        "Hungarian": "hu",
        "Icelandic": "is",
        "Igbo": "ig",
        "Indonesian": "id",
        "Irish": "ga",
        "Italian": "it",
        "Japanese": "ja",
        "Javanese": "jv",
        "Kannada": "kn",
        "Kazakh": "kk",
        "Khmer": "km",
        "Kinyarwanda": "rw",
        "Korean": "ko",
        "Kurdish": "ku",
        "Kyrgyz": "ky",
        "Lao": "lo",
        "Latin": "la",
        "Latvian": "lv",
        "Lithuanian": "lt",
        "Luxembourgish": "lb",
        "Macedonian": "mk",
        "Malagasy": "mg",
        "Malay": "ms",
        "Malayalam": "ml",
        "Maltese": "mt",
        "Maori": "mi",
        "Marathi": "mr",
        "Mongolian": "mn",
        "Myanmar": "my",
        "Nepali": "ne",
        "Norwegian": "no",
        "Odia": "or",
        "Oriya": "or",
        "Pashto": "ps",
        "Persian": "fa",
        "Polish": "pl",
        "Portuguese": "pt",
        "Punjabi": "pa",
        "Romanian": "ro",
        "Russian": "ru",
        "Samoan": "sm",
        "ScotsGaelic": "gd",
        "Serbian": "sr",
        "Sesotho": "st",
        "Shona": "sn",
        "Sindhi": "sd",
        "Sinhala": "si",
        "Slovak": "sk",
        "Slovenian": "sl",
        "Somali": "so",
        "Spanish": "es",
        "Sundanese": "su",
        "Swahili": "sw",
        "Swedish": "sv",
        "Tajik": "tg",
        "Tamil": "ta",
        "Tatar": "tt",
        "Telugu": "te",
        "Thai": "th",
        "Turkish": "tr",
        "Turkmen": "tk",
        "Ukrainian": "uk",
        "Urdu": "ur",
        "Uyghur": "ug",
        "Uzbek": "uz",
        "Vietnamese": "vi",
        "Welsh": "cy",
        "Xhosa": "xh",
        "Yiddish": "yi",
        "Yoruba": "yo",
        "Zulu": "zu",
    }
    argos_languages = {
        "Afrikaans": "af",
        "Albanian": "sq",
        "Amharic": "am",
        "Arabic": "ar",
        "Armenian": "hy",
        "Azerbaijani": "az",
        "Basque": "eu",
        "Belarusian": "be",
        "Bengali": "bn",
        "Bosnian": "bs",
        "Bulgarian": "bg",
        "Catalan": "ca",
        "Cebuano": "ceb",
        "Chichewa": "ny",
        "Chinese": "zh",
        "Corsican": "co",
        "Croatian": "hr",
        "Czech": "cs",
        "Danish": "da",
        "Dutch": "nl",
        "English": "en",
        "Esperanto": "es",
        "Estonian": "et",
        "Filipino": "tl",
        "Finnish": "fi",
        "French": "fr",
        "Frisian": "fy",
        "Galician": "gl",
        "Georgian": "ka",
        "German": "de",
        "Greek": "el",
        "Gujarati": "gu",
        "HaitianCreole": "ht",
        "Hausa": "ha",
        "Hawaiian": "haw",
        "Hebrew": "he",
        "Hindi": "hi",
        "Hmong": "hmn",
        "Hungarian": "hu",
        "Icelandic": "is",
        "Igbo": "ig",
        "Indonesian": "id",
        "Irish": "ga",
        "Italian": "it",
        "Japanese": "ja",
        "Javanese": "jv",
        "Kannada": "kn",
        "Kazakh": "kk",
        "Khmer": "km",
        "Kinyarwanda": "rw",
        "Korean": "ko",
        "Kurdish": "ku",
        "Kyrgyz": "ky",
        "Lao": "lo",
        "Latin": "la",
        "Latvian": "lv",
        "Lithuanian": "lt",
        "Luxembourgish": "lb",
        "Macedonian": "mk",
        "Malagasy": "mg",
        "Malay": "ms",
        "Malayalam": "ml",
        "Maltese": "mt",
        "Maori": "mi",
        "Marathi": "mr",
        "Mongolian": "mn",
        "Myanmar": "my",
        "Nepali": "ne",
        "Norwegian": "no",
        "Odia": "or",
        "Oriya": "or",
        "Pashto": "ps",
        "Persian": "fa",
        "Polish": "pl",
        "Portuguese": "pt",
        "Punjabi": "pa",
        "Romanian": "ro",
        "Russian": "ru",
        "Samoan": "sm",
        "ScotsGaelic": "gd",
        "Serbian": "sr",
        "Sesotho": "st",
        "Shona": "sn",
        "Sindhi": "sd",
        "Sinhala": "si",
        "Slovak": "sk",
        "Slovenian": "sl",
        "Somali": "so",
        "Spanish": "es",
        "Sundanese": "su",
        "Swahili": "sw",
        "Swedish": "sv",
        "Tajik": "tg",
        "Tamil": "ta",
        "Tatar": "tt",
        "Telugu": "te",
        "Thai": "th",
        "Turkish": "tr",
        "Turkmen": "tk",
        "Ukrainian": "uk",
        "Urdu": "ur",
        "Uyghur": "ug",
        "Uzbek": "uz",
        "Vietnamese": "vi",
        "Welsh": "cy",
        "Xhosa": "xh",
        "Yiddish": "yi",
        "Yoruba": "yo",
        "Zulu": "zu",
    }
    if "deep" in translator.lower():
        langs = deepl_languages
    elif "goo" in translator.lower():
        langs = google_languages
    elif "ar" in translator.lower():
        langs = argos_languages
    lang_found = strcmp(language, list(langs.keys()))[0]
    if lang_found in list(langs.keys()):
        return langs[lang_found]
    else:
        print(f"fail to find the {language} code in translator {translator}")
        return None


# language = "chinese"
# # Example usage:
# google_lang_code = get_language_code(language, "google")
# deepl_lang_code = get_language_code(language, "deepl")

# print(f"Google Translate Language Code for '{language}': {google_lang_code}")
# print(f"DeepL Translator Language Code for '{language}': {deepl_lang_code}")


def load_docx(filename):
    """
    Load a .docx file and return its content as a list of strings.
    """
    doc = docx.Document(filename)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return text


# # usage
# filename = "example.docx"  # Change to the path of your .docx file
# text = load_docx(filename)
# print("Document loaded successfully.")
# print("Text from the document:")
# print(text)


def load_pdf(filename, page="all", verbose=False):
    from PyPDF2 import PdfReader
    import numpy as np

    """
        Parameters:
        filename: The path to the PDF file to be loaded.
        page (optional): 
            Specifies which page or pages to extract text from. By default, it's set to "all", which means text from all 
            pages will be returned. It can also be an integer to specify a single page number or a list of integers to 
            specify multiple pages.
        verbose (optional): 
            If True, prints the total number of pages processed.
        Functionality:
        It initializes an empty dictionary text_dict to store page numbers as keys and their corresponding text as values.
        It iterates through each page of the PDF file using a for loop.
        For each page, it extracts the text using PyPDF2's extract_text() method and stores it in text_dict with the page number incremented by 1 as the key.
        If the page parameter is an integer, it converts it into a list containing that single page number to ensure consistency in handling.
        If the page parameter is a NumPy array, it converts it to a list using the tolist() method to ensure compatibility with list operations.
        If verbose is True, it prints the total number of pages processed.
        If page is a list, it combines the text of the specified pages into a single string combined_text and returns it.
        If page is set to "all", it returns the entire text_dict containing text of all pages.
        If page is an integer, it returns the text of the specified page number.
        If the specified page is not found, it returns the string "Page is not found".
    """

    text_dict = {}
    with open(filename, "rb") as file:
        pdf_reader = PdfReader(file)
        num_pages = len(pdf_reader.pages)
        for page_num in range(num_pages):
            page_ = pdf_reader.pages[page_num]
            text_dict[page_num + 1] = page_.extract_text()
    if isinstance(page, int):
        page = [page]
    elif isinstance(page, np.ndarray):
        page = page.tolist()
    if verbose:
        print(f"total pages: {page_num}")
    if isinstance(page, list):
        combined_text = ""
        for page_num in page:
            combined_text += text_dict.get(page_num, "")
        return combined_text
    elif "all" in page.lower():
        combined_text = ""
        for i in text_dict.values():
            combined_text += i
        return combined_text
    else:
        return text_dict.get(int(page), "Page is not found")


def split_text(text, method="sentence", limit=4500):
    """
    Split text into segments based on sentence boundaries or a specified length limit.
    """
    if "sent" in method.lower():
        res = re.findall(r"https?://\S+|[^.!?]+(?:[.!?](?:\s|$))?", text)
        print(f"There are {len(res)} sentences.")
        return res
    elif "len" in method.lower():
        return [text[i : i + limit] for i in range(0, len(text), limit)]
    else:
        return re.split(r"\{}".format(method), text)


def filter_errors(text):
    # handle bugs:
    # bug1: ".com" cannot be translated, but '..com' works
    text = text.replace(".com", "..come")
    return text


def merge_text(input, robust=True):
    """
    Convert a list of strings, tuple of strings, or numpy array of strings into a single concatenated string.

    Args:
        input (list, tuple, numpy.ndarray): A list, tuple, or numpy array of strings to be concatenated.
        robust (bool, optional): If True, handles non-supported types by converting them to string.
                                If False, directly converts the input to string. Default is True.

    Returns:
        str: The concatenated string.
    """
    supported_types = (list, tuple, np.ndarray)
    if not isinstance(input, supported_types):
        raise TypeError(f"{Input} must be {supported_types}.")
    if robust:
        # Convert each element to string if it's not already
        text = [str(item) for item in input]
        # Concatenate the strings
        return "".join(text)
    else:
        return str(input)


def replace_text(text, dict_replace=None, robust=True):
    """
    Replace specified substrings in the input text with provided replacements.
    Args:
        text (str): The input text where replacements will be made.
        dict_replace (dict, optional): A dictionary containing substrings to be replaced as keys
            and their corresponding replacements as values. Defaults to {".com": "..come", "\n": " ", "\t": " ", "  ": " "}.
        robust (bool, optional): If True, additional default replacements for newline and tab characters will be applied.
                                Default is False.
    Returns:
        str: The text after replacements have been made.
    """
    if not all(text):
        return ''
    # Default replacements for newline and tab characters
    default_replacements = {
        "\a": "",
        "\b": "",
        "\f": "",
        "\n": "",
        "\r": "",
        "\t": "",
        "\v": "",
        "\\": "",  # Corrected here
        # "\?": "",
        "�": "",
        "\\x": "",  # Corrected here
        "\\x hhhh": "",
        "\\ ooo": "",  # Corrected here
        "\xa0": "",
        "  ": " ",
    }

    # If dict_replace is None, use the default dictionary
    if dict_replace is None:
        dict_replace = {}

    # If robust is True, update the dictionary with default replacements
    if robust:
        dict_replace.update(default_replacements)

    # Iterate over each key-value pair in the dictionary and replace substrings accordingly
    for k, v in dict_replace.items():
        text = text.replace(k, v)
    return text


# # usage:
#     a = "kjkjk        (a, b, c)"
#     replace_text(a, {"(": "", ")": "", "        ": " "}, robust=False)


def merge_strings_every_n(strings_list, n=10):
    merged_list = []
    if n > 0:
        for i in range(0, len(strings_list), n):
            merged_string = "".join(strings_list[i : i + n])
            merged_list.append(merged_string)
        return merged_list, n
    else:
        return strings_list, n


def translate(
    text,
    lang="chinese",
    lang_src=None,
    method=methods(),
    service_urls=DEFAULT_SERVICE_URLS,
    user_agent=user_agent(),
    verbose=True,
    error_verbose=True,
    limit=5000,
):
    """
    Translate text to the target language using the specified translation method (Google Translate or DeepL).
    lang_src (str): e.g., 'english', or 'chinese' when there are two languages, then lang_src must be given
    """
    # error_verbose = verbose or False

    if isinstance(text, list):
        # text=merge_text(text)
        text_list = [
            translate(
                i,
                lang=lang,
                lang_src=lang_src,
                method=method,
                service_urls=service_urls,
                user_agent=user_agent,
                verbose=verbose,
                error_verbose=error_verbose,
                limit=limit,
            )
            for i in tqdm(
                text,
                desc="is translating",
                colour="green",
                bar_format="{l_bar}{bar} {n_fmt}/{total_fmt}",
            )
        ]
        return text_list
    text = replace_text(text)
    if lang_src is None:
        lang_src = detect_lang(text)
    try:
        if len(text) > limit:
            n = auto_chunk_size(text)
            text_segments = split_by_sent_n(text, n)
            translations = ""
            for segment in tqdm(text_segments, desc="is translating", colour="green"):
                segment = replace_text(merge_text(segment))
                translated_segment = translate_segment(
                    text=segment,
                    lang=lang,
                    lang_src=lang_src,
                    method=method,
                    user_agent=user_agent,
                    service_urls=service_urls,
                    verbose=verbose,
                    error_verbose=error_verbose,
                )
                time.sleep(1)
                if translated_segment:
                    translations += translated_segment
                else:
                    print("Error: Translation of one of the segments failed.")
                    translations += ""
            return translations
        else:
            return translate_segment(
                text=text,
                lang=lang,
                lang_src=lang_src,
                method=method,
                user_agent=user_agent,
                service_urls=service_urls,
                verbose=verbose,
                error_verbose=error_verbose,
            )
    except Exception as e:
        if error_verbose:
            print("(translate)Error during translation :", e)
        return ""


def translate_segment(
    text,
    lang="chinese",
    lang_src=None,
    method=methods(),
    service_urls=DEFAULT_SERVICE_URLS,
    user_agent=user_agent(),
    verbose=False,
    error_verbose=True,
):
    """
    Translate a text segment to the target language using the specified translation method (Google Translate or DeepL).
    """

    text_clean = filter_errors(text)
    text_clean = replace_text(text_clean)
    if lang_src is None:
        lang_src = detect_lang(text_clean)
    try:
        lang_src = get_language_code(lang_src, "google")
        lang_tgt = get_language_code(lang, "google")
        if "goog" in method.lower():
            Trstor = GoogleTranslator(service_urls=service_urls, user_agent=user_agent)
            txt = Trstor.translate(text_clean, src=lang_src, dest=lang_tgt).text
        elif "trans" in method.lower():
            lang_src = get_language_code(lang_src, "google")
            lang_tgt = get_language_code(lang, "google")
            translator = TranslateTranslator(
                from_lang=lang_src,
                to_lang=lang_tgt,
                provider="LibreTranslate",
                secret_access_key=None,
                base_url="https://translate.astian.org/",
            )
            txt = translator.translate(text_clean)
        elif "ar" in method.lower():
            lang_src = get_language_code(language=lang_src, translator="argos")
            lang_tgt = get_language_code(language=lang, translator="argos")
            argostranslate.package.update_package_index()
            available_packages = argostranslate.package.get_available_packages()
            package_to_install = next(
                filter(
                    lambda x: x.from_code == lang_src and x.to_code == lang_tgt,
                    available_packages,
                )
            )
            argostranslate.package.install_from_path(package_to_install.download())
            # Translate
            txt = argostranslate.translate.translate("Hello World", lang_src, lang_tgt)
        else:
            print("Error: Invalid translation method. supported: 'google' or 'deepl'.")
            return ""
        if verbose:
            print(txt)
        return txt
    except Exception as e:
        txt = translate_with_retry(
            text_clean,
            lang=lang,
            lang_src=lang_src,
            method=method,
            verbose=verbose,
            error_verbose=error_verbose,
            user_agent=user_agent,
            service_urls=service_urls,
        )
        return txt


def translate_with_retry(
    text,
    lang="chinese",
    lang_src=None,
    method=methods(),
    verbose=False,
    error_verbose=True,
    user_agent=user_agent(),
    service_urls=DEFAULT_SERVICE_URLS,
):
    """
    Translate a text to the target language, retrying with alternative service URLs on connection errors.
    """

    def try_translate(text, lang, lang_src, user_agent, service_url):
        try:
            translator_ = GoogleTranslator(
                user_agent=user_agent, service_urls=[service_url]
            )
            result = translator_.translate(text, dest=lang, src=lang_src)
            if result and hasattr(result, "text"):
                return result.text
            else:
                raise ValueError(f"Invalid response from {service_url}: {result}")
        except Exception as e:
            raise RuntimeError(f"Error using {service_url}: {e}")

    if lang_src is None:
        lang_src = detect_lang(text)
        lang_src = get_language_code(language=lang_src)
    lang = get_language_code(language=lang)
    try:
        return try_translate(
            text,
            lang=lang,
            lang_src=lang_src,
            user_agent=user_agent,
            service_url=service_urls[0],
        )
    except Exception as e:
        if error_verbose:
            print("Connection error:", e)
        try:
            time.sleep(1)
            return try_translate(
                text,
                lang=lang,
                lang_src=lang_src,
                user_agent=user_agent,
                service_url=service_urls[1],
            )
        except Exception as e:
            if error_verbose:
                print(
                    f"(translate_with_retry):Connection error with {service_urls}: {e}"
                )
        if error_verbose:
            print("All service URLs failed. Unable to translate the text.")
        return text


def trans_docx(
    filename,
    lang="english",
    lang_src=None,
    method=methods(),
    service_urls=["translate.google.de"],
    verbose=False,
):
    """
    load the docx file and translated it into target lang "lang",
    verbose: (default 'False', no display)to display the translated text in for loop
    Return (list):
        the translated text as a list
    """
    txt = load_docx(filename)
    trans_text = []
    for i in txt:
        # print(i)
        j = ""
        if len(i.strip()) < 3:
            pass
        else:
            i.join(j)
            trans_text_ = translate(
                i,
                lang=lang,
                lang_src=lang_src,
                method=method,
                service_urls=service_urls,
                verbose=verbose,
            )
            trans_text.append(trans_text_)
            # if verbose:
            #     print(trans_text_)
    if trans_text:
        return trans_text
    else:
        return None


def trans_pdf(
    filename,
    page="all",
    lang="english",
    lang_src=None,
    method="google",
    service_urls=["translate.google.de"],
    verbose=False,
):
    """load the pdf file and translated it into target lang "lang",
    verbose: (default 'False', no display)to display the translated text in for loop
    Return (list):
        the translated text as a list
    """
    txt = load_pdf(filename, page=page, verbose=verbose)
    trans_text = translate(
        txt,
        lang=lang,
        lang_src=lang_src,
        method=method,
        service_urls=service_urls,
        verbose=False,
    )
    return trans_text


def save_content(fpath, content):
    """
    Save content to a file.

    Parameters:
        fpath (str): The file path where content will be saved.
        content (str): The content to be saved.

    Returns:
        None
    """
    with open(fpath, "w") as file:
        file.write(content)


def save_file(fpath, content, kind=None, font_name="Arial", font_size=10, spacing=6):
    """
    Save content into a file with specified file type and formatting.

    Parameters:
        fpath (str): The file path where content will be saved.
        content (list of str): The content to be saved, where each string represents a paragraph.
        kind (str): The file type to save. Supported options: 'docx', 'txt', 'md', 'html', 'pdf'.
        font_name (str): The font name for text formatting (only applicable for 'docx', 'html', and 'pdf').
        font_size (int): The font size for text formatting (only applicable for 'docx', 'html', and 'pdf').

    Returns:
        None
    """
    file_types = [".docx", ".txt", ".md", ".html", ".pdf"]
    if kind is None:
        # Extract the file extension from fpath
        _, kind = os.path.splitext(fpath)
        kind = kind.lower()  # Convert extension to lowercase for comparison
    # Check if kind is valid
    if kind.lower() not in file_types:
        raise ValueError(f"Error:\n{kind} is not in the supported list {file_types}")
    if "docx" in kind.lower():
        if isinstance(content, str):
            content = split_text(content, "sentence")
        doc = docx.Document()
        for i, paragraph_text in enumerate(content):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(paragraph_text)
            font = run.font
            font.name = font_name
            font.size = docx.shared.Pt(font_size)
            if i != len(content) - 1:  # Add spacing for all but the last paragraph
                paragraph.space_after = docx.shared.Pt(spacing)
        doc.save(fpath)
    elif "txt" in kind.lower():
        save_content(fpath, "\n".join(content))
    elif "md" in kind.lower():
        save_content(fpath, "\n\n".join(content))
    elif "html" in kind.lower():
        html_content = "<html><body>"
        for paragraph_text in content:
            html_content += f'<p style="font-family:{font_name}; font-size:{font_size}px;">{paragraph_text}</p>'
        html_content += "</body></html>"
        save_content(fpath, html_content)
    elif "pdf" in kind.lower():
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font(font_name, size=font_size)
        for paragraph_text in content:
            pdf.cell(200, 10, txt=paragraph_text, ln=True)
        # Output PDF content as bytes
        pdf_bytes = pdf.output(dest="S").encode(
            "utf-8"
        )  # Encode PDF content to bytes using latin-1
        with open(fpath, "wb") as file:
            file.write(pdf_bytes)
    else:
        raise ValueError(f"Error:\n{kind} is not in the supported list {file_types}")


# if __name__ == "__main__":
#     text_to_translate = "Hello, how are you?"
#     lang = "chinese"
#     translated_text = translate(text_to_translate, lang=lang)
#     print(f"Detected language:{detected_language} \ntranslated into {lang}")
#     print("Translated text:\n", translated_text)
